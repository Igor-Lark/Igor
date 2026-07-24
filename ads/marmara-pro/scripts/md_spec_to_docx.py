#!/usr/bin/env python3
"""Convert Marmara/KlinkerPro Direct spec Markdown to readable Word (.docx)."""

from __future__ import annotations

import csv
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent


def strip_md(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("`", "")
    return text


def add_rich_run(paragraph, text: str, *, bold_default=False, code=False):
    """Add runs parsing **bold** segments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold_default
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def add_rich_paragraph(doc, text: str, style: str = "Normal", *, italic=False, code=False):
    text = strip_md(text)
    p = doc.add_paragraph(style=style)
    if italic:
        p.paragraph_format.left_indent = Inches(0.25)
    add_rich_run(p, text, code=code)
    if italic:
        for run in p.runs:
            run.italic = True
    return p


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [strip_md(c.strip()) for c in line.split("|")]


def is_table_sep(line: str) -> bool:
    s = line.strip().replace("|", "").replace(":", "").replace("-", "").strip()
    return s == "" and "|" in line and "-" in line


def set_table_header(table):
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_table(doc, rows: list[list[str]]):
    if len(rows) < 1:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            val = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_run(p, val)
    set_table_header(table)
    doc.add_paragraph()


def load_ads_from_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f, delimiter=";")
        return list(reader)


def add_ads_block(doc, ads: list[dict[str, str]]):
    doc.add_heading("Тексты объявлений (импорт CSV)", level=2)
    doc.add_paragraph(
        "Без указания цен. Лимиты: заголовок 1 — 56 символов, заголовок 2 — 30, текст — 81."
    )
    fields = [
        ("Заголовок 1", "Заголовок 1"),
        ("Заголовок 2", "Заголовок 2"),
        ("Текст", "Текст"),
        ("Ссылка", "Ссылка"),
    ]
    for i, ad in enumerate(ads, 1):
        doc.add_heading(f"Объявление {i}", level=3)
        rows = [["Поле", "Текст", "Символов"]]
        for label, key in fields:
            val = ad.get(key, "")
            rows.append([label, val, str(len(val)) if val else ""])
        dlink = ad.get("Отображаемая ссылка 1", "") + "/" + ad.get("Отображаемая ссылка 2", "")
        dlink = dlink.strip("/")
        if dlink:
            rows.append(["Отображаемая ссылка", dlink, ""])
        add_table(doc, rows)


def md_to_docx(
    md_path: Path,
    docx_path: Path,
    *,
    client_line: str,
    skip_line_patterns: tuple[str, ...] = (),
    extra_ads_csv: Path | None = None,
):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    section_num = 0
    i = 0
    title_done = False
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if any(stripped.startswith(p) for p in skip_line_patterns):
            i += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not title_done:
            doc.add_paragraph(stripped[2:].strip(), style="Title")
            add_rich_paragraph(doc, client_line, style="Normal")
            doc.add_paragraph()
            title_done = True
            i += 1
            continue

        if stripped.startswith("## "):
            section_num += 1
            heading = stripped[3:].strip()
            doc.add_heading(f"{section_num}. {heading}", level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if stripped.startswith("> "):
            add_rich_paragraph(doc, stripped[2:], italic=True)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_run(p, stripped[2:])
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_rich_run(p, m.group(2))
            i += 1
            while i < len(lines) and lines[i].startswith("   - "):
                sub = doc.add_paragraph(style="List Bullet 2")
                add_rich_run(sub, lines[i].strip()[2:].lstrip("- ").strip())
                i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(strip_md(stripped.strip("*")))
            run.italic = True
            i += 1
            continue

        if stripped:
            add_rich_paragraph(doc, stripped)
        i += 1

    if extra_ads_csv:
        ads = load_ads_from_csv(extra_ads_csv)
        if ads:
            doc.add_page_break()
            doc.add_heading(f"{section_num + 1}. Приложение: объявления из CSV", level=1)
            add_ads_block(doc, ads)

    doc.save(docx_path)


def main():
    name = sys.argv[1] if len(sys.argv) > 1 else "campaign-cpa"
    specs = {
        "campaign-cpa": (
            ROOT / "termopaneli-priozersky-campaign-cpa.md",
            ROOT / "termopaneli-priozersky-campaign-cpa.docx",
            "Клиент: КлинкерПрофи · ИП Баушев Дмитрий Викторович · г. Выборг",
            ("**Word:**",),
            ROOT / "termopaneli-priozersky-campaign-cpa-ads.csv",
        ),
        "adgroup": (
            ROOT / "termopaneli-priozersky-adgroup.md",
            ROOT / "termopaneli-priozersky-adgroup.docx",
            "Клиент: КлинкерПрофи · ИП Баушев Дмитрий Викторович · г. Выборг",
            (),
            None,
        ),
    }
    if name == "all":
        targets = ["campaign-cpa", "adgroup"]
    elif name in specs:
        targets = [name]
    else:
        print("Usage: md_spec_to_docx.py [campaign-cpa|adgroup|all]")
        sys.exit(1)

    for key in targets:
        md, docx, client, skip, ads_csv = specs[key]
        md_to_docx(md, docx, client_line=client, skip_line_patterns=skip, extra_ads_csv=ads_csv)
        print("Wrote", docx)


if __name__ == "__main__":
    main()
