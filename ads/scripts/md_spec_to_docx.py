#!/usr/bin/env python3
"""Convert advertising campaign spec Markdown to readable Word (.docx).

Used for any Yandex Direct / ad campaign in this repo. See ads/README.md.
"""

from __future__ import annotations

import argparse
import csv
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

DEFAULT_SKIP_PREFIXES = ("**Word:**", "**DOCX:**")


def strip_md(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("`", "")
    return text


def add_rich_run(paragraph, text: str, *, bold_default=False, code=False):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold_default
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def add_rich_paragraph(doc, text: str, style: str = "Normal", *, italic=False, code=False):
    text = strip_md(text)
    p = doc.add_paragraph(style=style)
    if italic:
        p.paragraph_format.left_indent = Inches(0.25)
    add_rich_run(p, text, code=code)
    if italic:
        for run in p.runs:
            run.italic = True
    return p


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [strip_md(c.strip()) for c in line.split("|")]


def is_table_sep(line: str) -> bool:
    s = line.strip().replace("|", "").replace(":", "").replace("-", "").strip()
    return s == "" and "|" in line and "-" in line


def set_table_header(table):
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            val = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_run(p, val)
    set_table_header(table)
    doc.add_paragraph()


def load_ads_from_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f, delimiter=";")
        return list(reader)


def add_ads_block(doc, ads: list[dict[str, str]]):
    doc.add_heading("Тексты объявлений (импорт CSV)", level=2)
    doc.add_paragraph(
        "Без указания цен. Лимиты Яндекс Директа: заголовок 1 — 56 символов, "
        "заголовок 2 — 30, текст — 81."
    )
    fields = [
        ("Заголовок 1", "Заголовок 1"),
        ("Заголовок 2", "Заголовок 2"),
        ("Текст", "Текст"),
        ("Ссылка", "Ссылка"),
    ]
    for i, ad in enumerate(ads, 1):
        doc.add_heading(f"Объявление {i}", level=3)
        rows = [["Поле", "Текст", "Символов"]]
        for label, key in fields:
            val = ad.get(key, "")
            rows.append([label, val, str(len(val)) if val else ""])
        d1 = ad.get("Отображаемая ссылка 1", "")
        d2 = ad.get("Отображаемая ссылка 2", "")
        dlink = f"{d1}/{d2}".strip("/") if (d1 or d2) else ""
        if dlink:
            rows.append(["Отображаемая ссылка", dlink, ""])
        add_table(doc, rows)


def extract_client_line(lines: list[str]) -> str:
    for line in reversed(lines):
        m = re.match(r"^\*Клиент:\s*(.+)\*\s*$", line.strip())
        if m:
            return f"Клиент: {m.group(1)}"
    return ""


def md_to_docx(
    md_path: Path,
    docx_path: Path,
    *,
    client_line: str | None = None,
    skip_line_patterns: tuple[str, ...] = DEFAULT_SKIP_PREFIXES,
    extra_ads_csv: Path | None = None,
    ads_appendix_title: str = "Приложение: объявления из CSV",
):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    client = client_line if client_line is not None else extract_client_line(lines)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    section_num = 0
    i = 0
    title_done = False
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if any(stripped.startswith(p) for p in skip_line_patterns):
            i += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not title_done:
            doc.add_paragraph(stripped[2:].strip(), style="Title")
            if client:
                add_rich_paragraph(doc, client, style="Normal")
                doc.add_paragraph()
            title_done = True
            i += 1
            continue

        if stripped.startswith("## "):
            section_num += 1
            doc.add_heading(f"{section_num}. {stripped[3:].strip()}", level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if stripped.startswith("> "):
            add_rich_paragraph(doc, stripped[2:], italic=True)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_run(p, stripped[2:])
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_rich_run(p, m.group(2))
            i += 1
            while i < len(lines) and lines[i].startswith("   - "):
                sub = doc.add_paragraph(style="List Bullet 2")
                add_rich_run(sub, lines[i].strip()[2:].lstrip("- ").strip())
                i += 1
            continue

        if (
            stripped.startswith("*")
            and stripped.endswith("*")
            and not stripped.startswith("**")
        ):
            p = doc.add_paragraph()
            run = p.add_run(strip_md(stripped.strip("*")))
            run.italic = True
            i += 1
            continue

        if stripped:
            add_rich_paragraph(doc, stripped)
        i += 1

    if extra_ads_csv and extra_ads_csv.is_file():
        ads = load_ads_from_csv(extra_ads_csv)
        if ads:
            doc.add_page_break()
            doc.add_heading(f"{section_num + 1}. {ads_appendix_title}", level=1)
            add_ads_block(doc, ads)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Build readable .docx from ad campaign Markdown.")
    parser.add_argument("markdown", type=Path, help="Path to spec .md")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output .docx (default: same name as .md)",
    )
    parser.add_argument(
        "--ads",
        type=Path,
        default=None,
        help="Optional ads CSV (semicolon-separated) for appendix tables",
    )
    parser.add_argument(
        "--client",
        type=str,
        default=None,
        help='Override client line under title (default: parse *Клиент: …* from md)',
    )
    parser.add_argument(
        "--skip",
        action="append",
        default=[],
        help="Skip md lines starting with this prefix (repeatable)",
    )
    args = parser.parse_args(argv)

    md_path = args.markdown.resolve()
    if not md_path.is_file():
        print(f"Not found: {md_path}", file=sys.stderr)
        return 1

    docx_path = args.output or md_path.with_suffix(".docx")
    skip = tuple(DEFAULT_SKIP_PREFIXES + tuple(args.skip))

    md_to_docx(
        md_path,
        docx_path,
        client_line=args.client,
        skip_line_patterns=skip,
        extra_ads_csv=args.ads,
    )
    print(docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
