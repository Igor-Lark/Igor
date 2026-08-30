#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Пересечения ключей четырёх групп Выборга — Word."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path("/workspace/reports/yandex-direct/2026-08-30_Выборг_ключи_пересечения.docx")
INBOX = Path("/workspace/inbox/Analiz_RK_Vyborg_klyuchi_peresecheniya_2026-08-30.docx")
CURSOR = Path("/workspace/cursor/Analiz_RK_Vyborg_klyuchi_peresecheniya_2026-08-30.docx")
GH = (
    "https://github.com/Igor-Lark/Igor/blob/"
    "cursor/yandex-rk-audit-aug16-bfbc/"
    "reports/yandex-direct/2026-08-30_%D0%92%D1%8B%D0%B1%D0%BE%D1%80%D0%B3_"
    "%D0%BA%D0%BB%D1%8E%D1%87%D0%B8_%D0%BF%D0%B5%D1%80%D0%B5%D1%81%D0%B5%D1%87%D0%B5%D0%BD%D0%B8%D1%8F.docx"
)


def shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def p(doc, text, size=11, bold=False, space_after=6, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = para.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return para


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    return para


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        r = para.add_run(htxt)
        set_run(r, size=9, bold=True, color=(255, 255, 255))
        shade(cell, "1A3A5C")
    for ri, row in enumerate(rows):
        fill = "F4F7FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            r = para.add_run(str(val))
            set_run(r, size=8)
            shade(cell, fill)
    doc.add_paragraph()
    return t


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = para.add_run(text)
    set_run(r, size=11)
    return para


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.4)

    p(doc, "КлинкерПрофи  ·  РК «Термопанели в Выборге» 712773255", size=10, color=(0x5A, 0x6A, 0x7A), space_after=2)
    p(doc, "Где пересекаются ключи четырёх групп", size=18, bold=True, space_after=4, color=(0x1A, 0x3A, 0x5C))
    p(
        doc,
        "Списки ключей из Директа, 30.08.2026. Сверено по тексту (регистр, ё/е, кавычки). "
        "Стратегию и CPA не трогаем. СПб-фразы в группе ЛО не удаляем. Транз не включаем. "
        "Фразы выключать, не удалять.",
        size=10,
        space_after=4,
    )
    p(doc, "Файл Word на GitHub:", size=10, space_after=0)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    r = para.add_run(GH)
    set_run(r, size=8, color=(0x0B, 0x5C, 0xA8))

    h(doc, "Коротко", 1)
    p(
        doc,
        "Точный дубль между группами один: «термопанели выборг» висит и в Гео, и в клинкере. "
        "Остальное — один и тот же корень с разными минусами, плюс ключи Выборга внутри группы ЛО "
        "(они отбирают трафик у Гео). Внутри клинкера 20 фраз с минусами в ключе — они и сами себя режут, "
        "и пересекаются с коммерческой. Это как раз те фразы, которые ещё 24.08 надо было выключить.",
        space_after=6,
    )
    table(
        doc,
        ["Группа", "№", "Ключей", "С минусом внутри"],
        [
            ["ЛО", "5773097196", "17", "2"],
            ["Гео", "5773898865", "25", "1"],
            ["Коммерческая", "5773901452", "27", "3"],
            ["Клинкер", "5773902142", "50", "20"],
            ["Всего", "—", "119", "26"],
        ],
    )

    h(doc, "1. Точный дубль", 1)
    table(
        doc,
        ["Фраза", "Где висит", "Что сделать"],
        [
            ['"термопанели выборг"', "Гео и Клинкер", "Оставить в Гео. В клинкере выключить."],
        ],
    )
    p(
        doc,
        "Это единственное полное совпадение текста. В клинкерной группе гео-ключу не место: "
        "он бьётся с группой Гео за один запрос, который в статистике 23–30.08 всё равно не всплывал.",
        space_after=8,
    )

    h(doc, "2. Один корень — две группы", 1)
    table(
        doc,
        ["Корень", "Группа А", "Группа Б", "Что сделать"],
        [
            [
                "термопанели ленобласть",
                "Гео: без минусов",
                "ЛО: −установка −доставка −купить",
                "В ЛО минус-хвост выключить. Корень оставить в ЛО (доставка / купить в Ленобласти уже есть). В Гео «термопанели ленобласть» выключить — это территория ЛО.",
            ],
            [
                "клинкерные термопанели от производителя",
                "Клинкер: чисто",
                "Коммерч.: −плитка",
                "Оставить в клинкере. Коммерческую с минусом выключить.",
            ],
            [
                "купить клинкерные термопанели",
                "Коммерч.: чисто",
                "Клинкер: −фасад −плитка",
                "Оставить в коммерческой. Клинкерную с минусом выключить.",
            ],
        ],
    )

    h(doc, "3. Выборг внутри группы ЛО — отбирает у Гео", 1)
    p(
        doc,
        "Группа ЛО нужна для СПб и области. Ключи с «Выборг» в ней дублируют Гео. "
        "СПб-фразы (Питер, СПб, санкт-петербург) не трогать: 29.08 с «купить в санкт петербурге» уже был клик.",
        space_after=4,
    )
    table(
        doc,
        ["Фраза в ЛО", "С чем пересекается в Гео", "Действие"],
        [
            ["Купить термопанели в Выборге", '"купить термопанели выборг", «выборг купить»', "Выключить в ЛО"],
            ["Термопанели с клинкерной плиткой в Выборге", "клинкерные термопанели выборг", "Выключить в ЛО"],
            ["термопанели доставка Выборг", "весь кластер «выборг» в Гео", "Выключить в ЛО"],
            ["термопанели Ленинградская область Выборг", "термопанели ленинградская область… + выборг", "Выключить в ЛО"],
            [
                "Термопанели с клинкерной плиткой в −питер −выборг −спб",
                "клинкер + гео, плюс минусы внутри ключа",
                "Выключить (и минусы, и бессмысленный хвост)",
            ],
        ],
    )
    p(doc, "В ЛО оставить:", size=11, bold=True, space_after=3)
    bullet(doc, "фасадные термопанели Питер; доставка Спб / Ленобласть; купить в Ленобласти;")
    bullet(doc, "изготовители / купить в санкт петербурге; с клинкерной плиткой в Питере / купить в спб;")
    bullet(doc, "Санкт-Петербург и область; с установкой в ленобласти; всеволжский район (опечатка: всеволожский — не переименовывать пачкой).")
    p(doc, "", space_after=2)

    h(doc, "4. Гео vs ЛО по области и чужая РК", 1)
    table(
        doc,
        ["Фраза в Гео", "Проблема", "Действие"],
        [
            ["термопанели ленобласть", "Тот же корень, что в ЛО", "Выключить в Гео"],
            ["фасадные термопанели ленинградская область", "Широкая ЛО, группа Гео про Выборг", "Выключить в Гео"],
            ["термопанели ленинградская область −фасадный", "Минус внутри + ЛО", "Выключить"],
            ["термопанели всеволожск", "Рядом с «всеволжский район» в ЛО", "Оставить в Гео или выключить — не дублировать в ЛО"],
            ["купить термопанели в гатчине", "Гатчина — ЛО, не Выборг", "Можно оставить (другой город), не копировать в ЛО"],
            ["термопанели приозерск", "Чужая кампания 713047408", "Выключить в Выборге"],
        ],
    )
    p(
        doc,
        "Каменногорск, Рощино, Светогорск, Выборгский район, Сосновый бор, Тихвин, Луга, Кингисепп — "
        "в Гео пересечений с другими группами этой РК нет. Сосновый Бор ≠ Сосново (Приозерский район). "
        "«термопанели приозерск» в Выборге не копировать в Приозерск повторно: там гео-фразы уже добавлены 24.08.",
        space_after=8,
    )

    h(doc, "5. Коммерческая vs клинкер (рядом, не точный дубль)", 1)
    table(
        doc,
        ["Коммерческая", "Клинкер", "Кто владеет темой"],
        [
            ["цена термопанелей фасадных под кирпич", "цена термопанелей под кирпич", "Клинкер. Коммерческую «под кирпич» выключить."],
            ["термопанели с плиткой купить", "термопанели с клинкерной плиткой купить", "Клинкер — про плитку. Коммерческую можно выключить."],
            ["термопанели с клинкерной плиткой от производителя", '"термопанели с клинкерной плиткой"', "Клинкер. Коммерческую выключить."],
            ["купить термопанели для фасада", "клинкерные термопанели для фасада купить", "Разные: коммерч. без «клинкер». Обе оставить."],
            ["фасадные панели с утеплением", "панели с утеплителем / с клинкером", "Разные. Обе оставить."],
        ],
    )
    p(
        doc,
        "В коммерческой самой с собой: «расчет термопанелей» и «расчет термопанелей» в кавычках — "
        "фраза и точное вхождение. Не дубль групп, можно оставить одну из двух, если мешает.",
        space_after=8,
    )

    h(doc, "6. Минусы внутри ключа — выключить все 26", 1)
    p(
        doc,
        "Пункт 1.6 от 24.08: фразы с −купить −цена −дом и т.п. внутри ключа выключены. В списках они всё ещё есть. "
        "Минусы в ключе режут показы и плодят «почти те же» фразы. Минус-слова — на уровне кампании, не в ключе.",
        space_after=4,
    )
    p(doc, "ЛО (2)", size=11, bold=True, space_after=2)
    bullet(doc, "термопанели ленобласть −установка −доставка −купить")
    bullet(doc, "Термопанели с клинкерной плиткой в −питер −выборг −спб")
    p(doc, "Гео (1)", size=11, bold=True, space_after=2)
    bullet(doc, "термопанели ленинградская область −фасадный")
    p(doc, "Коммерческая (3)", size=11, bold=True, space_after=2)
    bullet(doc, "термопанель кирпич цена фасадная −под")
    bullet(doc, "термопанели с плиткой от производителя −клинкерный")
    bullet(doc, "клинкерные термопанели от производителя −плитка")
    p(doc, "Клинкер (20) — все с «−» после фразы", size=11, bold=True, space_after=2)
    bullet(doc, "в т.ч. «клинкерные термопанели для фасада −купить −цена» (уже светилась в статистике 29.08);")
    bullet(doc, "хвосты −отделка −дом −фасад −плитка −наружный −м2 −утеплитель — выключить пачкой, чистые без минусов оставить.")
    p(doc, "", space_after=2)

    h(doc, "7. Что выключить сегодня (чеклист)", 1)
    p(doc, "ЛО 5773097196 — 6 фраз", size=12, bold=True, space_after=3)
    bullet(doc, "Купить термопанели в Выборге")
    bullet(doc, "Термопанели с клинкерной плиткой в Выборге")
    bullet(doc, "термопанели доставка Выборг")
    bullet(doc, "термопанели Ленинградская область Выборг")
    bullet(doc, "термопанели ленобласть −установка −доставка −купить")
    bullet(doc, "Термопанели с клинкерной плиткой в −питер −выборг −спб")
    p(doc, "Гео 5773898865 — 4 фразы", size=12, bold=True, space_after=3)
    bullet(doc, "термопанели ленобласть")
    bullet(doc, "фасадные термопанели ленинградская область")
    bullet(doc, "термопанели ленинградская область −фасадный")
    bullet(doc, "термопанели приозерск")
    p(doc, "Коммерческая 5773901452 — 6 фраз", size=12, bold=True, space_after=3)
    bullet(doc, "термопанель кирпич цена фасадная −под")
    bullet(doc, "термопанели с плиткой от производителя −клинкерный")
    bullet(doc, "клинкерные термопанели от производителя −плитка")
    bullet(doc, "цена термопанелей фасадных под кирпич")
    bullet(doc, "термопанели с плиткой купить")
    bullet(doc, "термопанели с клинкерной плиткой от производителя")
    p(doc, "Клинкер 5773902142", size=12, bold=True, space_after=3)
    bullet(doc, '"термопанели выборг"')
    bullet(doc, "все 20 с минусом внутри ключа")
    p(
        doc,
        "Не удалять и не переносить пачками. Не трогать СПб в ЛО. Не включать транз. "
        "Кавычки у «термопанели выборг» в Гео оставить. Новые ключи не добавлять на этой неделе.",
        space_after=8,
    )

    h(doc, "8. Что не делать", 1)
    bullet(doc, "не удалять СПб-фразы из ЛО;")
    bullet(doc, "не включать группу «транз»;")
    bullet(doc, "не вставлять «термопанели приозерск» в кампанию Приозерска повторно;")
    bullet(doc, "не чистить клинкер до нуля — колорадо, амстердам, чистые «под кирпич» и «с клинкерной плиткой» оставляем;")
    bullet(doc, "не менять стратегию и CPA.")
    p(doc, "", space_after=4)
    p(
        doc,
        "Источник: списки ключей из Директа, 30.08.2026. Связанный разбор статистики: "
        "inbox/Analiz_RK_Vyborg_Priozersk_dni_2026-08-29_30.docx.",
        size=9,
        color=(0x5A, 0x6A, 0x7A),
        space_after=0,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    INBOX.parent.mkdir(parents=True, exist_ok=True)
    CURSOR.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    doc.save(INBOX)
    doc.save(CURSOR)
    print(OUT)
    print(INBOX)
    print(CURSOR)


if __name__ == "__main__":
    main()
