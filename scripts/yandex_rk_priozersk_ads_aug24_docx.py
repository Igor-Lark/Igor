#!/usr/bin/env python3
"""Priozersk: geo keywords + new titles/texts. Print-ready."""

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path("reports/yandex-direct/2026-08-24_Приозерск_ключи_заголовки.docx")
INBOX = Path("inbox/Priozersk_klyuchi_zagolovki_2026-08-24.docx")
CURSOR = Path("cursor/Priozersk_klyuchi_zagolovki_2026-08-24.docx")

KEYS = [
    "термопанели приозерск",
    "купить термопанели приозерск",
    "термопанели приозерск цена",
    "клинкерные термопанели приозерск",
    "фасадные термопанели приозерск",
    "термопанели приозерский район",
    "доставка термопанелей приозерск",
    "утепление фасада приозерск термопанели",
    "термопанели сосново",
    "купить термопанели сосново",
    "термопанели ларионово",
    "термопанели лосево",
    "термопанели громово",
    "термопанели мельниково",
    "термопанели саперное",
    "термопанели кузнечное",
    "клинкерная плитка приозерск",
    "купить клинкерную плитку приозерск",
]

H1 = [
    "Термопанели для дома в Приозерске",
    "Термопанели в Сосново. Завод в Выборге",
    "Клинкерные термопанели, Приозерский р-н",
    "Доставка в Приозерск с завода в Выборге",
    "Термопанели Ларионово, Лосево, Громово",
    "Фасадные термопанели от 1 550 ₽/панель",
    "Утепление фасада термопанелями, Приозерск",
    "Термопанели и клинкерная плитка, Приозерск",
    "Купить термопанели в Приозерском районе",
]

TEXTS = [
    "Завод в Выборге. Доставка в Приозерск, Сосново, Ларионово. От 1 550 ₽ за панель.",
    "Клинкерные термопанели с утеплителем. Свой завод. Расчёт фасада за 1 день.",
    "Производство КлинкерПрофи, не стройбаза. Каталог и телефон — на сайте.",
]


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_p(doc, text, size=11, bold=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.12
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def paste_block(doc, lines, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run("\n".join(lines))
    set_run(run, size=size)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F7FB")
    shd.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shd)


def setup_print(doc):
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.left_margin = Cm(1.6)
    s.right_margin = Cm(1.6)
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.6)
    hp = s.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Приозерск РК 713047408  ·  24.08.2026")
    set_run(run, size=9, color=(89, 89, 89))


def build():
    doc = Document()
    setup_print(doc)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_p(
        doc,
        "Приозерск: ключи с гео и новые объявления",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_p(
        doc,
        "Пункт 2.7. Посадка — главная. Плитка в каталоге остаётся. Стратегию не трогать.",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    heading(doc, "Ключи с гео — вставить списком, тип «фраза»", 1)
    add_p(
        doc,
        "В группу «Приозерский район». Выборг сюда не класть. Не дублировать, если фраза уже есть.",
        size=10,
        space_after=4,
    )
    paste_block(doc, KEYS)

    heading(doc, "Заголовки (9 шт., лимит 56 знаков)", 1)
    add_p(doc, "В комбинатор: добавить к текущим. Слабый «Доставка в район» без города можно не удалять — Директ сам смешает.", size=10)
    paste_block(
        doc,
        [f"{t}   ({len(t)})" for t in H1],
    )

    heading(doc, "Тексты (3 шт., лимит 81 знак)", 1)
    paste_block(doc, [f"{t}   ({len(t)})" for t in TEXTS])

    heading(doc, "Как есть, не менять", 1)
    add_p(doc, "Ссылка объявлений — главная marmara-pro.ru. Карточки клинкерной плитки оставить: плитку спрашивают. Архив не трогать. Автотаргетинг: целевые, узкие, сопутствующие, свой бренд, без бренда. Бюджет 8 000 ₽. CPA 400/350.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    INBOX.parent.mkdir(parents=True, exist_ok=True)
    CURSOR.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    shutil.copy2(OUT, INBOX)
    shutil.copy2(OUT, CURSOR)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
