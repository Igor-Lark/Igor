#!/usr/bin/env python3
"""Print-ready: strategy lock + keywords for existing Vyborg groups."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path("reports/yandex-direct/2026-08-17_ключи_группы_Выборг.docx")
INBOX = Path("inbox/Klyuchi_gruppy_Termopaneli_Vyborg_2026-08-17.docx")
CURSOR = Path("cursor/Klyuchi_gruppy_Termopaneli_Vyborg_2026-08-17.docx")


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_p(doc, text, size=11, bold=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.12
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def fill_cell(cell, text, bold=False, size=10, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold or header, color=(255, 255, 255) if header else None)
    set_cell_border(cell)
    if header:
        shade_cell(cell, "1F4E79")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, header=True, size=10)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            fill_cell(table.rows[r_i + 1].cells[c_i], val, size=10, bold=c_i == 0)
            if r_i % 2 == 1:
                shade_cell(table.rows[r_i + 1].cells[c_i], "F2F2F2")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run(run, size=size)


def paste_block(doc, lines, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("\n".join(lines))
    set_run(run, size=size)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F7FB")
    shading.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shading)


def setup_print(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("КлинкерПрофи  ·  marmara-pro.ru  ·  17 августа 2026")
    set_run(run, size=9, color=(89, 89, 89))

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Стратегия + ключи групп «Термопанели в Выборге»  ·  стр. ")
    set_run(run, size=9, color=(89, 89, 89))
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r2 = fp.add_run()
    r2._r.append(fld1)
    r2._r.append(instr)
    r2._r.append(fld2)
    set_run(r2, size=9, color=(89, 89, 89))


def group_block(doc, title, meta, now, remove, keep, add, note=""):
    heading(doc, title, 2)
    add_p(doc, meta, size=10, space_after=4)
    add_p(doc, "Сейчас в группе", bold=True, size=11, space_after=2)
    for line in now:
        bullet(doc, line, size=10)
    add_p(doc, "УДАЛИТЬ", bold=True, size=11, space_after=2, space_before=4)
    if remove:
        for line in remove:
            bullet(doc, line, size=10)
    else:
        bullet(doc, "Ничего. Ключей почти нет — только добрать.", size=10)
    if keep:
        add_p(doc, "ОСТАВИТЬ", bold=True, size=11, space_after=2, space_before=4)
        for line in keep:
            bullet(doc, line, size=10)
    add_p(
        doc,
        "ДОБАВИТЬ — вставить списком, тип соответствия «фраза»",
        bold=True,
        size=11,
        space_after=3,
        space_before=4,
    )
    paste_block(doc, add)
    if note:
        add_p(doc, note, size=10, space_after=8)


def build():
    doc = Document()
    setup_print(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_p(
        doc,
        "Стратегия на все РК + ключи групп Выборга",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_p(
        doc,
        "Рабочий лист. Минус-слова уже вставлены, мусорные площадки в основном убраны — не трогать.",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    heading(doc, "1. Одинаково на все три кампании", 1)
    add_p(
        doc,
        "Термопанели в Выборге (712773255), Приозерский район Поиск (713047408) "
        "и третья РК аккаунта — те же поля. Не пакетная, каждая сама по себе.",
        space_after=6,
    )
    add_table(
        doc,
        ["Поле", "Значение"],
        [
            ["Стратегия", "Максимум конверсий"],
            ["Оплата", "За конверсии"],
            ["Недельный бюджет", "9 000 ₽"],
            ["Ограничение", "Цена конверсии"],
            ["Цель «Телефон» 582134942", "400 ₽"],
            ["Цель «Мессенджер» 582134994", "350 ₽"],
            ["Счётчик Метрики", "marmara-pro.ru · 110643713"],
            ["РСЯ: Connected TV / Smart TV", "Выключен"],
            ["Минус-слова кампании", "Уже стоят — не чистить и не дублировать"],
            ["Площадки РСЯ", "Мусор в основном убран — не открывать заново"],
        ],
    )
    add_p(
        doc,
        "400 / 350 ₽ лучше прежних 220–290 ₽. При оплате за конверсии Директ "
        "не спишет больше цели, но если за 7 дней на новых ключах показы около нуля — "
        "поднять только телефон до 700 ₽, остальное не трогать. После смены цифр — "
        "неделю не менять стратегию (иначе обучение сбрасывается).",
        space_after=6,
    )
    add_p(doc, "Ещё два правила, не стратегия:", bold=True, space_after=3)
    bullet(doc, "Автотаргетинг на Поиске — выключить в каждой группе. Иначе ключи снова не крутятся.")
    bullet(doc, "Ключи не дублировать между группами: гео — только с городом, коммерция — без города.")

    heading(doc, "2. Группы РК «Термопанели в Выборге»", 1)
    add_p(
        doc,
        "Пять существующих групп. Новые группы не создавать. "
        "За неделю 10–16.08 фразы почти не работали: автотаргетинг съел 94% показов. "
        "Гео-группа — 0 ключей, одна конверсия как раз с неё (РСЯ, автотаргетинг).",
        space_after=6,
    )
    add_table(
        doc,
        ["Группа", "№", "Показы", "Клики", "Фразы сейчас"],
        [
            ["Группа 1 — коммерческая", "5773901452", "240", "14", "5 фраз, 2 из них с минусами внутри"],
            ["Группа 2 — Гео (ключевая)", "5773898865", "200", "10", "0 фраз, только автотаргетинг"],
            ["Группа 3 — клинкер (продуктовая)", "5773902142", "258", "5", "1 фраза"],
            ["Термопанели (транз)", "5772834263", "453", "8", "4 фразы, уклон в СПб"],
            ["Термопанели (Лен. область)", "5773097196", "231", "5", "2 фразы, обе про СПб — не ЛО"],
        ],
    )

    group_block(
        doc,
        "Группа 1 — коммерческая  ·  5773901452",
        "Смысл: купить / цена / расчёт. Без городов (города — в группах 2, транз, ЛО).",
        now=[
            "купить клинкерные термопанели — 2 показа, оставить",
            "купить термопанели для дома — 1 показ, оставить",
            "термопанели -цена -производитель -купить -заказывать -доставка — 57 показов: минусы внутри ключа режут коммерцию",
            "фасадные термопанели цена -дом -кирпич — минус «дом» убивает целевой запрос",
            "фасадные панели для наружной отделки — широко, без «термо»",
            "автотаргетинг — 173 показа / 11 кликов",
        ],
        remove=[
            "термопанели -цена -производитель -купить -заказывать -доставка",
            "фасадные термопанели цена -дом -кирпич",
            "фасадные панели для наружной отделки",
        ],
        keep=[
            "купить клинкерные термопанели",
            "купить термопанели для дома",
        ],
        add=[
            "купить термопанели",
            "термопанели купить",
            "термопанели цена",
            "термопанели цена за м2",
            "фасадные термопанели купить",
            "купить фасадные термопанели",
            "термопанели для фасада купить",
            "термопанели с утеплителем купить",
            "термопанели от производителя",
            "сколько стоит обшить дом термопанелями",
            "расчёт термопанелей на дом",
            "обшить дом термопанелями цена",
        ],
    )

    group_block(
        doc,
        "Группа 2 — Гео (ключевая)  ·  5773898865",
        "Смысл: Выборг и Выборгский район. Сейчас ключей нет — это главная дыра. "
        "Единственная конверсия недели (14.08, муж. 45–54, смартфон) пришла из этой группы, но с автотаргетинга РСЯ.",
        now=["автотаргетинг — 200 показов / 10 кликов / 1 конверсия, фраз 0"],
        remove=[],
        keep=[],
        add=[
            "термопанели выборг",
            "купить термопанели выборг",
            "термопанели в выборге",
            "термопанели выборг цена",
            "клинкерные термопанели выборг",
            "фасадные термопанели выборг",
            "термопанели с клинкерной плиткой выборг",
            "производство термопанелей выборг",
            "клинкерпрофи выборг",
            "утепление фасада термопанелями выборг",
            "доставка термопанелей выборг",
            "термопанели выборгский район",
            "термопанели советский",
            "термопанели приморск",
            "термопанели рощино",
            "термопанели высоцк",
            "термопанели каменногорск",
            "термопанели светогорск",
        ],
        note="Регион показа группы — Выборг + Выборгский район, не вся РФ. Приозерск сюда не класть (другая РК).",
    )

    group_block(
        doc,
        "Группа 3 — клинкер (продуктовая)  ·  5773902142",
        "Смысл: вид панели, коллекции, «под кирпич». Без городов и без голого «купить термопанели».",
        now=[
            "термопанели с клинкерной плиткой купить — 1 показ, оставить",
            "автотаргетинг — 257 показов / 5 кликов",
        ],
        remove=[],
        keep=["термопанели с клинкерной плиткой купить"],
        add=[
            "клинкерные термопанели",
            "термопанели с клинкерной плиткой",
            "клинкерные панели для фасада",
            "панели с клинкерной плиткой для фасада",
            "термопанели под кирпич",
            "фасадные термопанели под кирпич",
            "термопанели колорадо",
            "термопанели амстердам",
            "фасадные термопанели с утеплителем",
            "термопанели пенополистирол клинкер",
            "термопанели с пенополистиролом",
        ],
        note="На Поиске у группы было 100 показов и 0 кликов — после ключей поправить заголовок: «Клинкерные термопанели, завод в Выборге, от 1 550 ₽».",
    )

    group_block(
        doc,
        "Термопанели (транз)  ·  5772834263",
        "Смысл: транзакционные запросы из СПб («купить / от производителя / с доставкой»), не Выборг и не Приозерск.",
        now=[
            "купить клинкерные термопанели для фасада в спб — оставить",
            "термопанели с клинкерной плиткой купить в спб — оставить здесь, убрать из группы ЛО",
            "фасадные термопанели для наружной отделки дома цена — оставить",
            "фасадные термопанели для наружной отделки дома -цена — удалить (минус «цена»)",
            "автотаргетинг — 447 показов / 8 кликов",
        ],
        remove=["фасадные термопанели для наружной отделки дома -цена"],
        keep=[
            "купить клинкерные термопанели для фасада в спб",
            "термопанели с клинкерной плиткой купить в спб",
            "фасадные термопанели для наружной отделки дома цена",
        ],
        add=[
            "купить термопанели спб",
            "термопанели купить санкт петербург",
            "клинкерные термопанели спб",
            "фасадные термопанели спб",
            "термопанели от производителя спб",
            "термопанели с доставкой спб",
            "производство термопанелей санкт петербург",
        ],
        note="Гео группы — Санкт-Петербург. Не ставить Выборг (группа 2) и не ставить Приозерск.",
    )

    group_block(
        doc,
        "Термопанели (Лен. область)  ·  5773097196",
        "Смысл: города ЛО кроме Выборгского и Приозерского районов. Сейчас обе фразы про СПб — это не эта группа.",
        now=[
            "термопанели с клинкерной плиткой купить в спб — дубль транз, удалить отсюда",
            "термопанель купить санкт петербург -фасадный — минус «фасадный» вредный, СПб не сюда",
            "автотаргетинг — 228 показов / 5 кликов",
        ],
        remove=[
            "термопанели с клинкерной плиткой купить в спб",
            "термопанель купить санкт петербург -фасадный",
        ],
        keep=[],
        add=[
            "термопанели ленинградская область",
            "термопанели ленобласть",
            "доставка термопанелей ленобласть",
            "термопанели всеволожск",
            "термопанели гатчина",
            "термопанели тосно",
            "термопанели сертолово",
            "термопанели кудрово",
            "термопанели мурино",
            "термопанели кириши",
            "термопанели тихвин",
            "термопанели кировск",
            "термопанели волхов",
            "термопанели коммунар",
            "термопанели сосновый бор",
        ],
        note="Не добавлять: выборг, приозерск, сосново, ларионово, лосево, громово — они в гео-группе или в РК Приозерска.",
    )

    heading(doc, "3. Как вставлять в Директе", 1)
    bullet(doc, "Группа → Ключевые фразы → Добавить. Вставить блок целиком. Тип: фраза.")
    bullet(doc, "Автотаргетинг в этой группе — выкл.")
    bullet(doc, "Минус-слова группы не копировать: они уже на кампании.")
    bullet(doc, "После добавления ключей стратегию и CPA не менять 7 дней.")
    add_p(
        doc,
        "Источник групп: выгрузка 10–16.08.2026, кампания 712773255. "
        "Минус-слова и чистка РСЯ в этом листе не повторяются — сделано.",
        size=9,
        space_after=0,
        space_before=8,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    INBOX.parent.mkdir(parents=True, exist_ok=True)
    CURSOR.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    doc.save(INBOX)
    doc.save(CURSOR)
    print("Wrote", OUT)
    print("Inbox", INBOX)


if __name__ == "__main__":
    build()
