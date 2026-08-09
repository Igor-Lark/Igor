#!/usr/bin/env python3
"""Yandex Direct CSV → Word report (vitaminki-style PPC audit)."""
import re
import sys
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

BASE_MINUS_WORDS = [
    "петрович",
    "вимос",
    "строим наш дом",
    "leroy",
    "леруа",
    "obi",
    "оби",
    "кастorama",
    "maxidom",
    "максидом",
    "dns",
    "mvideo",
    "каталог",
    "официальный сайт",
    "строймагазин",
    "стройбаза",
    "гипермаркет",
    "рассрочка",
    "распродажа",
    "пвх",
    "виниловый сайдинг",
    "сайдинг винил",
    "форум",
    "вакансии",
    "youtube",
    "ютуб",
    "б/у",
    "avito",
    "авито",
    "скачать",
    "технониколь",
    "hauberk",
    "baucenter",
    "анапа",
    "нижний новгород",
    "неводекор",
    "бакс",
    "сокол",
    "ленпанель",
    "мдф",
    "пластик",
]

BASE_MINUS_PHRASES = [
    "петрович выборг",
    "вимос выборг",
    "каталог товаров",
    "строительный магазин",
    "магазин петрович",
    "строим наш дом выборг",
    "термопанели в петровиче",
    "термопанели купить петрович",
]

OPTIONAL_MINUS_IF_NO_DRY_PANELS = [
    "без утеплителя",
    "вентфасад плитка",
]

OFF_QUERY_RE = re.compile(
    r"петрович|вимос|строим наш дом|каталог|рассрочк|распродаж|"
    r"нижн.{0,3}новгород|пвх|официальн|leroy|леруа|магазин строительный|"
    r"строймагазин|youtube|ютуб|форум|технониколь|baucenter|анапа|"
    r"неводекор|бакс|сокол|ленпанель|мдф|пластик",
    re.I,
)
TARGET_QUERY_RE = re.compile(r"термопан|клинкер|фасад.*панел|утепл.*фасад", re.I)

RSYA_BAN_NAME_RE = re.compile(
    r"dzen|\.zen|zen\.|mail\.ru|pogoda|weather|game|games|puzzle|block|"
    r"match3|tiletrip|cleaner|vkontakte|kidult|24smi|word\.|free\.games",
    re.I,
)

KEYWORD_TEMPLATES = {
    "коммерч": [
        "термопанели",
        "термопанели купить",
        "термопанели цена",
        "термопанели для фасада",
        "фасадные панели с утеплением",
        "купить термопанели",
        "термопанели от производителя",
    ],
    "гео": [
        "термопанели выборг",
        "термопанели в выборге",
        "клинкерные термопанели выборг",
        "термопанели выборгский район",
        "фасадные панели выборг",
    ],
    "транз": [
        "термопанели ленинградская область",
        "термопанели ленобласть",
        "термопанели приозерск",
        "термопанели всеволожск",
        "термопанели с доставкой",
    ],
    "клинкер": [
        "клинкерные термопанели",
        "термопанели под кирпич",
        "клинкерные панели для фасада",
        "термопанель кирпич фасад",
    ],
    "лен": [
        "термопанели производство ленобласть",
        "фасадные термопанели ленинградская область",
    ],
}


def num_series(df, col):
    if col not in df.columns:
        return pd.Series([0] * len(df))
    return (
        pd.to_numeric(
            df[col].astype(str).str.replace(",", ".").str.replace(" ", "").replace("-", ""),
            errors="coerce",
        )
        .fillna(0)
    )


def load_csv(path):
    df = pd.read_csv(path, encoding="utf-8-sig")
    for c in ["Расход, ₽", "Показы", "Клики", "Конверсии", "CPA, ₽", "CTR, %", "CPC, ₽"]:
        df[c] = num_series(df, c)
    total = df[df["День"].astype(str) == "Итого"]
    detail = df[df["День"].astype(str) != "Итого"].copy()
    return detail, total.iloc[0] if len(total) else None


def kpi_row(spend, shows, clicks, conv):
    ctr = (clicks / shows * 100) if shows else 0
    cpc = (spend / clicks) if clicks else 0
    cpa = (spend / conv) if conv else 0
    cr = (conv / clicks * 100) if clicks else 0
    return ctr, cpc, cpa, cr


def add_table(doc, headers, rows, col_widths=None):
    if not rows:
        doc.add_paragraph("Нет данных за период.")
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = str(val)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph()


def add_code_block(doc, title, lines):
    if title:
        doc.add_paragraph(title)
    p = doc.add_paragraph("\n".join(str(x) for x in lines))
    for run in p.runs:
        run.font.name = "Consolas"


def agg(detail, by, spend_min=0):
    if by not in detail.columns:
        return []
    g = (
        detail.groupby(by, dropna=False)
        .agg({"Расход, ₽": "sum", "Показы": "sum", "Клики": "sum", "Конверсии": "sum"})
        .reset_index()
    )
    if spend_min:
        g = g[g["Расход, ₽"] >= spend_min]
    g = g.sort_values("Расход, ₽", ascending=False)
    rows = []
    for _, r in g.iterrows():
        ctr, cpc, cpa, _cr = kpi_row(r["Расход, ₽"], r["Показы"], r["Клики"], r["Конверсии"])
        rows.append(
            [
                str(r[by])[:80],
                f"{r['Расход, ₽']:.2f}",
                int(r["Показы"]),
                int(r["Клики"]),
                int(r["Конверсии"]),
                f"{ctr:.2f}",
                f"{cpc:.2f}" if r["Клики"] else "—",
                f"{cpa:.2f}" if r["Конверсии"] else "—",
            ]
        )
    return rows


def search_query_stats(detail):
    qcol = "Поисковый запрос"
    search = detail[detail["Тип площадки"] == "Поиск"].copy()
    search = search[search[qcol].notna() & (search[qcol].astype(str).str.strip() != "")]
    if search.empty:
        return pd.DataFrame()
    return (
        search.groupby(qcol)
        .agg({"Расход, ₽": "sum", "Клики": "sum", "Конверсии": "sum", "Показы": "sum"})
        .reset_index()
    )


def minus_words_from_data(query_stats):
    extra = []
    for _, r in query_stats.iterrows():
        q = str(r["Поисковый запрос"])
        clicks = int(r["Клики"])
        if clicks < 2:
            continue
        if OFF_QUERY_RE.search(q):
            for token in [
                "петрович",
                "вимос",
                "каталог",
                "технониколь",
                "baucenter",
                "анапа",
                "неводекор",
                "мдф",
                "пластик",
            ]:
                if token in q.lower() and token not in extra:
                    extra.append(token)
    return extra


def target_queries_from_data(query_stats, limit=15):
    rows = []
    for _, r in query_stats.sort_values("Клики", ascending=False).iterrows():
        q = str(r["Поисковый запрос"])
        if TARGET_QUERY_RE.search(q) and not OFF_QUERY_RE.search(q):
            rows.append(q)
        if len(rows) >= limit:
            break
    return rows


def keywords_for_groups(detail, query_stats):
    groups = detail["Название группы"].dropna().unique()
    blocks = []
    for gname in groups:
        gl = str(gname).lower()
        keys = []
        for pattern, phrases in KEYWORD_TEMPLATES.items():
            if pattern in gl:
                keys.extend(phrases)
        if keys:
            blocks.append((str(gname), keys))
    from_data = target_queries_from_data(query_stats)
    if from_data:
        blocks.append(("Из отчёта (целевые запросы с кликами)", from_data))
    if not blocks:
        blocks.append(("Базовый набор", KEYWORD_TEMPLATES["коммерч"] + KEYWORD_TEMPLATES["гео"]))
    return blocks


def rsya_exclude_list(detail, min_clicks=1, top_n=25):
    if "Название площадки" not in detail.columns:
        return []
    rsy = detail[detail["Тип площадки"] == "Сети"]
    if rsy.empty:
        return []
    pl = (
        rsy.groupby("Название площадки")
        .agg({"Показы": "sum", "Клики": "sum", "Конверсии": "sum", "Расход, ₽": "sum"})
        .reset_index()
    )
    rows = []
    for _, r in pl.iterrows():
        name = str(r["Название площадки"])
        clicks = int(r["Клики"])
        shows = int(r["Показы"])
        conv = int(r["Конверсии"])
        if clicks < min_clicks and shows < 50:
            continue
        reasons = []
        if RSYA_BAN_NAME_RE.search(name):
            reasons.append("игры/Дзен/погода/соцсети")
        if name.startswith("com.") and "yandex" not in name.lower():
            reasons.append("мобильное приложение")
        if clicks >= 3 and conv == 0:
            reasons.append("клики без конверсий")
        if not reasons and clicks >= 5:
            reasons.append("много кликов — проверить релевантность")
        if reasons:
            rows.append([name[:55], shows, clicks, conv, "; ".join(reasons)])
    rows.sort(key=lambda x: (x[2], x[1]), reverse=True)
    return rows[:top_n]


def paid_search_condition_note(detail):
    search_paid = detail[(detail["Тип площадки"] == "Поиск") & (detail["Расход, ₽"] > 0)]
    if search_paid.empty or "Тип условия показа" not in search_paid.columns:
        return ""
    cond = search_paid.groupby("Тип условия показа")["Расход, ₽"].sum()
    top = cond.idxmax() if len(cond) else ""
    auto_share = cond.get("Автотаргетинг", 0) / cond.sum() * 100 if cond.sum() else 0
    if top == "Автотаргетинг" or auto_share > 70:
        return (
            f"На поиске {auto_share:.0f}% расхода идёт через «Автотаргетинг» — добавьте ключевые фразы "
            "(фразовое/точное соответствие) и сузьте широкий автотаргет; иначе бюджет уходит на "
            "нерелевантные запросы (Петрович, каталоги, конкуренты)."
        )
    return ""


def off_target_spend(query_stats):
    if query_stats.empty:
        return 0, 0
    off = query_stats[
        query_stats["Поисковый запрос"].astype(str).str.contains(OFF_QUERY_RE, na=False)
    ]
    return float(off["Расход, ₽"].sum()), int(off["Клики"].sum())


def add_action_block(doc, detail):
    doc.add_heading("8. Минус-слова, ключевые фразы, площадки к отключению", level=1)
    doc.add_paragraph(
        "Блок формируется автоматически: базовые списки для ниши фасадов/термопанелей "
        "+ данные из выгрузки (запросы и площадки РСЯ, если есть в отчёте)."
    )

    qs = search_query_stats(detail)
    extra_minus = minus_words_from_data(qs) if len(qs) else []
    minus_words = sorted(set(BASE_MINUS_WORDS + extra_minus))

    doc.add_heading("8.1. Минус-слова (уровень кампании)", level=2)
    add_code_block(doc, "Скопировать в Директ (по одному слову на строку):", minus_words)

    doc.add_heading("8.2. Минус-фразы", level=2)
    add_code_block(doc, "Рекомендуемые минус-фразы:", BASE_MINUS_PHRASES)

    doc.add_heading("8.3. Опционально (если не продаёте панели без утепления)", level=2)
    add_code_block(doc, "", OPTIONAL_MINUS_IF_NO_DRY_PANELS)

    if len(qs):
        off_paid = []
        for _, r in qs.iterrows():
            if r["Расход, ₽"] <= 0:
                continue
            q = str(r["Поисковый запрос"])
            if OFF_QUERY_RE.search(q):
                off_paid.append(
                    [q[:60], f"{r['Расход, ₽']:.0f}", int(r["Клики"]), int(r["Конверсии"]), "в минус-фразы"]
                )
        if off_paid:
            doc.add_heading("8.4. Платные запросы — отсечь", level=2)
            add_table(doc, ["Запрос", "₽", "Клики", "Конв.", "Действие"], off_paid[:30])

    doc.add_heading("8.5. Ключевые фразы по группам", level=2)
    auto_note = paid_search_condition_note(detail)
    if auto_note:
        doc.add_paragraph(auto_note)
    for gname, phrases in keywords_for_groups(detail, qs):
        doc.add_paragraph(f"Группа: {gname}", style="List Bullet")
        add_code_block(doc, "", phrases)

    doc.add_heading("8.6. Площадки РСЯ — убрать из показов", level=2)
    ban_rows = rsya_exclude_list(detail)
    if ban_rows:
        add_table(doc, ["Площадка", "Показы", "Клики", "Конв.", "Причина"], ban_rows)
        add_code_block(doc, "Список для копирования:", [r[0] for r in ban_rows])
    else:
        doc.add_paragraph(
            "В этой выгрузке нет колонки «Название площадки». Скачайте отчёт «Площадки» "
            "за тот же период и запретите: dzen.ru, mail.ru, игры, com.* (кроме Яндекса)."
        )

    doc.add_heading("8.7. Настройки (кратко)", level=2)
    for b in [
        "Поиск: фразовое + точное соответствие; не полагаться только на автотаргет.",
        "Сети: отдельная кампания с низким бюджетом или отключить до настройки поиска.",
        "Группы «транз» и «Лен. область» — проверить гео-таргетинг и ставки отдельно от Выборга.",
        "После добавления минус-слов — повторный анализ через 7–14 дней.",
    ]:
        doc.add_paragraph(b, style="List Bullet")


def build_report(csv_path, out_docx, campaign_name=None):
    detail, total = load_csv(csv_path)
    spend = float(total["Расход, ₽"]) if total is not None else detail["Расход, ₽"].sum()
    shows = int(total["Показы"]) if total is not None else int(detail["Показы"].sum())
    clicks = int(total["Клики"]) if total is not None else int(detail["Клики"].sum())
    conv = int(total["Конверсии"]) if total is not None else int(detail["Конверсии"].sum())
    ctr, cpc, cpa, cr = kpi_row(spend, shows, clicks, conv)

    if "Название кампании" in detail.columns:
        camp = detail["Название кампании"].iloc[0]
        camp_id = detail["№ Кампании"].iloc[0]
    else:
        camp = campaign_name or "Термопанели (marmara-pro.ru)"
        camp_id = "—"
    dmin, dmax = detail["День"].min(), detail["День"].max()

    qs = search_query_stats(detail)
    off_spend, off_clicks = off_target_spend(qs)

    doc = Document()
    title = doc.add_heading("Анализ кампании Яндекс Директ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Кампания: {camp} (№ {camp_id})\n"
        f"Сайт: marmara-pro.ru / КлинкерПрофи\n"
        f"Период: {dmin} — {dmax}\n"
        f"Источник: {Path(csv_path).name}\n"
        f"Дата отчёта: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Сводка (KPI)", level=1)
    add_table(
        doc,
        ["Показатель", "Значение"],
        [
            ["Расход, ₽", f"{spend:.2f}"],
            ["Показы", str(shows)],
            ["Клики", str(clicks)],
            ["Конверсии", str(conv)],
            ["CTR, %", f"{ctr:.2f}"],
            ["CPC, ₽", f"{cpc:.2f}"],
            ["CPA, ₽", f"{cpa:.2f}" if conv else "—"],
            ["CR, %", f"{cr:.2f}"],
            ["Нецелевой расход (оценка)", f"{off_spend:.0f} ₽ ({off_clicks} кл.)"],
        ],
        [8, 6],
    )

    doc.add_heading("2. Группы объявлений", level=1)
    add_table(
        doc,
        ["Группа", "₽", "Показы", "Клики", "Конв.", "CTR%", "CPC", "CPA"],
        agg(detail, "Название группы")[:12],
    )

    doc.add_heading("3. Категории поисковых запросов", level=1)
    if "Категория запроса" in detail.columns:
        add_table(
            doc,
            ["Категория", "₽", "Показы", "Клики", "Конв.", "CTR%", "CPC", "CPA"],
            agg(detail, "Категория запроса")[:10],
        )
        doc.add_paragraph(
            "«Альтернативные» и «Широкие» — часто запросы к конкурентам и строймагазинам; "
            "«Целевые» — основной потенциал для масштабирования."
        )

    doc.add_heading("4. Условия показа (поиск)", level=1)
    if "Тип условия показа" in detail.columns:
        search = detail[detail["Тип площадки"] == "Поиск"]
        add_table(
            doc,
            ["Условие", "₽", "Показы", "Клики", "Конв.", "CTR%", "CPC", "CPA"],
            agg(search, "Тип условия показа"),
        )

    doc.add_heading("5. Поисковые запросы (с расходом)", level=1)
    qcol = "Поисковый запрос"
    q = detail[
        detail[qcol].notna()
        & (detail[qcol].astype(str).str.strip() != "")
        & (detail["Расход, ₽"] > 0)
    ]
    add_table(
        doc,
        ["Запрос", "₽", "Показы", "Клики", "Конв.", "CTR%", "CPC", "CPA"],
        agg(q, qcol)[:35],
    )

    doc.add_heading("6. Заголовки объявлений", level=1)
    hrows = [
        r
        for r in agg(detail, "Заголовок", spend_min=0)
        if float(r[1].replace(",", ".")) > 0 or int(r[4]) > 0
    ][:15]
    add_table(
        doc,
        ["Заголовок", "₽", "Показы", "Клики", "Конв.", "CTR%", "CPC", "CPA"],
        hrows,
    )

    doc.add_heading("7. Сети (РСЯ) и размещения", level=1)
    rsy = detail[detail["Тип площадки"] == "Сети"]
    doc.add_paragraph(
        f"РСЯ: {int(rsy['Показы'].sum())} показов, {int(rsy['Клики'].sum())} кликов, "
        f"расход {float(rsy['Расход, ₽'].sum()):.2f} ₽."
    )
    if "Вид размещения" in detail.columns:
        add_table(
            doc,
            ["Вид размещения", "₽", "Показы", "Клики", "Конв.", "CTR%", "CPC", "CPA"],
            agg(detail, "Вид размещения")[:10],
        )

    doc.add_heading("8. Динамика по дням", level=1)
    daily = (
        detail.groupby("День")
        .agg({"Расход, ₽": "sum", "Показы": "sum", "Клики": "sum", "Конверсии": "sum"})
        .reset_index()
        .sort_values("День")
    )
    add_table(
        doc,
        ["День", "₽", "Показы", "Клики", "Конв."],
        [
            [
                str(r["День"]),
                f"{r['Расход, ₽']:.2f}",
                int(r["Показы"]),
                int(r["Клики"]),
                int(r["Конверсии"]),
            ]
            for _, r in daily.iterrows()
        ],
    )

    doc.add_heading("9. Выводы и рекомендации", level=1)
    bullets = [
        f"За период потрачено {spend:.0f} ₽, получено {conv} конверсий (CPA ≈ {cpa:.0f} ₽). "
        "Сверьте цели Метрики с реальными заявками (звонок, форма, чат).",
    ]
    auto_note = paid_search_condition_note(detail)
    if auto_note:
        bullets.append(auto_note)
    if off_spend > 50:
        bullets.append(
            f"Оценочно {off_spend:.0f} ₽ ({off_clicks} кл.) ушло на нецелевые запросы "
            "(Петрович, каталоги, конкуренты, пластик/МДФ) — добавьте минус-слова из раздела 10."
        )
    bullets.extend(
        [
            "Единственная конверсия с явным гео-запросом: «термопанели в выборге купить» (100 ₽) — "
            "усилить гео-группу и ставки на «выборг», «ленина 11», «собственное производство».",
            "Группа «Термопанели (транз)» — самый большой расход (≈862 ₽), но 1 конверсия; "
            "проверить гео (не уходит ли бюджет за пределы ЛО) и минус-слова по СПб без «выборг».",
            "Платные запросы «петрович», «строим наш дом», «технониколь» — в минус-фразы; "
            "клиент ищет чужой магазин, не производителя.",
            "Целевые запросы с расходом без конверсий — проверить посадочную /termo и AI-чат на сайте.",
            "РСЯ: если нет отдельного отчёта по площадкам — скачать и отключить игры, Дзен, приложения.",
            "Масштабировать заголовки с лучшим CTR и запросы категории «Целевые».",
        ]
    )
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_page_break()
    add_action_block(doc, detail)

    doc.add_page_break()
    doc.add_heading("11. Техническая справка", level=1)
    doc.add_paragraph(
        "Отчёт: scripts/yandex_direct_report_docx.py. Формат vitaminki-style PPC audit. "
        "Раздел 10 (минус-слова, ключи, площадки) формируется при каждом запуске."
    )

    doc.save(out_docx)
    return out_docx


if __name__ == "__main__":
    csv_p = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("input.csv")
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else csv_p.with_suffix(".docx")
    camp = sys.argv[3] if len(sys.argv) > 3 else None
    build_report(csv_p, out, camp)
    print(out)
