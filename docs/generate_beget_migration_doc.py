#!/usr/bin/env python3
"""Generate Beget VPS migration guide for boat + klinker bots."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).parent / "beget-migration-boat-klinker.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # Title
    title = doc.add_heading("Перенос ботов Boat Sochi и KlinkerPro на VPS Beget", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "Подробная инструкция по миграции двух ИИ-ботов с текущего VPS (webtaxi2.ru) "
        "на виртуальный сервер Beget. Документ подготовлен: 21 августа 2026 г.",
        italic=True,
    )
    doc.add_paragraph()

    # TOC-like section list
    add_heading(doc, "Содержание", 1)
    toc = [
        "1. Что переносим и зачем",
        "2. Что заказать на Beget",
        "3. Домены и DNS",
        "4. Первичная настройка VPS",
        "5. Установка Node.js, nginx, certbot, pm2",
        "6. Перенос бота Boat Sochi (boat-sochi.ru)",
        "7. Перенос бота KlinkerPro (marmara-pro.ru)",
        "8. Nginx: конфигурации для обоих ботов",
        "9. SSL-сертификаты (Let's Encrypt)",
        "10. Обновление виджетов на Tilda",
        "11. Cron: мониторинг и фоновые задачи",
        "12. План переключения (cutover) со старого VPS",
        "13. Проверка после миграции",
        "14. Частые проблемы",
        "15. Полезные ссылки",
    ]
    for item in toc:
        add_bullet(doc, item)

    # Section 1
    add_heading(doc, "1. Что переносим и зачем", 1)
    add_para(
        doc,
        "На одном VPS работают два независимых Node.js-приложения — ИИ-чаты для сайтов на Tilda. "
        "Каждый бот слушает свой порт, снаружи доступен через HTTPS и поддомен.",
    )
    add_table(
        doc,
        ["Бот", "Сайт", "Текущий URL", "Порт", "Ветка GitHub", "Назначение"],
        [
            [
                "Boat Sochi",
                "boat-sochi.ru",
                "https://boat.webtaxi2.ru",
                "3000",
                "cursor/boat-contact-route-5814",
                "Чат на сайте аренды яхт; заявки в MAX; отзывы Avito",
            ],
            [
                "KlinkerPro",
                "marmara-pro.ru/termo",
                "https://klinker.webtaxi2.ru",
                "3001",
                "cursor/termopaneli-bot-bfbc",
                "Чат по термопанелям; заявки в MAX (пока отключены)",
            ],
        ],
    )
    add_para(doc, "Репозиторий: https://github.com/Igor-Lark/Igor", bold=False)
    add_para(
        doc,
        "После переноса старый VPS (webtaxi2.ru) можно отключить, когда новые адреса "
        "протестированы и виджеты на Tilda обновлены.",
    )

    # Section 2
    add_heading(doc, "2. Что заказать на Beget", 1)
    add_heading(doc, "2.1. Тариф VPS", 2)
    add_bullet(doc, "Раздел Beget → «VPS/VDS» → Ubuntu 22.04 или 24.04 LTS.")
    add_bullet(doc, "Минимум: 1 vCPU, 1–2 GB RAM, 10 GB SSD — достаточно для двух лёгких Node.js-ботов.")
    add_bullet(doc, "Рекомендуется: 2 vCPU, 2 GB RAM — запас под пики YandexGPT и certbot.")
    add_bullet(doc, "Регион: любой с низкой задержкой до России (Москва/СПб, если доступен).")

    add_heading(doc, "2.2. Что понадобится заранее", 2)
    add_bullet(doc, "Доступ по SSH (логин root или пользователь с sudo).")
    add_bullet(doc, "Файл .env с текущего VPS для boat (YandexGPT, MAX, Telegram — если используется).")
    add_bullet(doc, "Доступ к панели DNS Beget (домены уже есть у вас).")
    add_bullet(doc, "Доступ к Tilda: boat-sochi.ru и marmara-pro.ru — для смены URL embed.js.")

    # Section 3
    add_heading(doc, "3. Домены и DNS", 1)
    add_para(
        doc,
        "Ботам нужны отдельные поддомены с HTTPS. Два варианта — выберите один.",
    )

    add_heading(doc, "3.1. Вариант A — поддомены на вашем домене в Beget (рекомендуется)", 2)
    add_para(
        doc,
        "Если у вас на Beget есть домен, например example.ru, создайте поддомены второго уровня:",
    )
    add_table(
        doc,
        ["Тип", "Имя (поддомен)", "Значение", "Назначение"],
        [
            ["A", "boat-bot", "IP нового VPS Beget", "Boat Sochi bot"],
            ["A", "klinker-bot", "IP нового VPS Beget", "KlinkerPro bot"],
        ],
    )
    add_para(doc, "Итоговые URL:")
    add_bullet(doc, "Boat: https://boat-bot.example.ru")
    add_bullet(doc, "Klinker: https://klinker-bot.example.ru")
    add_para(
        doc,
        "В панели Beget: «Домены» → ваш домен → «Поддомены» или «DNS-зона» → добавить A-записи. "
        "TTL: 300–600 сек. Распространение DNS: 5–30 минут, иногда до 2 часов.",
    )

    add_heading(doc, "3.2. Вариант B — оставить webtaxi2.ru", 2)
    add_para(
        doc,
        "Если зона webtaxi2.ru тоже на Beget, можно просто сменить A-записи boat и klinker "
        "на IP нового VPS (вместо старого сервера). URL не меняются — Tilda трогать не нужно.",
    )
    add_table(
        doc,
        ["Тип", "Имя", "Новое значение"],
        [
            ["A", "boat", "IP VPS Beget"],
            ["A", "klinker", "IP VPS Beget"],
        ],
    )

    add_heading(doc, "3.3. Проверка DNS", 2)
    add_code(
        doc,
        "dig +short boat-bot.example.ru\n"
        "dig +short klinker-bot.example.ru\n"
        "# или для webtaxi2.ru:\n"
        "dig +short boat.webtaxi2.ru\n"
        "dig +short klinker.webtaxi2.ru",
    )
    add_para(doc, "Должен вернуться IP вашего нового VPS Beget.")

    # Section 4
    add_heading(doc, "4. Первичная настройка VPS", 1)
    add_para(doc, "Подключитесь по SSH (IP и пароль/ключ из письма Beget после заказа VPS):")
    add_code(doc, "ssh root@ВАШ_IP_BEGET")
    add_para(doc, "Обновление системы и базовые пакеты:")
    add_code(
        doc,
        "apt update && apt upgrade -y\n"
        "apt install -y git curl wget ufw fail2ban unzip",
    )
    add_para(doc, "Часовой пояс (важно для напоминаний boat и cron):")
    add_code(
        doc,
        "timedatectl set-timezone Europe/Moscow\ntimedatectl",
    )
    add_para(doc, "Создайте пользователя для деплоя (не обязательно root):")
    add_code(
        doc,
        "adduser deploy\n"
        "usermod -aG sudo deploy\n"
        "# скопируйте SSH-ключ:\n"
        "mkdir -p /home/deploy/.ssh\n"
        "cp /root/.ssh/authorized_keys /home/deploy/.ssh/\n"
        "chown -R deploy:deploy /home/deploy/.ssh",
    )
    add_para(doc, "Firewall (открыть SSH, HTTP, HTTPS):")
    add_code(
        doc,
        "ufw allow OpenSSH\n"
        "ufw allow 80/tcp\n"
        "ufw allow 443/tcp\n"
        "ufw enable\n"
        "ufw status",
    )

    # Section 5
    add_heading(doc, "5. Установка Node.js, nginx, certbot, pm2", 1)
    add_heading(doc, "5.1. Node.js 20 LTS", 2)
    add_code(
        doc,
        "curl -fsSL https://deb.nodesource.com/setup_20.x | bash -\n"
        "apt install -y nodejs\n"
        "node -v   # v20.x\n"
        "npm -v",
    )

    add_heading(doc, "5.2. pm2 (менеджер процессов)", 2)
    add_code(doc, "npm install -g pm2\npm2 startup\n# выполните команду, которую выведет pm2 startup")

    add_heading(doc, "5.3. nginx", 2)
    add_code(doc, "apt install -y nginx\nsystemctl enable nginx\nsystemctl start nginx")

    add_heading(doc, "5.4. certbot (Let's Encrypt)", 2)
    add_code(
        doc,
        "apt install -y certbot python3-certbot-nginx\n"
        "# сертификаты выпустим после настройки nginx (раздел 9)",
    )

    # Section 6 - Boat
    add_heading(doc, "6. Перенос бота Boat Sochi (boat-sochi.ru)", 1)
    add_heading(doc, "6.1. Клонирование кода", 2)
    add_code(
        doc,
        "mkdir -p /var/www\n"
        "cd /var/www\n"
        "git clone https://github.com/Igor-Lark/Igor.git boat-sochi-bot\n"
        "cd boat-sochi-bot\n"
        "git fetch origin cursor/boat-contact-route-5814\n"
        "git checkout cursor/boat-contact-route-5814\n"
        "git pull origin cursor/boat-contact-route-5814",
    )
    add_para(
        doc,
        "Примечание: boat-sochi-bot — отдельная копия репозitorия в корне ветки "
        "(не папка bots/). На старом VPS путь был /var/www/boat-sochi-bot — сохраняем ту же структуру.",
    )

    add_heading(doc, "6.2. Зависимости", 2)
    add_code(doc, "cd /var/www/boat-sochi-bot\nnpm install")

    add_heading(doc, "6.3. Файл .env", 2)
    add_para(doc, "Скопируйте .env со старого VPS или создайте из шаблона:")
    add_code(doc, "cp .env.example .env\nnano .env")
    add_para(doc, "Обязательно измените PUBLIC_URL на новый адрес:")
    add_code(
        doc,
        "# --- AI: YandexGPT ---\n"
        "YANDEX_API_KEY=AQVN...\n"
        "YANDEX_FOLDER_ID=b1g...\n"
        "YANDEX_MODEL=yandexgpt-lite\n\n"
        "# --- MAX (заявки менеджеру) ---\n"
        "MAX_BOT_TOKEN=...\n"
        "MAX_CHAT_ID=...\n"
        "MAX_USER_ID=...\n"
        "AVITO_ITEM_URL=https://m.avito.ru/sochi/...\n"
        "AVITO_NOTIFY_ALWAYS=0\n\n"
        "# --- Telegram (если используется) ---\n"
        "TELEGRAM_BOT_TOKEN=...\n\n"
        "# --- Сервер ---\n"
        "PORT=3000\n"
        "PUBLIC_URL=https://boat-bot.example.ru\n"
        "# или: PUBLIC_URL=https://boat.webtaxi2.ru\n"
        "BOT_NAME=Boat Sochi\n"
        "TZ=Europe/Moscow\n\n"
        "BOOKING_REMINDERS_ENABLED=0\n"
        "NO_CONTACT_IDLE_MINUTES=10\n"
        "AI_ALERT_COOLDOWN_MINUTES=30",
    )
    add_para(
        doc,
        "Копирование .env со старого сервера (выполнить на своём ПК или со старого VPS):",
    )
    add_code(doc, "scp root@СТАРЫЙ_IP:/var/www/boat-sochi-bot/.env deploy@НОВЫЙ_IP:/var/www/boat-sochi-bot/.env")
    add_para(doc, "После копирования отредактируйте PUBLIC_URL на новый домен.")

    add_heading(doc, "6.4. Тестовый запуск", 2)
    add_code(
        doc,
        "cd /var/www/boat-sochi-bot\n"
        "npm start\n"
        "# в другом окне SSH:\n"
        "curl -s http://127.0.0.1:3000/health",
    )
    add_para(
        doc,
        "В ответе health должны быть флаги: seaRoute, contactCallback, wake, groupSailing, "
        "groupFishing, delfinCharter, streamingUi. Остановите тест: Ctrl+C.",
    )

    add_heading(doc, "6.5. Запуск через pm2", 2)
    add_code(
        doc,
        "cd /var/www/boat-sochi-bot\n"
        "pm2 start src/index.js --name boat-sochi\n"
        "pm2 save",
    )

    add_heading(doc, "6.6. Альтернатива: systemd", 2)
    add_para(doc, "Если предпочитаете systemd вместо pm2:")
    add_code(
        doc,
        "cat > /etc/systemd/system/boat-sochi-bot.service << 'EOF'\n"
        "[Unit]\n"
        "Description=Boat Sochi AI bot\n"
        "After=network.target\n\n"
        "[Service]\n"
        "Type=simple\n"
        "User=deploy\n"
        "WorkingDirectory=/var/www/boat-sochi-bot\n"
        "Environment=NODE_ENV=production\n"
        "Environment=TZ=Europe/Moscow\n"
        "ExecStart=/usr/bin/node src/index.js\n"
        "Restart=always\n"
        "RestartSec=5\n\n"
        "[Install]\n"
        "WantedBy=multi-user.target\n"
        "EOF\n\n"
        "systemctl daemon-reload\n"
        "systemctl enable boat-sochi-bot\n"
        "systemctl start boat-sochi-bot",
    )

    # Section 7 - Klinker
    add_heading(doc, "7. Перенос бота KlinkerPro (marmara-pro.ru)", 1)
    add_heading(doc, "7.1. Клонирование (отдельная копия или один репозиторий)", 2)
    add_para(doc, "Вариант 1 — один репозиторий для обоих ботов (удобнее обновлять):")
    add_code(
        doc,
        "cd /var/www\n"
        "git clone https://github.com/Igor-Lark/Igor.git igor\n"
        "cd igor\n"
        "git fetch origin cursor/termopaneli-bot-bfbc\n"
        "git checkout cursor/termopaneli-bot-bfbc\n"
        "git pull origin cursor/termopaneli-bot-bfbc\n"
        "cd bots/klinkerpro-bot\n"
        "npm install",
    )
    add_para(doc, "Вариант 2 — если boat уже в /var/www/boat-sochi-bot, можно клонировать igor отдельно.")

    add_heading(doc, "7.2. Файл .env", 2)
    add_code(
        doc,
        "cd /var/www/igor/bots/klinkerpro-bot\n"
        "cp .env.example .env\n"
        "nano .env",
    )
    add_code(
        doc,
        "YANDEX_API_KEY=...          # можно те же, что у boat\n"
        "YANDEX_FOLDER_ID=...\n"
        "YANDEX_MODEL=yandexgpt-lite\n\n"
        "MAX_NOTIFY_ENABLED=false    # заявки в MAX пока выключены\n\n"
        "PORT=3001\n"
        "PUBLIC_URL=https://klinker-bot.example.ru\n"
        "# или: PUBLIC_URL=https://klinker.webtaxi2.ru\n"
        "BOT_NAME=КлинкерПрофи\n"
        "TZ=Europe/Moscow\n"
        "NO_CONTACT_IDLE_MINUTES=10\n"
        "AI_ALERT_COOLDOWN_MINUTES=30",
    )

    add_heading(doc, "7.3. Запуск", 2)
    add_code(
        doc,
        "cd /var/www/igor/bots/klinkerpro-bot\n"
        "curl -s http://127.0.0.1:3001/health   # после npm start или pm2\n\n"
        "pm2 start src/index.js --name klinkerpro\n"
        "pm2 save\n"
        "pm2 list",
    )

    # Section 8 - Nginx
    add_heading(doc, "8. Nginx: конфигурации для обоих ботов", 1)
    add_para(
        doc,
        "Замените boat-bot.example.ru и klinker-bot.example.ru на ваши реальные имена.",
    )

    add_heading(doc, "8.1. Boat Sochi", 2)
    add_code(
        doc,
        "cat > /etc/nginx/sites-available/boat-bot << 'EOF'\n"
        "server {\n"
        "    listen 80;\n"
        "    server_name boat-bot.example.ru;\n\n"
        "    location / {\n"
        "        proxy_pass http://127.0.0.1:3000;\n"
        "        proxy_http_version 1.1;\n"
        "        proxy_set_header Host $host;\n"
        "        proxy_set_header X-Real-IP $remote_addr;\n"
        "        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n"
        "        proxy_set_header X-Forwarded-Proto $scheme;\n"
        "        proxy_read_timeout 120s;\n"
        "    }\n"
        "}\n"
        "EOF\n\n"
        "ln -sf /etc/nginx/sites-available/boat-bot /etc/nginx/sites-enabled/\n"
        "nginx -t && systemctl reload nginx",
    )

    add_heading(doc, "8.2. KlinkerPro", 2)
    add_code(
        doc,
        "cat > /etc/nginx/sites-available/klinker-bot << 'EOF'\n"
        "server {\n"
        "    listen 80;\n"
        "    server_name klinker-bot.example.ru;\n\n"
        "    location / {\n"
        "        proxy_pass http://127.0.0.1:3001;\n"
        "        proxy_http_version 1.1;\n"
        "        proxy_set_header Host $host;\n"
        "        proxy_set_header X-Real-IP $remote_addr;\n"
        "        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n"
        "        proxy_set_header X-Forwarded-Proto $scheme;\n"
        "        proxy_read_timeout 120s;\n"
        "    }\n"
        "}\n"
        "EOF\n\n"
        "ln -sf /etc/nginx/sites-available/klinker-bot /etc/nginx/sites-enabled/\n"
        "nginx -t && systemctl reload nginx",
    )

    add_para(doc, "Удалите дефолтный сайт nginx, если мешает:")
    add_code(doc, "rm -f /etc/nginx/sites-enabled/default\nnginx -t && systemctl reload nginx")

    # Section 9 - SSL
    add_heading(doc, "9. SSL-сертификаты (Let's Encrypt)", 1)
    add_para(doc, "DNS должен уже указывать на VPS. Выпуск сертификатов:")
    add_code(
        doc,
        "certbot --nginx -d boat-bot.example.ru\n"
        "certbot --nginx -d klinker-bot.example.ru\n"
        "# для webtaxi2.ru:\n"
        "# certbot --nginx -d boat.webtaxi2.ru -d klinker.webtaxi2.ru",
    )
    add_para(doc, "Автопродление:")
    add_code(doc, "certbot renew --dry-run\nsystemctl status certbot.timer")

    # Section 10 - Tilda
    add_heading(doc, "10. Обновление виджетов на Tilda", 1)
    add_para(
        doc,
        "Если меняли домены (вариант A), обновите скрипт embed.js на обоих сайтах. "
        "Если оставили webtaxi2.ru (вариант B) — этот шаг можно пропустить.",
    )

    add_heading(doc, "10.1. boat-sochi.ru", 2)
    add_para(doc, "Tilda → Настройки сайта → HTML-код перед </body>:")
    add_code(doc, '<script src="https://boat-bot.example.ru/embed.js"></script>')
    add_para(doc, "Опубликуйте сайт. Проверьте чат на любой странице boat-sochi.ru.")

    add_heading(doc, "10.2. marmara-pro.ru", 2)
    add_code(doc, '<script src="https://klinker-bot.example.ru/embed.js"></script>')
    add_para(doc, "Страница каталога: https://marmara-pro.ru/termo — кнопка чата в углу.")

    add_heading(doc, "10.3. Telegram webhook (если используется)", 2)
    add_para(
        doc,
        "После смены PUBLIC_URL перезапустите boat — webhook Telegram обновится автоматически "
        "при старте (если TELEGRAM_BOT_TOKEN задан и PUBLIC_URL — HTTPS).",
    )
    add_code(doc, "pm2 restart boat-sochi\ncurl -s https://boat-bot.example.ru/health")

    # Section 11 - Cron
    add_heading(doc, "11. Cron: мониторинг и фоновые задачи", 1)
    add_para(doc, "Открыть crontab пользователя deploy:")
    add_code(doc, "crontab -e -u deploy")
    add_para(doc, "Рекомендуемые задачи для boat:")
    add_code(
        doc,
        "# Health ping каждые 10 минут (алерт в MAX при падении)\n"
        "*/10 * * * * cd /var/www/boat-sochi-bot && npm run health:ping -- --notify >> /var/log/boat-health.log 2>&1\n\n"
        "# Проверка новых отзывов Avito (раз в 6 часов)\n"
        "0 */6 * * * cd /var/www/boat-sochi-bot && npm run check:avito >> /var/log/boat-avito.log 2>&1\n\n"
        "# Напоминания о бронировании (если включите BOOKING_REMINDERS_ENABLED=1)\n"
        "# */15 * * * * cd /var/www/boat-sochi-bot && npm run remind:bookings >> /var/log/boat-remind.log 2>&1",
    )
    add_para(doc, "Для klinker (опционально):")
    add_code(
        doc,
        "*/10 * * * * cd /var/www/igor/bots/klinkerpro-bot && npm run health:ping -- --notify >> /var/log/klinker-health.log 2>&1",
    )

    # Section 12 - Cutover
    add_heading(doc, "12. План переключения (cutover) со старого VPS", 1)
    add_para(doc, "Рекомендуемый порядок — минимум простоя:")
    add_bullet(doc, "Шаг 1. Развернуть оба бота на Beget, проверить health по IP/новому домену (до смены DNS).")
    add_bullet(doc, "Шаг 2. Временно прописать в /etc/hosts на своём ПК новый IP для теста домена.")
    add_bullet(doc, "Шаг 3. Проверить чат на Tilda (локально через hosts) или временный поддомен.")
    add_bullet(doc, "Шаг 4. Сменить A-записи DNS (или обновить Tilda, если новые домены).")
    add_bullet(doc, "Шаг 5. pm2 restart all на новом VPS; curl health снаружи.")
    add_bullet(doc, "Шаг 6. Обновить embed.js на Tilda и опубликовать.")
    add_bullet(doc, "Шаг 7. 24–48 часов понаблюдать логи: pm2 logs, journalctl.")
    add_bullet(doc, "Шаг 8. Остановить боты на старом VPS: pm2 stop all; отключить автозапуск.")
    add_bullet(doc, "Шаг 9. Через неделю — отменить старый VPS у прежнего хостера.")

    add_heading(doc, "12.1. Бэкап перед миграцией", 2)
    add_code(
        doc,
        "# на старом VPS:\n"
        "tar czf ~/backup-bots-$(date +%Y%m%d).tar.gz \\\n"
        "  /var/www/boat-sochi-bot/.env \\\n"
        "  /var/www/boat-sochi-bot/knowledge \\\n"
        "  ~/igor/bots/klinkerpro-bot/.env 2>/dev/null\n"
        "scp root@СТАРЫЙ_IP:~/backup-bots-*.tar.gz .",
    )

    # Section 13 - Verification
    add_heading(doc, "13. Проверка после миграции", 1)
    add_table(
        doc,
        ["Проверка", "Команда / действие", "Ожидание"],
        [
            ["Boat health", "curl -s https://boat-bot.example.ru/health", "JSON, ai: yandex или openai"],
            ["Klinker health", "curl -s https://klinker-bot.example.ru/health", "JSON, ai настроен"],
            ["Boat embed", "curl -sI https://boat-bot.example.ru/embed.js | head -3", "HTTP/2 200"],
            ["Klinker embed", "curl -sI https://klinker-bot.example.ru/embed.js | head -3", "HTTP/2 200"],
            ["Boat чат", "Вопрос на boat-sochi.ru", "Ответ про яхты/катера"],
            ["Klinker чат", "Вопрос на marmara-pro.ru/termo", "Ответ про термопанели"],
            ["Заявка boat", "«Хочу забронировать, +7...»", "Сообщение в MAX"],
            ["pm2", "pm2 list", "boat-sochi и klinkerpro online"],
        ],
    )

    add_heading(doc, "13.1. Обновление кода в будущем", 2)
    add_para(doc, "Boat:")
    add_code(
        doc,
        "cd /var/www/boat-sochi-bot\n"
        "git fetch origin cursor/boat-contact-route-5814\n"
        "git pull origin cursor/boat-contact-route-5814\n"
        "pm2 restart boat-sochi\n"
        "curl -s https://boat-bot.example.ru/health",
    )
    add_para(doc, "Klinker:")
    add_code(
        doc,
        "cd /var/www/igor\n"
        "bash bots/klinkerpro-bot/scripts/deploy-knowledge.sh\n"
        "# или вручную:\n"
        "git pull origin cursor/termopaneli-bot-bfbc\n"
        "pm2 restart klinkerpro",
    )

    # Section 14 - Troubleshooting
    add_heading(doc, "14. Частые проблемы", 1)
    add_table(
        doc,
        ["Симптом", "Причина", "Решение"],
        [
            ["502 Bad Gateway", "Node не запущен или неверный порт", "pm2 list; systemctl status; nginx proxy_pass"],
            ["ai: none в health", "Нет Yandex ключей", "Проверить .env, перезапустить pm2"],
            ["Виджет не появляется", "Старый URL в Tilda или не опубликовано", "F12 → Network → embed.js"],
            ["SSL ошибка", "DNS ещё не обновился", "dig +short; подождать; certbot снова"],
            ["Telegram не отвечает", "Webhook на старый URL", "PUBLIC_URL + pm2 restart boat-sochi"],
            ["CORS / чат молчит", "HTTP вместо HTTPS на сайте", "Tilda только HTTPS"],
            ["Порт занят", "Два процесса на 3000/3001", "ss -tlnp | grep 300"],
        ],
    )

    # Section 15 - Links
    add_heading(doc, "15. Полезные ссылки", 1)
    links = [
        ("Репозиторий Igor", "https://github.com/Igor-Lark/Igor"),
        ("Ветка boat", "https://github.com/Igor-Lark/Igor/tree/cursor/boat-contact-route-5814"),
        ("PR boat", "https://github.com/Igor-Lark/Igor/pull/17"),
        ("Ветка klinker", "https://github.com/Igor-Lark/Igor/tree/cursor/termopaneli-bot-bfbc"),
        ("DEPLOY klinker (markdown)", "https://github.com/Igor-Lark/Igor/blob/cursor/termopaneli-bot-bfbc/bots/klinkerpro-bot/DEPLOY.md"),
        ("Сайт boat-sochi.ru", "https://boat-sochi.ru"),
        ("Сайт marmara-pro.ru/termo", "https://marmara-pro.ru/termo"),
        ("Yandex Cloud", "https://console.cloud.yandex.ru"),
        ("MAX dev", "https://dev.max.ru"),
        ("Beget VPS", "https://beget.com/ru/vps"),
        ("Текущий boat (старый VPS)", "https://boat.webtaxi2.ru/health"),
        ("Текущий klinker (старый VPS)", "https://klinker.webtaxi2.ru/health"),
    ]
    for name, url in links:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{name}: ")
        r1.bold = True
        r2 = p.add_run(url)
        r2.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

    doc.add_paragraph()
    add_para(
        doc,
        "При выборе поддоменов на Beget напишите, какой домен используете — "
        "можно сразу подставить готовые A-записи и nginx-конфиги под ваши имена.",
        italic=True,
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
