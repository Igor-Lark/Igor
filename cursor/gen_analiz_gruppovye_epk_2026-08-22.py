#!/usr/bin/env python3
"""Дополнение ЕПК групповых прогулок: CSV 20–22.08 + РСЯ за неделю."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Analiz_RK_Gruppovye_progulki_EPK_2026-08-22.docx"
LOCAL_DIR = ROOT.parent / "локальная"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else (13 if level == 2 else 12),
            bold=True,
            color=RGBColor(0x0B, 0x3D, 0x5C),
        )


def add_p(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
            if r_i % 2 == 1:
                shade_cell(cell, "F2F7FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.6)
    s.bottom_margin = Cm(1.6)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.6)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("АНАЛИЗ КАМПАНИИ · ЯНДЕКС ДИРЕКТ")
    set_run_font(r, size=12, bold=True, color=RGBColor(0x0B, 0x3D, 0x5C))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ЕПК «Групповые прогулки» — дополнение 20–22.08")
    set_run_font(r, size=18, bold=True, color=RGBColor(0x0B, 0x3D, 0x5C))

    for line in [
        "Кампания: «ЕПК Групповые прогулки» · № 713449981",
        "Новый файл: 2026-08-22_21-40-00_vitaminki21 · дни 20–22.08 (20.08 теперь полный)",
        "Сшивка с отчётом 14–20.08: календарная неделя Пн–Сб 17–22.08 (вс 23.08 ещё нет)",
        "Отдельно: РСЯ за неделю — наблюдение, паузу не делали",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=10, color=RGBColor(0x44, 0x44, 0x44))

    add_heading(doc, "1. Что добавил этот CSV", 1)
    add_p(
        doc,
        "Прошлый срез обрывался 20.08 днём (52 показа). Этот файл закрывает 20.08 целиком и даёт 21–22.08 до ~21:40. "
        "Итого в базе агента теперь: 11–22.08, последний полный вечер — 22.08.",
    )
    add_table(
        doc,
        ["День", "Расход", "Показы", "Клики", "Конв.", "CTR", "Источник"],
        [
            ["20.08 было (срез днём)", "0", "52", "3", "0", "5,8%", "CSV 20.08 15:33"],
            ["20.08 стало (полный)", "0", "122", "5", "0", "4,1%", "этот файл"],
            ["21.08", "0", "109", "6", "0", "5,5%", "этот файл"],
            ["22.08 до 21:40", "145,20", "69", "4", "1", "5,8%", "этот файл"],
            ["Итого 20–22.08", "145,20", "300", "15", "1", "5,0%", "CPA 145 ₽"],
        ],
        [4.4, 2.0, 1.8, 1.6, 1.6, 1.6, 3.6],
    )
    add_p(
        doc,
        "Единственная новая конверсия: 22.08, запрос «выход в море сочи», широкие, автотаргет, "
        "Поиск / прочее, заголовок «Закат … 2 500 ₽», смартфон 25–34. Расход 145,20 ₽ = цена этой конверсии. "
        "Интент широкий (не Сириус / не билет) — не мусор уровня Дубай, но и не ядро. Минусовать не стоит.",
    )

    add_heading(doc, "2. Неделя Пн–Сб 17–22.08 (сшивка)", 1)
    add_p(
        doc,
        "17–19.08 из прошлого CSV, 20–22.08 из нового. Воскресенье 23.08 ещё впереди — порог обучения 10 конв./нед. пока 8.",
    )
    add_table(
        doc,
        ["День", "Расход", "Показы", "Клики", "Конв.", "CTR"],
        [
            ["17.08", "410", "210", "19", "3", "9,1%"],
            ["18.08", "436", "190", "20", "3", "10,5%"],
            ["19.08", "138", "116", "12", "1", "10,3%"],
            ["20.08 (после минусов 5.1)", "0", "122", "5", "0", "4,1%"],
            ["21.08", "0", "109", "6", "0", "5,5%"],
            ["22.08", "145", "69", "4", "1", "5,8%"],
            ["17–22.08", "1 129", "816", "66", "8", "8,1%"],
        ],
        [4.8, 2.2, 2.0, 1.8, 1.8, 2.0],
    )
    add_bullets(
        doc,
        [
            "CR за 17–22: 12,1% · CPA ~141 ₽. CTR Поиска держится выше, чем на обучении (2,6%).",
            "После минусов п. 5.1 (вечер 20.08) две суток без конверсий и без расхода — стратегия не крутит мусор, но и ядро реже стреляет.",
            "Дубай / Лазаревское / прокат / ДР в 20–22.08 не всплыли. Минусы работают.",
            "Ещё кликают «снять яхту сочи» и «аренда яхты в сириусе» (конв. 0, расход 0). Слово «аренда» и «снять» добить в минусах, если ещё нет.",
        ],
    )

    add_heading(doc, "3. РСЯ за неделю — отдельно", 1)
    add_p(
        doc,
        "Решение 20.08: РСЯ не паузили, наблюдаем. Ниже — факт наблюдения.",
        bold=True,
    )
    add_heading(doc, "3.1. Цифры РСЯ", 2)
    add_table(
        doc,
        ["Период", "Показы", "Клики", "Конв.", "Расход", "CTR"],
        [
            ["14–20.08 (прошлый отчёт)", "245", "10", "0", "0 ₽", "4,1%"],
            ["Пн–Сб 17–22.08", "192", "3", "0", "0 ₽", "1,6%"],
            ["Только 20–22.08 (этот файл)", "110", "1", "0", "0 ₽", "0,91%"],
            ["16–22.08 (7 дней)", "260", "4", "0", "0 ₽", "1,5%"],
        ],
        [5.8, 2.0, 1.8, 1.8, 2.0, 2.0],
    )
    add_table(
        doc,
        ["День", "Показы РСЯ", "Клики", "Конв.", "Расход"],
        [
            ["16.08", "68", "1", "0", "0"],
            ["17.08", "41", "1", "0", "0"],
            ["18.08", "23", "1", "0", "0"],
            ["19.08", "18", "0", "0", "0"],
            ["20.08 полный", "50", "1", "0", "0"],
            ["21.08", "37", "0", "0", "0"],
            ["22.08", "23", "0", "0", "0"],
        ],
        [3.6, 3.2, 2.2, 2.2, 2.4],
    )
    add_heading(doc, "3.2. Как выглядит РСЯ", 2)
    add_bullets(
        doc,
        [
            "Конверсий нет ни за неделю, ни с запуска. При оплате за конверсии расход РСЯ = 0 ₽. Бюджет Поиска не ест.",
            "Показы падают: 50 → 37 → 23 за 20–22.08. Стратегия сама режет сети, раз заявок нет.",
            "CTR 20–22.08 = 0,91% — на нижней границе плана (≥ 0,8–1,2%). Клик 20.08: десктоп, 45–54, заголовок «линия 2», глубина 1, без цели.",
            "Форматы 20–22: ТГО 99 показов / 1 клик; видео 7 / 0; карусель 4 / 0.",
            "Аудитория РСЯ смещена: мужчины, 35+ и 55+, Android. Ядро Поиска — смартфоны 25–44.",
        ],
    )
    add_heading(doc, "3.3. РСЯ по площадкам (CSV 2026-08-22_21-49-50)", 2)
    add_p(
        doc,
        "Выгрузка 11–22.08: 56 площадок, 339 показов, 10 кликов, 0 конв., 0 ₽. Ни одна площадка в РК ещё не выключена.",
    )
    add_table(
        doc,
        ["Площадка", "Показы", "Клики", "Конв.", "Запретить?", "Оценка"],
        [
            ["yandex.ru", "139", "5", "0", "Нет", "Половина кликов РСЯ. Не банится."],
            ["dzen.ru", "63", "1", "0", "Да", "Дзен. В исключения."],
            ["m.images.yandex.ru", "19", "1", "0", "Да", "Картинки, 55+. Слабый интент."],
            ["tabor.ru", "4", "1", "0", "Да", "Клик 20.08. Развлекательный мусор."],
            ["ria.ru", "2", "1", "0", "Да", "Новости, не билет."],
            ["fooby.ru", "1", "1", "0", "Да", "Мусор, 1 клик."],
            ["com.vkontakte.android", "12", "0", "0", "Да", "VK-приложение."],
            ["video.yandex.ru", "10", "0", "0", "Да", "Видео."],
            ["ru.zen.android / zen", "8", "0", "0", "Да", "Дзен-приложение."],
            ["topwar.ru", "5", "0", "0", "Да", "Военные новости."],
            ["погода (gismeteo, pogoda, weather*)", "6", "0", "0", "Да", "Как в плане ЕПК — в исключения."],
        ],
        [5.2, 1.6, 1.5, 1.5, 2.0, 4.8],
    )
    add_p(doc, "Клики по дням:", bold=True)
    add_bullets(
        doc,
        [
            "14.08: yandex.ru, dzen.ru, картинки Яндекса, ещё yandex.ru (отказ 100%)",
            "15.08: yandex.ru (отказ 100%, 55+), fooby.ru",
            "16.08: ria.ru",
            "17–18.08: только yandex.ru (2 клика) — эту площадку запретить нельзя",
            "20.08: tabor.ru (десктоп 45–54, глубина 1) — единственный клик после 5.1",
            "21–22.08: кликов РСЯ нет",
        ],
    )
    add_p(doc, "В исключения РСЯ сейчас (не пауза всей сети):", bold=True)
    add_p(
        doc,
        "dzen.ru, ru.zen.android, com.yandex.zen, tabor.ru, fooby.ru, ria.ru, topwar.ru, topcor.ru, "
        "podolyaka.ru, pikabu.ru, popcornnews.ru, adme.media, 24smi.info, mom.life, motogonki.ru, "
        "gismeteo.ru, gismeteo.by, m.pogoda.yandex.ru, ru.yandex.weatherplugin, ru.yandex.mobile.weather, "
        "nur.kz, fontanka.ru, 93.ru, v102.ru, samaragovorit.ru, gazetametro.ru, tehnoomsk.ru, www1.ru",
        size=10,
    )
    add_p(
        doc,
        "Не трогать: yandex.ru / ya.ru (не банятся), sochi1.ru и sochi-fornia.ru (локальные). "
        "Avito / Otzovik — пока оставить, показов мало.",
    )
    add_heading(doc, "3.4. Вывод по РСЯ", 2)
    add_p(
        doc,
        "Пользы нет, вреда по деньгам нет. Всю РСЯ не паузим — наблюдаем. "
        "Имеет смысл точечно запретить Дзен, tabor, fooby, погоду и новостной/военный хвост. "
        "yandex.ru запретить нельзя: 5 из 10 кликов, все без заявок.",
        bold=True,
    )

    add_heading(doc, "4. Поиск 20–22.08", 1)
    add_table(
        doc,
        ["Место", "Показы", "Клики", "Конв.", "Расход", "CTR"],
        [
            ["Спецразмещение", "97", "8", "0", "0", "8,3%"],
            ["Карты", "45", "2", "0", "0", "4,4%"],
            ["Поиск / прочее", "27", "2", "1", "145", "7,4%"],
            ["Динамические места", "21", "2", "0", "0", "9,5%"],
            ["РСЯ", "110", "1", "0", "0", "0,9%"],
        ],
        [4.2, 2.0, 1.8, 1.8, 2.0, 1.8],
    )
    add_p(
        doc,
        "Спецразмещение кликает (CTR 8,3%), но без заявок за три дня. Единственная оплата — «прочее» (галерея/остаток) по запросу «выход в море сочи». "
        "Карты без конверсий; «водные прогулки» ещё показываются (15 показов, 0 кликов) — минус не ставили, верно.",
    )
    add_p(doc, "Клики без конверсии, которые добить минусом:", bold=True)
    add_bullets(
        doc,
        [
            "снять яхту в сочи / снять яхту сочи",
            "аренда яхты в сириусе",
            "неватрип (чужой бренд, альтернативные, 0 кликов — всё равно в минус)",
        ],
    )

    add_heading(doc, "5. Что делать", 1)
    add_bullets(
        doc,
        [
            "РСЯ целиком не паузить. В исключения — Дзен, tabor.ru, fooby.ru, погода, военные/новостные из списка 3.3.",
            "Минусы поиска: аренда, снять, неватрип. Не минусовать «выход в море» и «водные прогулки».",
            "Стратегию, бюджет 7 000, цены целей не менять.",
            "Следующий срез: полный 17–23.08 или 23–30.08.",
        ],
    )

    add_heading(doc, "6. Файлы", 1)
    add_bullets(
        doc,
        [
            "Этот отчёт: cursor/Analiz_RK_Gruppovye_progulki_EPK_2026-08-22.docx",
            "Копия: локальная/Analiz_RK_Gruppovye_progulki_EPK_2026-08-22.docx",
            "Предыдущий: Analiz_RK_Gruppovye_progulki_EPK_2026-08-20.docx",
        ],
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "Агент ЕПК групповых прогулок · 22.08.2026 · /progulki_na_yacht · 1 800 / 2 500 ₽ · +7 918 304-40-00"
    )
    set_run_font(r, size=9, color=RGBColor(0x66, 0x66, 0x66))

    doc.save(OUT)
    LOCAL_DIR.mkdir(parents=True, exist_ok=True)
    (LOCAL_DIR / OUT.name).write_bytes(OUT.read_bytes())
    print("wrote", OUT)


if __name__ == "__main__":
    build()
