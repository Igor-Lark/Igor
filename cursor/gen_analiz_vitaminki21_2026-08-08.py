#!/usr/bin/env python3
"""Full Yandex Direct CSV analysis -> Word doc."""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CSV_PATH = Path("/home/ubuntu/.cursor/projects/workspace/uploads/2026-08-08_11-24-39_vitaminki21_5641.csv")
OUT_PATH = Path("/workspace/cursor/Analiz_RK_vitaminki21_2026-08-08.docx")


def shade(cell, color="D9E2F3"):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(el)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    return h


def bullet(doc, text, bold=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold:
        r = p.add_run(bold)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i])
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def num(s):
  if pd.isna(s) or s in ("-", "", "—"):
    return 0.0
  if isinstance(s, str):
    s = s.replace(" ", "").replace(",", ".")
  try:
    return float(s)
  except Exception:
    return 0.0


def agg(df, group_cols):
    g = df.groupby(group_cols, dropna=False).agg(
        cost=("cost", "sum"),
        impr=("impr", "sum"),
        clicks=("clicks", "sum"),
        conv=("conv", "sum"),
    ).reset_index()
    g["ctr"] = np.where(g.impr > 0, g.clicks / g.impr * 100, 0)
    g["cr"] = np.where(g.clicks > 0, g.conv / g.clicks * 100, 0)
    g["cpc"] = np.where(g.clicks > 0, g.cost / g.clicks, 0)
    g["cpa"] = np.where(g.conv > 0, g.cost / g.conv, np.nan)
    return g.sort_values("cost", ascending=False)


def fmt_money(x):
    return f"{x:,.0f}".replace(",", " ") + " ₽"


def fmt_pct(x):
    return f"{x:.2f}%"


def top_table_rows(g, n=15, min_cost=0):
    g = g[g.cost >= min_cost].head(n)
    rows = []
    for _, r in g.iterrows():
        cpa = fmt_money(r.cpa) if r.conv > 0 and not np.isnan(r.cpa) else "—"
        rows.append([
            str(r.iloc[0]) if len(g.columns) == 6 else " / ".join(str(x) for x in r[group_cols].values),
            fmt_money(r.cost),
            int(r.impr),
            int(r.clicks),
            int(r.conv),
            fmt_pct(r.ctr),
            fmt_pct(r.cr),
            fmt_money(r.cpc),
            cpa,
        ])
    return rows


TABLE_HDR = ["Срез", "Расход", "Показы", "Клики", "Конв.", "CTR", "CR", "CPC", "CPA"]


def load():
    df = pd.read_csv(CSV_PATH, encoding="utf-8-sig")
    df = df[df["День"] != "Итого"].copy()
    rename = {
        "Расход, ₽": "cost",
        "Показы": "impr",
        "Клики": "clicks",
        "Конверсии": "conv",
        "Название группы": "group",
        "Тип площадки": "platform",
        "Вид размещения": "placement",
        "Тип устройства": "device",
        "Пол": "gender",
        "Возраст": "age",
        "Поисковый запрос": "query",
        "Категория запроса": "query_cat",
        "Тип условия показа": "cond_type",
        "Ключевая фраза": "keyword",
        "Заголовок": "title",
        "Тип объявления": "ad_type",
        "День": "day",
        "Уровень платежеспособности": "solvent",
        "Тип операционной системы": "os",
        "Ср. позиция показа": "pos_show",
        "Ср. позиция клика": "pos_click",
    }
    for k, v in rename.items():
        if k in df.columns:
            df[v] = df[k]
    for c in ["cost", "impr", "clicks", "conv", "pos_show", "pos_click"]:
        if c in df.columns:
            df[c] = df[c].apply(num)
    df["day"] = pd.to_datetime(df["day"], dayfirst=True, errors="coerce")
    return df


def minus_phrases(df):
    bad_patterns = [
        "дешев", "бесплат", "скидк", "акци", "промокод", "купить", "продаж",
        "работа", "ваканс", "зарплат", "кино", "театр", "музей", "аквапарк",
        "зоопарк", "банан", "парасейл", "гидроцикл", "jetski", "рыбал",
        "экскурс", "билет", "теплоход", "катамаран", "парусник", "погода",
        "камера", "онлайн", "метеор", "маршрут", "автобус", "поезд",
        "отель", "гостиниц", "санатор", "школ", "дет", "гид ",
        "где поесть", "ресторан", "кафе", "одеон", "звезд", "знаменит",
        "что посетить", "куда сходить", "развлечен", "погулять",
        "интересные места", "достопримечатель", "набережн",
    ]
    q = df[df["query"].notna() & (df["query"] != "-")].copy()
    qg = agg(q, ["query"])
    qg = qg[(qg.cost > 0) | (qg.clicks > 0)]
    minus = []
    for _, r in qg.iterrows():
        ql = str(r["query"]).lower()
        if r.conv > 0:
            continue
        if r.clicks >= 2 and r.cost >= 20:
            minus.append((r["query"], r.cost, r.clicks, "клики без конверсий"))
            continue
        for p in bad_patterns:
            if p in ql and r.cost > 0:
                minus.append((r["query"], r.cost, r.clicks, f"нерелевант: {p}"))
                break
    # dedupe
    seen = set()
    out = []
    for item in sorted(minus, key=lambda x: -x[1]):
        if item[0] not in seen:
            seen.add(item[0])
            out.append(item)
    return out[:40]


def main():
    df = load()
    total = {
        "cost": df.cost.sum(),
        "impr": df.impr.sum(),
        "clicks": df.clicks.sum(),
        "conv": df.conv.sum(),
    }
    total["ctr"] = total["clicks"] / total["impr"] * 100 if total["impr"] else 0
    total["cr"] = total["conv"] / total["clicks"] * 100 if total["clicks"] else 0
    total["cpc"] = total["cost"] / total["clicks"] if total["clicks"] else 0
    total["cpa"] = total["cost"] / total["conv"] if total["conv"] else 0

    doc = Document()
    t = doc.add_heading("Полный анализ рекламных кампаний", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("boat-sochi.ru · vitaminki21 · выгрузка 08.08.2026")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    period = f"{df.day.min().strftime('%d.%m.%Y')} — {df.day.max().strftime('%d.%m.%Y')}"
    doc.add_paragraph(
        f"Период данных: {period}. Строк в отчёте: {len(df):,}. "
        f"В выгрузке одна группа: «{df.group.iloc[0]}». Площадка: только Поиск."
    )

    heading(doc, "1. Сводка за период", 1)
    doc.add_paragraph(
        "За 5 дней (03–07.08.2026) кампания «Закат на море» в Поиске: "
        f"{int(total['clicks'])} кликов, {int(total['conv'])} конверсии, CPA {fmt_money(total['cpa'])}. "
        f"CTR {fmt_pct(total['ctr'])} — нормально для Поиска; CPC {fmt_money(total['cpc'])}. "
        "Основной «слив» — автотаргетинг по информационным запросам («куда сходить», «экскурсии», «банан»)."
    )
    add_table(doc, ["Метрика", "Значение"], [
        ["Расход", fmt_money(total["cost"])],
        ["Показы", f"{int(total['impr']):,}".replace(",", " ")],
        ["Клики", str(int(total["clicks"]))],
        ["Конверсии", str(int(total["conv"]))],
        ["CTR", fmt_pct(total["ctr"])],
        ["CR", fmt_pct(total["cr"])],
        ["CPC", fmt_money(total["cpc"])],
        ["CPA", fmt_money(total["cpa"]) if total["conv"] else "—"],
    ], [5, 11])

    heading(doc, "2. По группам объявлений", 1)
    g_groups = agg(df, ["group"])
    rows = []
    for _, r in g_groups.iterrows():
        cpa = fmt_money(r.cpa) if r.conv > 0 else "—"
        rows.append([r.group, fmt_money(r.cost), int(r.impr), int(r.clicks), int(r.conv), fmt_pct(r.ctr), fmt_pct(r.cr), fmt_money(r.cpc), cpa])
    add_table(doc, TABLE_HDR, rows, [4.5, 2, 1.5, 1.2, 1, 1.2, 1.2, 1.5, 1.5])

    heading(doc, "3. По дням", 1)
    g_day = agg(df, ["day"])
    g_day["day"] = g_day["day"].dt.strftime("%d.%m.%Y")
    rows = []
    for _, r in g_day.iterrows():
        cpa = fmt_money(r.cpa) if r.conv > 0 else "—"
        rows.append([r.day, fmt_money(r.cost), int(r.impr), int(r.clicks), int(r.conv), fmt_pct(r.ctr), fmt_pct(r.cr), fmt_money(r.cpc), cpa])
    add_table(doc, TABLE_HDR, rows)

    heading(doc, "4. Площадки и размещения", 1)
    for title, col in [("Тип площадки", "platform"), ("Вид размещения", "placement")]:
        heading(doc, title, 2)
        g = agg(df, [col])
        rows = []
        for _, r in g.iterrows():
            cpa = fmt_money(r.cpa) if r.conv > 0 else "—"
            rows.append([getattr(r, col), fmt_money(r.cost), int(r.impr), int(r.clicks), int(r.conv), fmt_pct(r.ctr), fmt_pct(r.cr), fmt_money(r.cpc), cpa])
        add_table(doc, TABLE_HDR, rows)

    heading(doc, "5. Устройства, пол, возраст", 1)
    for title, col in [("Устройство", "device"), ("Пол", "gender"), ("Возраст", "age"), ("Платежеспособность", "solvent")]:
        heading(doc, title, 2)
        g = agg(df, [col])
        rows = []
        for _, r in g.head(12).iterrows():
            cpa = fmt_money(r.cpa) if r.conv > 0 else "—"
            rows.append([getattr(r, col), fmt_money(r.cost), int(r.impr), int(r.clicks), int(r.conv), fmt_pct(r.ctr), fmt_pct(r.cr), fmt_money(r.cpc), cpa])
        add_table(doc, TABLE_HDR, rows)

    heading(doc, "6. Поисковые запросы", 1)
    qmask = df["query"].notna() & (df["query"] != "-")
    gq = agg(df[qmask], ["query"])
    gq_conv = gq[gq.conv > 0].head(10)
    gq_bad = gq[(gq.conv == 0) & (gq.clicks > 0)].head(20)

    heading(doc, "Конверсионные запросы", 2)
    if len(gq_conv):
        rows = []
        for _, r in gq_conv.iterrows():
            note = ""
            ql = str(r["query"]).lower()
            if any(x in ql for x in ["абхаз", "снять", "цена", "купить", "аренда без"]):
                note = " ⚠ в минус"
            rows.append([r["query"] + note, fmt_money(r.cost), int(r.clicks), int(r.conv), fmt_money(r.cpa)])
        add_table(doc, ["Запрос", "Расход", "Клики", "Конв.", "CPA"], rows, [6, 2.5, 1.5, 1.5, 2.5])
    else:
        doc.add_paragraph("Конверсионных поисковых запросов в выгрузке не зафиксировано (возможно, цели с других страниц/мессенджеров).")

    heading(doc, "Запросы с кликами без конверсий (кандидаты в минус)", 2)
    rows = []
    for _, r in gq_bad.iterrows():
        rows.append([r["query"], fmt_money(r.cost), int(r.clicks), fmt_pct(r.ctr)])
    add_table(doc, ["Запрос", "Расход", "Клики", "CTR"], rows[:20], [6, 2.5, 1.5, 2])

    heading(doc, "7. Автотаргетинг и категории запросов", 1)
    g_cat = agg(df[df.query_cat.notna() & (df.query_cat != "-")], ["query_cat"])
    rows = []
    for _, r in g_cat.iterrows():
        cpa = fmt_money(r.cpa) if r.conv > 0 else "—"
        rows.append([r["query_cat"], fmt_money(r.cost), int(r.impr), int(r.clicks), int(r.conv), fmt_pct(r.ctr), cpa])
    add_table(doc, ["Категория", "Расход", "Показы", "Клики", "Конв.", "CTR", "CPA"], rows)

    g_cond = agg(df, ["cond_type"])
    heading(doc, "Тип условия показа", 2)
    rows = []
    for _, r in g_cond.iterrows():
        cpa = fmt_money(r.cpa) if r.conv > 0 else "—"
        rows.append([r.cond_type, fmt_money(r.cost), int(r.impr), int(r.clicks), int(r.conv), fmt_pct(r.ctr), cpa])
    add_table(doc, ["Условие", "Расход", "Показы", "Клики", "Конв.", "CTR", "CPA"], rows)

    heading(doc, "8. Объявления (заголовки)", 1)
    ga = agg(df[df.title.notna() & (df.title != "-")], ["title"])
    rows = []
    for _, r in ga.head(15).iterrows():
        cpa = fmt_money(r.cpa) if r.conv > 0 else "—"
        rows.append([r.title[:60], fmt_money(r.cost), int(r.clicks), int(r.conv), fmt_pct(r.ctr), cpa])
    add_table(doc, ["Заголовок", "Расход", "Клики", "Конв.", "CTR", "CPA"], rows, [6, 2, 1.5, 1.2, 1.2, 2])

    heading(doc, "9. Позиции в поиске", 1)
    pos = df[df.platform == "Поиск"].copy()
    if len(pos):
        avg_show = pos[pos.pos_show > 0].pos_show.mean()
        avg_click = pos[pos.pos_click > 0].pos_click.mean()
        doc.add_paragraph(
            f"Средняя позиция показа (Поиск): {avg_show:.2f}. Средняя позиция клика: {avg_click:.2f}. "
            "Цель для коммерческих групп — удерживать клик в зоне 1–3 при приемлемом CPA."
        )
        gpos = agg(pos, ["group"])
        rows = []
        for _, r in gpos.iterrows():
            subp = pos[pos.group == r.group]
            ps = subp[subp.pos_show > 0].pos_show.mean()
            pc = subp[subp.pos_click > 0].pos_click.mean()
            rows.append([r.group, f"{ps:.2f}" if not np.isnan(ps) else "—", f"{pc:.2f}" if not np.isnan(pc) else "—", int(r.clicks), int(r.conv)])
        add_table(doc, ["Группа", "Поз. показа", "Поз. клика", "Клики", "Конв."], rows)

    # Recommendations
    heading(doc, "10. Рекомендации: что изменить", 1)

    recs = []
    # Analyze specific patterns
    zakat = g_groups[g_groups.group.str.contains("Закат", na=False)]
    tigger = g_groups[g_groups.group.str.contains("Tigger|Тиггер|тиггер", case=False, na=False)]
    delf = g_groups[g_groups.group.str.contains("дельфин|Дельфин", na=False)]
    rsya = df[df.platform == "Рекламная сеть Яндекса"]
    search = df[df.platform == "Поиск"]
    maps = df[df.platform == "Яндекс Карты"]

    if len(zakat):
        z = zakat.iloc[0]
        recs.append(
            f"Группа «Закат на море»: расход {fmt_money(z.cost)}, {int(z.clicks)} кликов, {int(z.conv)} конв. "
            "Автотаргетинг тянет нерелевант («куда сходить», «экскурсии», «банан») — отключить категорию «Альтернативные» и «Сопутствующие» или сузить до «Целевые»."
        )
    if len(tigger):
        t = tigger.iloc[0]
        recs.append(f"Группа Tigger: расход {fmt_money(t.cost)}, CPA {fmt_money(t.cpa) if t.conv else 'без конверсий'}. Проверить цели и качество лидов.")
    if len(delf):
        d = delf.iloc[0]
        recs.append(f"Дельфины: расход {fmt_money(d.cost)}, CTR {fmt_pct(d.ctr)}, конв. {int(d.conv)}.")

    auto_alt = agg(df[df.query_cat == "Альтернативные запросы"], ["query_cat"])
    auto_sop = agg(df[df.query_cat == "Сопутствующие запросы"], ["query_cat"])
    if len(auto_alt) and auto_alt.iloc[0].cost > 50 and auto_alt.iloc[0].conv == 0:
        recs.append("Категория автотаргетинга «Альтернативные запросы» — расход без конверсий. Отключить в группах Поиска.")
    if len(auto_sop) and auto_sop.iloc[0].cost > 100:
        recs.append("«Сопутствующие запросы» — много информационного трафика. Оставить только при достаточном бюджете; иначе отключить.")

    if len(rsya) and rsya.cost.sum() > 0:
        r = rsya.cost.sum(); c = rsya.conv.sum(); cl = rsya.clicks.sum()
        recs.append(f"РСЯ: {fmt_money(r)}, {int(cl)} кликов, {int(c)} конв. Если CPA высокий — снизить ставки или отключить слабые группы (рыбалка, отдых на яхте).")

    if total["cpa"] > 200:
        recs.append(f"Средний CPA {fmt_money(total['cpa'])} — при целевой цене контакта 150–200 ₽ ужесточить цели в стратегии или снизить долю автотаргетинга.")

    recs.extend([
        "В выгрузке только «Закат на море» (Поиск). Tigger/дельфины/РСЯ — в других кампаниях, смотреть отдельно.",
        "4 конверсии, но 2 запроса нерелевантны («яхты абхазия снять цены», широкое «сириус яхта») — добавить в минус и сузить автотаргетинг.",
        "Целевой запрос «прогулка на закат по морю для двоих» — масштабировать: отдельный ключ + заголовок под пару/романтику.",
        "Поиск: добавить минус-слова по информационным запросам (см. раздел 11).",
        "Для «Закат на море»: оставить ключи с «закат», «вечерняя прогулка», «яхта/катер + сириус/адлер»; убрать широкий автотаргетинг.",
        "Проверить посадочную #services — UTM zakat_my должен вести на блок заката с ценой и кнопкой.",
        "Карты: если расход есть без броней — снизить корректировку или отключить до стабилизации Поиска.",
        "Масштабировать заголовки/объявления с лучшим CTR и хотя бы 1 конверсией; слабые — на паузу.",
    ])

    for r in recs:
        bullet(doc, r)

    heading(doc, "11. Минус-фразы (добавить)", 1)
    doc.add_paragraph("Рекомендуется добавить на уровень кампании или группы «Закат на море»:")

    existing_minus = [
        "-!на крыше", "-аквапарк", "-анапа", "-без экипажа", "-вакансия", "-водопад",
        "-геленджик", "-зарплата", "-катер !в лизинг", "-купить", "-купить катер",
        "-лоо", "-новороссийск", "-продажа катера", "-продать", "-работа", "-ремонт",
        "-ресторан",
    ]
    bullet(doc, "Текущий список (сохранить): " + ", ".join(existing_minus))

    new_minus = [
        "-банан", "-парасейл", "-камера", "-онлайн", "-метеор", "-маршрут", "-автобус",
        "-кинотеатр", "-музей", "-театр", "-зоопарк", "-аквариум", "-одеон",
        "-дешев", "-бесплатн", "-скидк", "-акци", "-промокод", "-билет",
        "-экскурс", "-гид ", "-теплоход", "-катамаран", "-парусник",
        "-куда сходить", "-что посетить", "-достопримечатель", "-развлечен",
        "-набережн", "-погода", "-погулять", "-интересные места",
        "-где поесть", "-кафе", "-отель", "-гостиниц", "-звезд", "-знаменитост",
        "-абхаз", "-снять", "-цены", "-армения", "-грузия", "-крым", "-турци",
        "-школ", "-детск", "-гидроцикл", "-jetski", "-рыбал",
    ]
    for m in new_minus:
        bullet(doc, m)

    heading(doc, "Запросы из отчёта — в минус (точечно)", 2)
    mins = minus_phrases(df)
    if mins:
        rows = [[q, fmt_money(c), cl, reason] for q, c, cl, reason in mins[:25]]
        add_table(doc, ["Запрос", "Расход", "Клики", "Причина"], rows, [5, 2, 1.5, 5])
    else:
        doc.add_paragraph("Явных кандидатов по порогу кликов нет — ориентируйтесь на список выше.")

    heading(doc, "12. План на 7 дней", 1)
    plan = [
        "День 1–2: отключить «Альтернативные» в автотаргетинге; добавить минус-фразы из раздела 11.",
        "День 3–4: отчёт «Поисковые запросы» — точечные минуса; оставить 5–10 лучших ключей в «Закат».",
        "День 5: проверить CPA по группам; снизить цену цели там, где CPA > 250 ₽ без броней.",
        "День 6–7: A/B заголовков с «закат 18:00» и «Имеретинский порт»; масштабировать победителя.",
    ]
    for p in plan:
        bullet(doc, p)

    doc.add_paragraph()
    f = doc.add_paragraph(f"Файл: {OUT_PATH.name}")
    f.runs[0].italic = True
    doc.save(OUT_PATH)
    print(OUT_PATH)
    print("TOTAL", total)


if __name__ == "__main__":
    main()
