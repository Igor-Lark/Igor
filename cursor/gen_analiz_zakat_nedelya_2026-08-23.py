#!/usr/bin/env python3
"""Анализ РК «Закат на море» 20–23.08.2026 → Word."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Analiz_RK_Zakat_na_more_POISK_20-23_2026-08-23.docx"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else (13 if level == 2 else 12),
            bold=True,
            color=RGBColor(0x0B, 0x3D, 0x5C),
        )
    return p


def add_p(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if r_i % 2 == 1:
                shade_cell(cell, "F2F7FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("АНАЛИЗ КАМПАНИИ")
    set_run_font(r, size=12, bold=True, color=RGBColor(0x0B, 0x3D, 0x5C))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Закат на море (Поиск) — 20–23.08")
    set_run_font(r, size=18, bold=True, color=RGBColor(0x0B, 0x3D, 0x5C))

    for line in [
        "Группа 5779888441 · vitaminki21 · код ZK · только Поиск",
        "CSV 2026-08-23_19-00-15_vitaminki21 · 4 дня после чистки",
        "Стратегию, бюджет 12 000 и цели 140/130/110/120 не менять",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=10, color=RGBColor(0x44, 0x44, 0x44))

    add_heading(doc, "1. Короткий вердикт", 1)
    add_bullets(
        doc,
        [
            "20–23.08: 0 ₽ · 40 показов · 4 клика · 0 конв. · CTR 10%. Объём почти встал.",
            "Чужие города (Астрахань / Екб / Волга) в этом срезе не всплыли — минусы гео, похоже, сработали.",
            "Автотаргет всё ещё 36/3 из 40/4. Ключи-фразы: 4 показа / 1 клик.",
            "Единственный клик с ключа: «яхта на закате адлер» ← ключ «закат на яхте адлер -аренда». Минус -аренда всё ещё висит — снять.",
            "Два клика автотаргета — «аренда яхты сириус / в сириусе»: это дневной флот Сириуса 704503370, не ZK.",
            "Один свой клик автотаргета: «прогулки на катере с имеретинской набережной на закат» — оформить ключом.",
            "Новые ключи «закат в море адлер/сириус» ловят «во сколько закат» и «время заката» — не бронь. Минусовать.",
            "Стратегию не трогать. Сначала ключи ядра и минусы справочных запросов.",
        ],
    )

    add_heading(doc, "2. Три среза подряд", 1)
    add_table(
        doc,
        ["Метрика", "10–16.08", "16–21.08", "20–23.08"],
        [
            ["Расход", "182 ₽", "0 ₽", "0 ₽"],
            ["Показы", "555", "341", "40"],
            ["Клики", "80", "32", "4"],
            ["Конверсии", "1", "0", "0"],
            ["CTR", "14,4%", "9,4%", "10%"],
            ["Ср. позиция", "3,72", "4,04", "3,60"],
            ["Объём трафика", "67%", "65%", "52%"],
        ],
        col_widths=[4, 4, 4, 4],
    )
    add_table(
        doc,
        ["День", "Показы", "Клики", "Конв."],
        [
            ["20.08", "10", "1", "0"],
            ["21.08", "8", "2", "0"],
            ["22.08", "10", "0", "0"],
            ["23.08", "12", "1", "0"],
        ],
        col_widths=[4, 4, 4, 4],
    )
    add_p(
        doc,
        "После 17.08 альтернативные сжались, но ядро ключей не подставили — "
        "Директ почти перестал показывать. 10–12 показов в день мало для обучения.",
    )

    add_heading(doc, "3. Откуда трафик", 1)
    add_table(
        doc,
        ["Источник", "Показы", "Клики"],
        [
            ["Автотаргетинг", "36", "3"],
            ["Ключи-фразы", "4", "1"],
            ["Целевые (категория)", "29", "3"],
            ["Узкие", "8", "1"],
            ["Сопутствующие (категория запроса)", "3", "0"],
        ],
        col_widths=[8, 4, 4],
    )
    add_p(
        doc,
        "«Сопутствующие» здесь — класс запроса к ключу «закат в море…» "
        "(«во сколько закат», «время заката»), не обязательно включённый автотаргет. "
        "Динамические места 26/3, спецразмещение только 7/1.",
    )
    add_table(
        doc,
        ["Ключ", "Показы", "Клики"],
        [
            ["(автотаргет, пустой ключ)", "36", "3"],
            ["закат на яхте адлер -аренда", "1", "1"],
            ["закат в море сириус", "2", "0"],
            ["закат в море адлер", "1", "0"],
        ],
        col_widths=[8, 4, 4],
    )

    add_heading(doc, "4. Четыре клика — разбор", 1)
    add_table(
        doc,
        ["День", "Запрос", "Откуда", "Вердикт"],
        [
            ["20.08", "яхта на закате адлер", "ключ заката", "Своё. Усилить ключом без -аренда"],
            ["21.08", "прогулки на катере с имеретинской набережной на закат", "автотаргет, узкие", "Своё. Добавить ключ"],
            ["21.08", "аренда яхты сириус", "автотаргет", "Сириус 704503370, не ZK"],
            ["23.08", "аренда яхты в сириусе", "автотаргет", "То же — дневной флот"],
        ],
        col_widths=[2.2, 7.3, 3.5, 4],
    )

    add_heading(doc, "5. Минусы — добавить", 1)
    add_p(doc, "Справочные (ключ «закат в море» ловит погоду, не бронь):", bold=True)
    add_p(doc, '-время заката -"во сколько"')
    add_p(doc, "Чужой продукт / Сириус-день:", bold=True)
    add_p(doc, "-рыбалк -удочк -стеклянн -алиса")
    add_p(
        doc,
        "Гео-минусы прошлого раза в этом срезе не всплыли — не снимать. "
        "Не минусовать: закат, вечерняя, для двоих, целиком, аренда на закатных ключах.",
    )

    add_heading(doc, "6. Ключи — что сделать", 1)
    add_heading(doc, "6.1. Исправить", 2)
    add_p(
        doc,
        "Снять «-аренда» с «закат на яхте адлер -аренда». "
        "Один рабочий клик недели сидел именно на этой фразе. Аренда целиком — наш оффер.",
        bold=True,
    )

    add_heading(doc, "6.2. Добавить фразовым (оформить ключом)", 2)
    add_bullets(
        doc,
        [
            '"яхта на закате адлер" · "яхта на закате сочи"',
            '"прогулки на катере с имеретинской набережной на закат" — или короче: "катер на закат имеретинский"',
            '"вечерняя морская прогулка на яхте сириус" (2 показа)',
            '"вечерняя прогулка на яхте сочи"',
            '"покататься на яхте катере сочи вдвоем" → "яхта сочи вдвоем"',
            '"прогулка на закат по морю для двоих" — ядра из пакета всё ещё нет в ключах',
        ],
    )

    add_heading(doc, "6.3. Не масштабировать", 2)
    add_bullets(
        doc,
        [
            "морская прогулка сириус / морские прогулки сириус — SR 704503370",
            "аренда яхты сириус — то же, без слова «закат»",
            "время заката / во сколько закат — не покупатель",
        ],
    )

    add_heading(doc, "7. Объявления", 1)
    add_p(
        doc,
        "26 из 40 показов — «заголовок из сниппета» (динамические места). "
        "Свои заголовки «для двоих / целиком» почти не выходят: мало ключей, не плохие тексты. "
        "Тексты не трогать. Посадка всё ещё #zakat.",
    )

    add_heading(doc, "8. Чего не делать", 1)
    add_bullets(
        doc,
        [
            "Не менять стратегию, 12 000 ₽ и цены целей — за 4 дня 0 конверсий, поднимать CPA бессмысленно.",
            "Не включать РСЯ, галерею, альтернативные/широкие автотаргеты «чтобы было больше показов».",
            "Не минусовать «аренда» на закатных фразах.",
        ],
    )

    add_heading(doc, "9. План", 1)
    add_table(
        doc,
        ["Когда", "Действие"],
        [
            ["Сегодня", "Снять -аренда; минусы «время заката / во сколько / рыбалк / стеклянн»"],
            ["Сегодня", "Добавить ключи п. 6.2 фразовым"],
            ["3–5 дней", "Новая выгрузка: доля закат-кликов vs «аренда яхты сириус»"],
            ["После 10 конв./нед", "Тогда цели 160–180 ₽, не раньше"],
        ],
        col_widths=[4, 13],
    )

    add_heading(doc, "10. Источники", 1)
    add_bullets(
        doc,
        [
            "CSV: 2026-08-23_19-00-15_vitaminki21.csv (группа 5779888441, 20–23.08)",
            "Предыдущие: неделя 10–16 и ключи 16–21.08",
            "Память: cursor/Zakat_campaign_memory.md · MP 712465896 снята",
        ],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
