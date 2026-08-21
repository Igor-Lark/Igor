#!/usr/bin/env python3
"""Анализ РК «Яхта Tigger» срез 16–21.08.2026 → Word на печать."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Analiz_RK_Tigger_srez_2026-08-16_21_na_pechat.docx"
NAVY = (0x1F, 0x4E, 0x79)
GRAY = (0x59, 0x59, 0x59)


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_p(doc, text, size=11, bold=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.12
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def fill_cell(cell, text, bold=False, size=10, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold or header, color=(255, 255, 255) if header else None)
    set_cell_border(cell)
    if header:
        shade_cell(cell, "1F4E79")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, header=True, size=10)
    for r_i, row in enumerate(rows):
        prevent_row_split(table.rows[r_i + 1])
        for c_i, val in enumerate(row):
            fill_cell(table.rows[r_i + 1].cells[c_i], val, size=10, bold=c_i == 0)
            if r_i % 2 == 1:
                shade_cell(table.rows[r_i + 1].cells[c_i], "F2F2F2")
    if col_widths:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblW = OxmlElement("w:tblW")
        total = int(sum(col_widths) * 567)
        tblW.set(qn("w:w"), str(total))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*NAVY)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run(run, size=size)


def add_page_field(paragraph):
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    set_run(run, size=9, color=GRAY)


def setup_print(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Яхта Tigger  ·  Яндекс Директ  ·  срез 16–21.08.2026  ·  на печать")
    set_run(run, size=9, color=GRAY)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Внутренний анализ РК  ·  страница ")
    set_run(run, size=9, color=GRAY)
    add_page_field(fp)
    run = fp.add_run("  ·  файл Analiz_RK_Tigger_srez_2026-08-16_21_na_pechat.docx")
    set_run(run, size=9, color=GRAY)


MINUS_NEW = [
    "группы",
    "прогулки",
    "прогулка",
    "круизы",
    "taurica",
    "yachtjoy",
    "алисия",
    "хармони",
    "снафу",
    "внжела",
    "shalawaty",
    "азимут ивушка",
    "азимут 40",
    "бассейн азимут",
    "для двоих",
]


def build():
    doc = Document()
    setup_print(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_p(
        doc,
        "Анализ рекламной кампании «Яхта Tigger»",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_p(
        doc,
        "Срез вс 16.08 — пт 21.08.2026  ·  календарная неделя 17–23.08 ещё не закрыта",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_p(
        doc,
        "Кампания в отчёте: «Поиск+Карты - Яхта Tigger»  ·  CSV 2026-08-21_21-41-10_vitaminki21",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    heading(doc, "1. Короткий вердикт", 1)
    bullet(
        doc,
        "Первая конверсия: 17.08, 1 шт., расход 137,50 ₽. Запрос «туры на частной яхте адлер», "
        "М 18–24, платёжеспособность «Остальные», глубина 1. Не раздувать как VIP-успех, "
        "пока не видно цель в Метрике.",
    )
    bullet(
        doc,
        "Минусы 16.08 сработали: автотаргет 104→28 показов, мусорных «целевых» меньше, CTR 9,9%→12%.",
    )
    bullet(
        doc,
        "Объём просел. Файл 16–21: 100 показов / 12 кликов. Пн–пт 17–21: 69 / 10 / 1 конв. "
        "Прошлая полная неделя пн–вс 10–16: 232 / 23 / 0.",
    )
    bullet(
        doc,
        "Главный риск: в CSV 60 показов и 6 кликов висят на объявлениях со статусом «Остановленные». "
        "Сначала проверить, что кампания/группа/объявления включены — иначе низкий объём не про CPA.",
    )
    bullet(
        doc,
        "CPA, бюджет, РСЯ, каюты — не трогать. Доминусовать новые дыры столбиком ниже. "
        "Полная неделя — выгрузка вс 23.08.",
    )

    heading(doc, "2. Цифры: прошлое vs этот срез", 1)
    add_table(
        doc,
        ["Метрика", "Пн–вс 10–16.08", "Вс–пт 16–21.08", "Пн–пт 17–21.08"],
        [
            ["Показы", "232", "100", "69"],
            ["Клики", "23", "12", "10"],
            ["CTR", "9,91%", "12,00%", "14,5%"],
            ["Конверсии", "0", "1", "1"],
            ["Расход", "0 ₽", "137,50 ₽", "137,50 ₽"],
            ["CPA", "—", "137,50 ₽", "137,50 ₽"],
            ["Отказы", "27%", "18%", "—"],
            ["Глубина", "1,14", "1,18", "—"],
            ["Ср. позиция", "~3,1–3,3", "~3,4", "—"],
            ["Объём трафика", "~80%", "83%", "—"],
        ],
        col_widths=[3.8, 4.4, 4.6, 4.8],
    )
    add_p(
        doc,
        "16.08 в этом файле — хвост прошлой календарной недели (после выгрузки 19:00). "
        "Сравнивать «неделя к неделе» можно будет только после вс 23.08. "
        "Оплата за конверсии: 137,50 ₽ ниже потолка целей (250–300 / профиль 210) — Директ списал меньше капа.",
    )

    heading(doc, "2.1. По дням", 2)
    add_table(
        doc,
        ["День", "Показы", "Клики", "Конв.", "Расход", "Заметка"],
        [
            ["Вс 16.08", "31", "2", "0", "0", "Хвост прошлой недели; объявления уже «остановленные»"],
            ["Пн 17.08", "28", "3", "1", "137,50", "Единственная конверсия"],
            ["Вт 18.08", "18", "5", "0", "0", "Пик кликов, CTR высокий"],
            ["Ср 19.08", "8", "1", "0", "0", "Спад"],
            ["Чт 20.08", "3", "0", "0", "0", "Почти ноль"],
            ["Пт 21.08", "12", "1", "0", "0", "Отскок слабый"],
        ],
        col_widths=[2.8, 2.0, 1.8, 1.8, 2.2, 7.0],
    )

    heading(doc, "3. Первая конверсия — разбор", 1)
    add_table(
        doc,
        ["Поле", "Значение"],
        [
            ["Дата", "17.08.2026, пн"],
            ["Запрос", "туры на частной яхте адлер"],
            ["Категория / условие", "Широкие · Фраза · Спецразмещение"],
            ["Заголовок", "Аренда яхты премиум-класса «Tigger» в Сочи. От 2-х часов"],
            ["Аудитория", "Мужской · 18–24 · Остальные"],
            ["Поведение", "Отказ 0% · глубина 1,00"],
            ["Статус объявления", "Активные"],
            ["Списание", "137,50 ₽"],
        ],
        col_widths=[5.0, 12.6],
    )
    add_p(
        doc,
        "«Частной яхте» близко к private charter — интент не худший. Но 18–24 и «Остальные» "
        "слабо бьются с чеком от 100 000 ₽ / 2 часа. Глубина 1 = клик сразу в цель, без просмотра яхты. "
        "Не минусовать «туры» из‑за этой одной конверсии. Не поднимать CPA «потому что обучились». "
        "В Метрике 94713538 проверить: какая цель 17.08 и был ли контакт.",
        bold=True,
    )

    heading(doc, "4. Что изменилось после минусов 16.08", 1)
    add_table(
        doc,
        ["Срез", "10–16.08", "16–21.08", "Вывод"],
        [
            ["Автотаргет, показы/клики", "104 / 5", "28 / 1", "Чище, оставить целевые"],
            ["Целевые (категория Директа)", "97 / 3", "23 / 1", "Меньше ложных, дыры остались"],
            ["Альтернативные", "23 / 2", "11 / 0", "Ещё крутятся — перепроверить выкл."],
            ["Сопутствующие", "3 / 0", "5 / 0", "Азимут-отель/бассейн — минусовать"],
            ["Фраза", "128 / 18", "72 / 11 + 1 конв.", "Рабочее ядро"],
            ["Карты", "16 / 0", "5 / 0", "Без кликов, не выключать"],
        ],
        col_widths=[5.2, 3.6, 3.8, 5.0],
    )
    add_p(
        doc,
        "16.08 вы сказали, что альтернативные и сопутствующие выключены. В срезе они ещё дают показы "
        "(в т.ч. после 17.08: азимут 40, азимут ивушка/бассейн, прогулка имеретинка). "
        "Открыть группу → автотаргетинг и убедиться, что галки сняты, не только на уровне рекомендации.",
    )

    heading(doc, "4.1. Объявления «Остановленные» — проверить в Директе", 2)
    add_table(
        doc,
        ["Статус в CSV", "Показы", "Клики", "Конв.", "Расход"],
        [
            ["Остановленные", "60", "6", "0", "0 ₽"],
            ["Активные", "40", "6", "1", "137,50 ₽"],
        ],
        col_widths=[4.5, 3.2, 3.0, 3.0, 3.9],
    )
    add_p(
        doc,
        "Среди остановленных были сильные запросы: «премиум вип», «люксовую яхту 56 метров» (отказ 100% — "
        "другая длина), «на день», «круизы». Если это автостоп комбинаторных объявлений — норма. "
        "Если на паузе кампания или группа — включить: иначе бюджет 10 000 ₽/нед не на что учиться. "
        "Название кампании в отчёте сменилось на «Поиск+Карты - Яхта Tigger» — тоже сверить, что смотрим ту же РК.",
    )

    heading(doc, "5. Клики: что оставлять, что резать", 1)
    add_p(doc, "Оставлять (ядро VIP / private):", bold=True, space_after=3)
    add_table(
        doc,
        ["Запрос", "Кто", "Вердикт"],
        [
            ["аренда яхты с гидроциклом сочи", "Ж 25–34, 6–10%, отказ 0%", "Ключ. Не минусовать"],
            ["аренда яхты на неделю сириус", "М 35–44, 6–10%, глуб. 2", "Неделя — ок"],
            ["большая яхта … премиум вип", "М 35–44, 2–5%, отказ 0%", "VIP-интент"],
            ["аренда яхты морской порт сочи", "М 25–34, 2–5%, глуб. 2", "Гео причала"],
            ["яхта тайгер сочи снять", "автотаргет, без клика", "Опечатка Tigger — оставить"],
        ],
        col_widths=[6.4, 6.2, 5.0],
    )
    add_p(doc, "Мусор / не Tigger (новые дыры после минусов 16.08):", bold=True, space_after=3)
    add_table(
        doc,
        ["Запрос", "Показы/клики", "Почему"],
        [
            ["сочи прогулки … на яхте группы", "2 / 1", "Групповой продукт, не 50 000 ₽/ч"],
            ["яхта taurica / yachtjoy / алисия / хармони / снафу / внжела / shalawaty", "пачками 1–3", "Чужие суда, автотаргет"],
            ["азимут ивушка / бассейн азимут / азимут 40", "2+1+2", "Отель и другая модель"],
            ["яхта для двоих в сириусе", "3 / 0", "Дешёвый «для двоих»"],
            ["круизы на яхте из сочи", "1 клик", "Круиз ≠ private charter"],
            ["арендовать люксовую яхту 56 метров", "1 клик, отказ 100%", "Tigger 21,6 м, не 56"],
        ],
        col_widths=[7.4, 3.4, 6.8],
    )

    heading(doc, "6. Рекомендации — приоритет", 1)
    heading(doc, "6.1. Сегодня", 2)
    add_p(doc, "A. В Директе открыть кампанию и проверить статус. Если на паузе — включить. Каюты не трогать.", bold=True)
    add_p(doc, "B. Доминусовать только новое (столбик, копировать в Директ):", bold=True, space_after=3)
    for phrase in MINUS_NEW:
        bullet(doc, phrase)
    add_p(
        doc,
        "Не минусовать: туры, гидроцикл, тайгер, tigger, vip, премиум, азимут (без 40/бассейн/ивушка), "
        "на неделю, на ночь, закат.",
        space_before=4,
    )
    add_p(
        doc,
        "C. Перепроверить автотаргет группы: альтернативные и сопутствующие должны быть выкл. "
        "«Целевые» оставить, минусами закрыть чужие яхты из столбика.",
    )

    heading(doc, "6.2. Не делать до вс 23.08", 2)
    bullet(doc, "Не поднимать CPA (тел 300 / TG 250 / MAX 300 / профиль 210).")
    bullet(doc, "Не включать РСЯ и галерею услуг.")
    bullet(doc, "Не менять тип стратегии из‑за одной конверсии и 0/10 обучения.")
    bullet(doc, "Не править заголовок «3 каюты» — решение 16.08.")
    bullet(doc, "Не минусовать «туры» из‑за конверсии 17.08.")

    heading(doc, "7. План до конца календарной недели", 1)
    add_table(
        doc,
        ["Когда", "Действие"],
        [
            ["Пт 21.08", "Статус объявлений + минусы из §6.1 столбиком"],
            ["Сб–вс 22–23.08", "Не трогать стратегию. Чистить новые запросы, если всплывут"],
            ["Вс 23.08 вечером", "Выгрузка пн–вс 17–23.08 — тогда сравнивать неделю к неделе"],
            ["После 23.08", "Если объявления живые, ядро чистое, а VIP-показов мало — тогда разговор про CPA +15–20%"],
        ],
        col_widths=[4.2, 13.4],
    )

    heading(doc, "8. Источники", 1)
    bullet(doc, "CSV: 2026-08-21_21-41-10_vitaminki21.csv · 99 строк + итого · 16–21.08.2026.")
    bullet(doc, "Сравнение: Analiz_RK_Tigger_nedelya_2026-08-10_16_na_pechat.docx.")
    bullet(doc, "Память: cursor/Tigger_campaign_memory.md · минусы столбиком: RK_agents_conventions.md.")

    add_p(
        doc,
        "Документ для печати, A4, внутренний. Не для клиента.",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
        space_after=0,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
