#!/usr/bin/env python3
"""Баннер на поиске — групповые прогулки. Word: как создать РК в Директе."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Banner_Poisk_Gruppovye_Progulki_2026-08-27.docx"

PHONE = "+7 918 304-40-00"
LANDING = "https://boat-sochi.ru/progulki_na_yacht"
COUNTER_SITE = "94713538"
COUNTER_ORG = "103116887"
UTM = "utm_source=yandex&utm_medium=cpc&utm_campaign=banner_poisk_gruppovye"
AD_URL = (
    f"{LANDING}?{UTM}&utm_content={{ad_id}}&utm_term={{keyword}}"
)
GIF_REPO = "cursor/banners-search/gp-240x400/gp-search-240x400.gif"
GIF_RAW = (
    "https://raw.githubusercontent.com/Igor-Lark/Igor/"
    "cursor/banner-na-poiske-4385/cursor/banners-search/gp-240x400/gp-search-240x400.gif"
)
NAVY = RGBColor(0x0B, 0x3D, 0x5C)


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13,
            bold=True,
            color=NAVY,
        )
    return p


def add_p(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)


def add_column(doc, items):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if r_i % 2 == 1:
                shade_cell(cell, "F2F7FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


KEYS = [
    '"групповая прогулка на яхте сочи"',
    '"групповая прогулка на яхте сириус"',
    '"групповые прогулки на яхте сочи"',
    '"билет на яхту сочи"',
    '"билет на яхту сириус"',
    '"прогулка на яхте за человека"',
    '"прогулка на яхте за человека сочи"',
    '"прогулки под парусом сириус"',
    '"прогулка под парусом сочи"',
    '"парусная прогулка сириус"',
    '"парусная прогулка сочи"',
    '"место на яхте сочи"',
    '"групповая морская прогулка сириус"',
    '"яхта сириус групповая"',
]

DO_NOT_MINUS = [
    "группа",
    "билет",
    "за человека",
    "1800",
    "парус",
    "парусная",
    "до 11",
    "расписание",
    "купание",
]

MINUS = [
    "аренда яхты целиком",
    "аренда катера",
    "аренда яхты",
    "прокат катера",
    "прокат яхты",
    "снять яхту",
    "снять катер",
    "яхта сочи аренда",
    "катер на час",
    "на час",
    "целиком",
    "для двоих",
    "вип",
    "vip",
    "премиум",
    "алексум",
    "tigger",
    "тиггер",
    "сутки",
    "на сутки",
    "круиз",
    "рыбалка",
    "трофейн",
    "гидроцикл",
    "теплоход",
    "7500",
    "частная аренда",
    "бизнес класс",
    "яхт клуб",
    "работа",
    "вакансия",
    "купить яхту",
    "продажа яхты",
    "чертеж",
    "своими руками",
    "симулятор",
    "игра",
    "бесплатно",
    "петрозаводск",
    "калининград",
    "крым",
    "севастополь",
    "абхаз",
    "абхазия",
    "грузия",
    "турци",
    "дельфинарий",
    "океанариум",
    "шоу дельфинов",
    "плавать с дельфинами",
]


def set_margins(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(1.6)
        sec.bottom_margin = Cm(1.6)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)


def build():
    doc = Document()
    set_margins(doc)

    add_heading(doc, "Баннер на поиске — групповые прогулки", 0)
    add_p(
        doc,
        "Яндекс Директ · vitaminki21 · 27.08.2026 · код BP-GP. "
        "Это не текстовая ЕПК и не HTML5 для РСЯ.",
        size=11,
    )

    add_heading(doc, "0. Зачем отдельная кампания", 1)
    add_p(
        doc,
        "Контекстный баннер на Поиске — отдельный тип РК. "
        "В текстовую ЕПК групповых его не добавить. "
        "Рыбалку, катер и закат сюда группами не класть: бюджет и стратегия общие, "
        "Директ съест более дешёвую тему. На каждую тему — своя баннерная РК. "
        "Сейчас только групповые.",
    )
    add_table(
        doc,
        ["", "Текстовая GP", "Этот баннер BP-GP"],
        [
            ["Тип", "ЕПК Поиск + РСЯ + Карты", "Баннер на Поиске"],
            ["Где видно", "Блоки текста", "Справа от выдачи, десктоп"],
            ["Файл", "Заголовки и 8 быстрых ссылок", "GIF/JPG/PNG 240×400 ≤120 КБ"],
            ["Посадка", LANDING, LANDING],
            ["Телефон", f"{PHONE} Наталья", "На сайте, на баннере нет"],
            ["Бюджет/нед", "10–12 тыс. ₽", "4–5 тыс. ₽ отдельно"],
            ["Галерея услуг", "ВЫКЛ", "Формат не про галерею"],
        ],
        col_widths=[4.2, 6.2, 6.4],
    )

    add_heading(doc, "1. Что создать в кабинете", 1)
    add_table(
        doc,
        ["Поле", "Значение"],
        [
            ["Название", "Banner_Poisk_Gruppovye_Progulki"],
            ["Режим", "Эксперт"],
            ["Цель мастера", "Заметность на Поиске Яндекса"],
            ["Если нет этой цели", "Контекстный баннер в Поиске"],
            ["Стратегия", "Оптимизация кликов"],
            ["Средняя цена клика", "80 ₽"],
            ["Недельный бюджет", "4 000 ₽ (можно 5 000)"],
            ["Регион", "Сочи, Сириус, Адлер (агломерация)"],
            ["Время", "08:00–22:00, Москва"],
            ["Счётчики Метрики", f"{COUNTER_SITE} и {COUNTER_ORG}"],
            ["Автотаргет", "Нет / не включать"],
            ["Параметры URL кампании", "Пусто — UTM в ссылке объявления"],
        ],
        col_widths=[5.5, 11.3],
    )
    add_p(
        doc,
        "Официальные стратегии этого типа — оптимизация кликов и ручные ставки. "
        "Не выбирай «максимум конверсий» / ЕПК: это другой мастер, туда GIF 240×400 не грузится.",
    )

    add_heading(doc, "2. Клики в Директе — кампания", 1)
    add_numbered(
        doc,
        [
            "direct.yandex.ru → аккаунт vitaminki21.",
            "Добавить → Кампанию.",
            "Режим эксперта. Цель: «Заметность на Поиске Яндекса» "
            "(или «Контекстный баннер в Поиске», если так подписано).",
            "Название: Banner_Poisk_Gruppovye_Progulki. Дата начала — сегодня, окончание не ставь.",
            "Регион: Сочи + Сириус + Адлер. Не вся Россия.",
            "Стратегия: Оптимизация кликов. Средняя цена клика 80 ₽. Недельный бюджет 4000 ₽.",
            "Расписание: 08:00–22:00, часовой пояс Москва.",
            "Метрика: оба счётчика 94713538 и 103116887.",
            "Блок параметров URL кампании не заполняй.",
            "Сохранить и перейти к группе.",
        ],
    )

    add_heading(doc, "3. Клики — группа", 1)
    add_p(doc, "Одна группа на старт. Название: G1 · Билет парус 1800", bold=True)
    add_numbered(
        doc,
        [
            "Регион группы — как у кампании (не шире).",
            "Ключевые фразы — вставь столбик из раздела 5. Кавычки не снимай: это фразовое совпадение.",
            "Минус-фразы группы — столбик из раздела 7. По одной строке.",
            "Автотаргет, если поле есть, — выкл.",
            "Корректировки на старте не ставь.",
            "Сохранить и перейти к объявлению.",
        ],
    )

    add_heading(doc, "4. Клики — объявление (креатив)", 1)
    add_numbered(
        doc,
        [
            "В одно объявление — один файл. Не ZIP, не HTML5.",
            "Скачай GIF по ссылке ниже. Проверь: 240×400 и меньше 120 КБ (сейчас ~61 КБ).",
            "Загрузи файл с компьютера.",
            "Ссылка в объявлении — целиком из раздела 8 (посадка + UTM).",
            "Телефон, WhatsApp и быстрые ссылки в этом формате не добавляются — так и должно быть.",
            "Отправить на модерацию. Показы после оплаты.",
        ],
    )
    add_p(doc, "Креатив GIF (скачать):", bold=True)
    add_p(doc, GIF_RAW, size=10)
    add_p(doc, f"В репозитории: {GIF_REPO}")
    add_bullets(
        doc,
        [
            "На баннере уже есть: Сириус · причал 2, прогулки под парусом, 1 800 ₽/чел, "
            "ротация «до 11 / 1,5 часа / купание / заявка», кнопка Расписание.",
            "Не ставь второй креатив с «аренда яхты», дельфинами «гарантируем» или wa.me.",
            "Имя яхты на баннере не писать.",
        ],
    )

    add_heading(doc, "5. Ключевые фразы — копировать столбиком", 1)
    add_p(
        doc,
        "Только групповой интент. Широкие «морские прогулки», «прогулки на яхте сочи», "
        "«яхта сочи аренда» — в РК катера (SR), сюда не ставить.",
    )
    add_column(doc, KEYS)
    add_p(doc, "")

    add_heading(doc, "6. Не минусовать", 1)
    add_p(doc, "Это ядро билета. Если минусанёшь — баннер ослепнет.")
    add_column(doc, DO_NOT_MINUS)
    add_p(doc, "")

    add_heading(doc, "7. Минус-фразы — копировать столбиком", 1)
    add_column(doc, MINUS)
    add_p(doc, "")

    add_heading(doc, "8. Ссылка объявления", 1)
    add_p(doc, "Вставь как есть, фигурные скобки Директ подставит сам:", bold=True)
    add_p(doc, AD_URL, size=10)
    add_p(
        doc,
        "Посадка только /progulki_na_yacht. Не главная, не /delfin, не рыбалка, не Tigger. "
        "Телефон на лендинге: " + PHONE + ".",
    )

    add_heading(doc, "9. После запуска", 1)
    add_bullets(
        doc,
        [
            "Пришли агенту BP номер кампании — запишет в память.",
            "Неделя Пн–Вс: CSV поисковых запросов + скрин стратегии и бюджета.",
            "Минусы по факту запросов — столбиком. Стратегию и 80 ₽ без отдельного «да» не менять.",
            "Если кликов мало — не расширяй ключи до «морские прогулки». Лучше ставка/бюджет, чем чужой интент.",
            "Текстовую GP не останавливай: баннер её не кормит на мобильных.",
            "Рыбалка и катер — новые баннерные РК, не новые группы здесь.",
        ],
    )

    add_heading(doc, "10. Чего не делать", 1)
    add_bullets(
        doc,
        [
            "Не грузить HTML5 ZIP из РСЯ в этот тип.",
            "Не ставить на баннер телефон и wa.me — ломает конверсию в Директе.",
            "Не обещать дельфинов и не называть яхту.",
            "Не включать галерею услуг в текстовой GP «за компанию».",
            "Не мешать в одной баннерной РК прогулки и рыбалку.",
        ],
    )

    add_p(
        doc,
        "Агент BP · ветка cursor/banner-na-poiske-4385 · пакет 27.08.2026",
        size=9,
        space_after=0,
    )

    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build()
