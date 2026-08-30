#!/usr/bin/env python3
"""Пакетная стратегия RB + групповые прогулки (GP) для обучения."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Paketnaya_strategiya_RB_GP_Popaj_2026-08-30.docx"
OUT_LOCAL = Path(__file__).resolve().parent.parent / "локальная" / OUT.name

NAVY = RGBColor(0x0B, 0x3D, 0x5C)
GRAY = RGBColor(0x33, 0x33, 0x33)
RED = RGBColor(0x8B, 0x1A, 0x1A)


def set_run(run, *, bold=False, size=11, color=GRAY, font="Arial"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font


def add_p(doc, text, *, bold=False, size=11, color=GRAY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    p.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=True, size=16 if level == 1 else 13, color=NAVY)
    p.paragraph_format.space_before = Pt(12)


def add_bullets(doc, items, *, color=GRAY):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        set_run(run, size=11, color=color)


def build():
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Пакетная стратегия · RB + групповые прогулки")
    set_run(run, bold=True, size=18, color=NAVY)
    add_p(doc, "План пользователя 30.08.2026 · vitaminki21 · boat-sochi.ru", size=10)

    add_heading(doc, "1. Зачем")
    add_p(
        doc,
        "RB застряла на обучении (0/10 конверсий). Пакетная стратегия объединяет несколько ЕПК "
        "в один «пакет» — алгоритм учится на конверсиях **всех** кампаний пакета, "
        "быстрее набирает порог 10 конв./нед.",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "RB — билет 2 800 ₽, рыбалка, /gruppovaja_ribbalka",
            "GP (групповые прогулки) — билет 1 800 / закат 2 500 ₽, /progulki_na_yacht",
            "Общее: один телефон +7 918 304-40-00, один счётчик 94713538, билетный формат, Сириус линия 2",
        ],
    )

    add_heading(doc, "2. Условия — когда имеет смысл")
    add_bullets(
        doc,
        [
            "GP **уже запущена** и даёт конверсии (хотя бы 2–3 в неделю) — иначе пакет не ускорит RB.",
            "Обе РК — **максимум конверсий**, оплата **за конверсии**, схожие цели (телефон / MAX / форма).",
            "После перевода в пакет обе РК **снова идут в обучение** — это нормально, закладывать 1–2 недели.",
        ],
    )

    add_heading(doc, "3. Как включить (кабинет)")
    add_bullets(
        doc,
        [
            "Директ → кампания RB (713632237) → **Стратегия**.",
            "Тип: **Пакетная** (вместо «Обычная»).",
            "Выбрать в пакет: **ЕПК «Групповые прогулки»** (уточнить номер GP в кабинете).",
            "Проверить: общий бюджет пакета ≥ сумма бюджетов (RB 8 000 + GP 10–12 000 ₽/нед).",
            "Цели и цены конверсий — **не менять** в день переключения (или менять только вместе, один раз).",
            "GP тоже перевести на **пакетную** и указать тот же пакет.",
        ],
    )

    add_heading(doc, "4. Бюджет пакета (ориентир)")
    add_bullets(
        doc,
        [
            "RB: **8 000 ₽ / нед** (как сейчас)",
            "GP: **10 000–12 000 ₽ / нед** (из handoff групповых прогулок)",
            "Пакет суммарно: **18 000–20 000 ₽ / нед** — алгоритм сам распределит между RB и GP",
            "Не резать GP до нуля — иначе пакет «ослепнет» для RB",
        ],
    )

    add_heading(doc, "5. Антиканнибализация (обязательно)")
    add_p(doc, "Пакет ≠ разрешение смешивать ключи. Минусы оставить:", bold=True)
    add_bullets(
        doc,
        [
            "В GP: −рыбалка −улов −барабулька −попай −морская рыбалка",
            "В RB: −групповая прогулка −парусная −1800 −2500 −дельфинарий −прогулка на яхте (билет)",
            "Посадки разные: RB → /gruppovaja_ribbalka · GP → /progulki_na_yacht",
        ],
        color=RED,
    )

    add_heading(doc, "6. Что сделать ДО переключения")
    add_bullets(
        doc,
        [
            "☐ Записать номер кампании GP в Директе",
            "☐ GP — проверить, что идут конверсии за последнюю неделю",
            "☐ RB — закрыть архивное G1 (новое активное объявление)",
            "☐ RB — G10 без текстов «барабулька/пеламида» (качество сигналов)",
            "☐ G12 **не заводить** — пеламида в G5",
            "☐ Скрин стратегии RB и GP **до** и **после** — для памяти агента",
        ],
    )

    add_heading(doc, "7. Что НЕ делать")
    add_bullets(
        doc,
        [
            "Не добавлять в пакет VIP/Tigger/аренду целиком — другой продукт и CPA",
            "Не менять в один день: пакет + бюджет + цены целей + тексты всех групп",
            "Не ждать, что RB сразу получит конверсии — нужен объём GP в пакете",
        ],
        color=RED,
    )

    add_heading(doc, "8. После переключения — мониторинг")
    add_bullets(
        doc,
        [
            "7–10 дней: смотреть конверсии **пакета** и **отдельно RB / GP**",
            "Если GP забирает весь бюджет без RB — поднять минимальный недельный бюджет RB или CPA RB чуть ниже GP (осторожно)",
            "Если RB пошли конверсии — не трогать 2 недели",
            "Выгрузка запросов RB отдельно — как обычно",
        ],
    )

    add_heading(doc, "9. Альтернатива")
    add_p(
        doc,
        "Если GP ещё не даёт конверсий — **не** переводить в пакет. "
        "Сначала вывести GP на 5+ конв./нед, потом подключать RB.",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    try:
        OUT_LOCAL.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUT_LOCAL)
    except OSError:
        pass
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
