#!/usr/bin/env python3
"""Сводка: групповые рыбалки в Сочи + катамаран «Моряк-Попай» (boat-sochi)."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Rybalka_Sochi_Moryak_Popaj_2026-08-13.docx"
SITE = "https://boat-sochi.ru/"
PHONE_NATALYA = "+7 918 304-40-00"
PHONE_OLEG = "+7 917 675-05-55"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else (13 if level == 2 else 12),
            bold=True,
            color=RGBColor(0x0B, 0x3D, 0x5C),
        )
    return p


def add_p(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if r_i % 2 == 1:
                shade_cell(cell, "F0F5F8")
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.8)

    add_heading(doc, "Рыбалка в Сочи: групповые выходы и катамаран «Моряк-Попай»", 1)
    add_p(doc, "Сводка на 13.08.2026. Источники: boat-sochi.ru, база знаний бота, рынок агрегаторов.")
    add_p(
        doc,
        "Важно: у boat-sochi нет продукта «групповая рыбалка по билетам с человека». "
        "На рынке Сочи такой формат распространён у конкурентов/агрегаторов. "
        "У нас — аренда катамарана целиком (обычная и трофейная рыбалка).",
        bold=True,
    )

    # --- Market ---
    add_heading(doc, "1. Групповые рыбалки в Сочи (рынок)", 1)
    add_p(
        doc,
        "Типичный сборный продукт у агрегаторов (Sputnik8, Tripster, ExtraGuide, ek-sochi и др.): "
        "билет за человека, выход на катере/катамаране/тримаране с капитаном и снастями.",
    )
    add_table(
        doc,
        ["Параметр", "Обычно на рынке"],
        [
            ["Формат", "Сборная группа, оплата за человека"],
            ["Длительность", "3 часа"],
            ["Цена", "~2 800–3 400 ₽ / чел."],
            ["Размер группы", "до 10–11 человек"],
            ["Расписание", "часто 09:00 / 12:00 / 15:00 (не у всех)"],
            ["Старт", "Адлер / Сириус / Кудепста / порт Сочи"],
            ["Включено", "снасти, капитан; часто приготовление улова на борту"],
            ["Типичный улов", "барабулька, ставрида, бычок и др.; гарантии нет"],
        ],
    )
    add_heading(doc, "Примеры предложений", 2)
    add_bullets(
        doc,
        [
            "Sputnik8 — «Групповая морская рыбалка в Адлере», ~3 060–3 400 ₽/чел., 3 ч, до 11 чел., Морской бульвар, 1",
            "ExtraGuide — то же направление, ~3 060 ₽/чел., ежедневно 9:00 / 12:00 / 15:00",
            "ek-sochi — групповая ~3 000 ₽/чел. (дети до 5 лет бесплатно у ряда офферов); индивидуальная от ~20 000 ₽ до 11 чел.",
            "Tripster / travel-обзоры — групповая рыбалка на катамаране ~3 ч, до ~10 чел., старт Кудепста / Сириус",
        ],
    )
    add_p(
        doc,
        "Это конкурентный/рыночный контур. Не путать с тарифами и форматом boat-sochi.",
        bold=True,
    )

    # --- boat-sochi fishing ---
    add_heading(doc, "2. Рыбалка у boat-sochi.ru", 1)
    add_p(
        doc,
        "На главной сайте — блок «Трофейная рыбалка» (попап). Отдельной посадочной /rybalka нет (404).",
    )
    add_bullets(
        doc,
        [
            "Цель трофейной: катран (черноморская акула) и скат — морской кот (хвостокол) / морская лисица (колючий скат)",
            "Катрана обычно ищут на глубине ориентир 40–70 м (у Имеретинки берег быстро уходит на глубину)",
            "На борту: снасти, жилеты; капитан помогает с ловлей",
            "Улов при удаче можно приготовить на борту катамарана",
            "В свободное время: купание, загар, виды",
            "Судно: катамаран «Моряк-Попай»",
            "В базе бота: есть обычная и трофейная рыбалка; детальный прайс описан для трофейной",
            "В Директе группа «Трофейная рыбалка» — остановлена (09.08.2026)",
        ],
    )
    add_table(
        doc,
        ["Параметр", "Значение (наш продукт)"],
        [
            ["Формат", "Аренда судна целиком (не билет с человека)"],
            ["Цена трофейной", "от 20 000 ₽ за выход"],
            ["Минимум", "3 часа"],
            ["Вместимость", "до 8 человек (полный состав не обязателен)"],
            ["Причал", "Сириус, Имеретинский порт, Парусная 1, линия 1"],
            ["Посадочная", f"{SITE} (блок на главной; отдельного URL нет)"],
            ["Контакты", f"капитан / Наталья {PHONE_NATALYA} (не капитан Олег — это «Сириус»)"],
        ],
    )
    add_p(
        doc,
        "Цены «от 7–8 тыс. ₽/час» с чужих сайтов проката катамарана — не наш тариф на трофейную рыбалку "
        "(зафиксировано в базе знаний бота).",
        bold=True,
    )

    # --- Papay ---
    add_heading(doc, "3. Катамаран «Моряк-Попай» (отдельно)", 1)
    add_p(
        doc,
        "Написание в наших источниках: «Моряк-Попай» / «Моряк Попай». "
        "Модель в открытых данных: Sunset Bay Maurell 230 (понтонный катамаран).",
    )

    add_heading(doc, "3.1. Наш продукт (boat-sochi + база бота)", 2)
    add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Название", "«Моряк-Попай»"],
            ["Модель", "Sunset Bay Maurell 230, понтонный катамаран"],
            ["Размер (ориентир)", "~7,5 м × ~2,5 м, малая осадка"],
            ["Причал", "Сириус / Имеретинский порт, Парусная 1, линия 1"],
            ["Вместимость у нас", "до 8 человек"],
            ["Время выхода (сайт)", "с 6:00 до 20:00"],
            ["Время выхода (бот)", "с 6:00 до 22:00 — уточнить актуальную норму"],
            ["Форматы", "обычная рыбалка + трофейная; также отдых / купание / закат"],
            ["Цена трофейной", "от 20 000 ₽, минимум 3 часа (за весь выход/судно)"],
            ["На борту", "снасти, жилеты, гальюн; капитан помогает с ловлей"],
            ["Улов", "гарантии конкретного улова нет; при удаче — приготовление на борту"],
            ["Плюсы катамарана", "устойчивее катера, меньше качки, больше палубы"],
            ["Контакты", f"капитан / Наталья {PHONE_NATALYA}; WhatsApp / Telegram / MAX — у Натальи"],
            ["Не путать", f"капитан Олег {PHONE_OLEG} — катер «Сириус», не Попай"],
        ],
    )
    add_heading(doc, "Что писать клиенту (коротко, из базы бота)", 3)
    add_bullets(
        doc,
        [
            "Рыбалка есть — обычная и трофейная",
            "Попай: катамаран в Сириусе, устойчивый, снасти на борту, от 20 000 ₽ / мин. 3 часа, катран и скат",
            "Детали и бронь — у капитана / Натальи",
        ],
    )

    add_heading(doc, "3.2. Чужие витрины проката (справочно, не наш прайс)", 2)
    add_p(
        doc,
        "На сторонних сайтах аренды (пример: ялт-адлер.рф / product/katamaran-moryak-popaj) "
        "тот же борт фигурирует с другими цифрами:",
    )
    add_table(
        doc,
        ["Параметр", "Чужая витрина"],
        [
            ["Модель", "Sunset Bay Maurell 230"],
            ["Класс", "Катамаран"],
            ["Длина / ширина", "7,5 м / 2,5 м"],
            ["Вместимость", "до 10 пассажиров"],
            ["Каюты", "0"],
            ["Гальюн", "1"],
            ["Старт", "Имеретинский порт Адлера"],
            ["Витринная цена", "от 7 200 ₽/час"],
            ["Заявленные плюсы", "устойчивость, простор, рыбалка, музыка, кейтеринг по запросу"],
        ],
    )
    add_p(
        doc,
        "Не выдавать чужие цены как наши. Для трофейной рыбалки ориентир — от 20 000 ₽ / мин. 3 часа за судно.",
        bold=True,
    )

    # --- Ops / ads ---
    add_heading(doc, "4. Реклама и операционка", 1)
    add_bullets(
        doc,
        [
            "Группа Директа «Трофейная рыбалка» — СТОП (память кампаний на 09.08.2026)",
            "В ЕПК групповых парусных прогулок рыбалка вынесена в минус-слова (не смешивать продукты)",
            "Отдельной посадочной под рыбалку/Попай на boat-sochi.ru нет — только блок на главной",
            "Бронь: предоплата; детали — Наталья / капитан судна",
        ],
    )

    add_heading(doc, "5. Источники", 1)
    add_bullets(
        doc,
        [
            f"Сайт: {SITE} — блок «Трофейная рыбалка»",
            "База бота (ветки boat-sochi-max / telegram): knowledge/faq-extra.md, knowledge/llms-full.txt",
            "Память кампаний: cursor/Morskie_progulki_campaign_memory.md",
            "Рынок: Sputnik8, ExtraGuide, ek-sochi, Tripster/обзоры",
            "Чужая витрина Попай: каталоги аренды Имеретинского порта (от 7 200 ₽/час — не наш трофейный тариф)",
        ],
    )

    add_heading(doc, "6. Краткий вывод", 1)
    add_bullets(
        doc,
        [
            "«Групповая рыбалка» в Сочи на рынке ≈ 3 тыс. ₽/чел., 3 часа, сборные группы",
            "У boat-sochi такого билетного продукта нет",
            "Наш продукт — private charter «Моряк-Попай» под обычную/трофейную рыбалку",
            "Трофей: катран и скат, от 20 000 ₽, минимум 3 часа, до 8 гостей, линия 1 Сириус",
        ],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
