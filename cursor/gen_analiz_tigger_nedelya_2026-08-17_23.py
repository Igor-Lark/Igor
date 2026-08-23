#!/usr/bin/env python3
"""Анализ РК «Яхта Tigger» календарная неделя 17–23.08.2026 → Word."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Analiz_RK_Tigger_nedelya_2026-08-17_23.docx"
NAVY = (0x1F, 0x4E, 0x79)
GRAY = (0x59, 0x59, 0x59)


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_p(doc, text, size=11, bold=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.12
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def fill_cell(cell, text, bold=False, size=10, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold or header, color=(255, 255, 255) if header else None)
    set_cell_border(cell)
    if header:
        shade_cell(cell, "1F4E79")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, header=True, size=10)
    for r_i, row in enumerate(rows):
        prevent_row_split(table.rows[r_i + 1])
        for c_i, val in enumerate(row):
            fill_cell(table.rows[r_i + 1].cells[c_i], val, size=10, bold=c_i == 0)
            if r_i % 2 == 1:
                shade_cell(table.rows[r_i + 1].cells[c_i], "F2F2F2")
    if col_widths:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(int(sum(col_widths) * 567)))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*NAVY)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run(run, size=size)


def add_page_field(paragraph):
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    set_run(run, size=9, color=GRAY)


def setup_print(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Яхта Tigger  ·  Яндекс Директ  ·  календарная неделя 17–23.08.2026")
    set_run(run, size=9, color=GRAY)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Внутренний анализ РК  ·  страница ")
    set_run(run, size=9, color=GRAY)
    add_page_field(fp)
    run = fp.add_run("  ·  Analiz_RK_Tigger_nedelya_2026-08-17_23.docx")
    set_run(run, size=9, color=GRAY)


MINUS_NEW = [
    "круиз",
    "катание",
    "отель",
    "виноградная",
    "батуми",
    "грузии",
    "виктория",
    "рахиль",
    "globalsailor",
]

PLUS_KEYS = [
    "яхта тайгер",
    "яхта tigger",
    "аренда яхты tigger сочи",
    "azimut 68 сочи",
    "аренда vip яхты сочи",
]


def build():
    doc = Document()
    setup_print(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_p(
        doc,
        "Анализ рекламной кампании «Яхта Tigger»",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_p(
        doc,
        "Календарная неделя Директа: понедельник 17.08 — воскресенье 23.08.2026",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_p(
        doc,
        "Файл 23.08 = срез 20–23.08. Полная неделя собрана: 17–19 из выгрузки 21.08 + 20–23 из этой.",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    heading(doc, "1. Короткий вердикт", 1)
    bullet(
        doc,
        "Неделя 17–23: 87 показов / 14 кликов / CTR 16% / 1 конверсия / 137,50 ₽. "
        "Обучение 1/10 — снова не дотянуто, обучение остановится.",
    )
    bullet(
        doc,
        "Минусы и выкл. широких/альтернативных/сопутствующих сжали мусор и объём: "
        "232 показа на прошлой неделе → 87. После правок пт 21.08 сб–вс всего 10 показов / 2 клика.",
    )
    bullet(
        doc,
        "Качество лучше: CTR 9,9% → 16%. Появилось ядро «яхта тайгер» (автотаргет, 8 показов, 0 кликов). "
        "Позиция хуже: ~3,3 → ~3,9, объём трафика 80% → 59%.",
    )
    bullet(
        doc,
        "Единственная конверсия — всё ещё 17.08 «туры на частной яхте адлер», М 18–24. "
        "Сб–вс новых конверсий нет.",
    )
    bullet(
        doc,
        "Что улучшить сейчас: плюс-фразы на бренд/VIP + минусы новых дыр столбиком. "
        "CPA не поднимать на этой же неделе после пятничных правок — ждать пн–вс 24–30.08.",
    )

    heading(doc, "2. Неделя к неделе", 1)
    add_table(
        doc,
        ["Метрика", "Пн–вс 10–16", "Пн–вс 17–23", "Срез файла 20–23"],
        [
            ["Показы", "232", "87", "33"],
            ["Клики", "23", "14", "5"],
            ["CTR", "9,91%", "16,1%", "15,2%"],
            ["Конверсии", "0", "1", "0"],
            ["Расход", "0 ₽", "137,50 ₽", "0 ₽"],
            ["CPA", "—", "137,50 ₽", "—"],
            ["Ср. позиция", "~3,1–3,3", "~3,9", "3,88"],
            ["Объём трафика", "~80%", "~59%", "59%"],
            ["Отказы / глубина", "27% / 1,14", "— / ~1,2", "20% / 1,20"],
        ],
        col_widths=[4.0, 4.2, 4.6, 4.8],
    )

    heading(doc, "2.1. По дням 17–23.08", 2)
    add_table(
        doc,
        ["День", "Показы", "Клики", "Конв.", "Расход", "Заметка"],
        [
            ["Пн 17.08", "28", "3", "1", "137,50", "Единственная конверсия недели"],
            ["Вт 18.08", "18", "5", "0", "0", "Пик кликов"],
            ["Ср 19.08", "8", "1", "0", "0", "Спад"],
            ["Чт 20.08", "3", "0", "0", "0", "Почти ноль"],
            ["Пт 21.08", "20", "3", "0", "0", "Отскок; вечером широкие выкл. + минусы"],
            ["Сб 22.08", "6", "1", "0", "0", "После правок: двухпалубная яхта Адлер"],
            ["Вс 23.08", "4", "1", "0", "0", "морпорт (опечатка), отказ 100%"],
        ],
        col_widths=[2.6, 2.0, 1.8, 1.8, 2.2, 7.2],
    )
    add_p(
        doc,
        "«Широкие» в столбце категории CSV — это класс запроса Директа, не галка автотаргета. "
        "Галку широких вы выключили; фразовые ключи по-прежнему могут помечаться как широкие. Это норма.",
    )

    heading(doc, "3. Что уже хорошо", 1)
    bullet(doc, "Автотаргет 104→20 показов за неделю; сб–вс автотаргет без кликов, в основном «тайгер».")
    bullet(doc, "Клики пт–вс ближе к ядру: капитан, морской порт, двухпалубная.")
    bullet(doc, "Кампания работает. РСЯ выкл. Каюты не трогаем.")
    bullet(doc, "Не минусовали «туры» — правильно: единственная конверсия на «туры на частной яхте».")

    heading(doc, "4. Что улучшить", 1)
    heading(doc, "4.1. Сегодня — плюс-фразы (объём без РСЯ и без широкого автотаргета)", 2)
    add_p(
        doc,
        "После чистки стратегия почти не видит ядро. «яхта тайгер» уже 8 показов и 0 кликов — "
        "добавить как ключевые (фраза / точная), чтобы не жить только на автотаргете.",
        bold=True,
        space_after=4,
    )
    add_p(doc, "Добавить в группу столбиком:", bold=True, space_after=3)
    for k in PLUS_KEYS:
        bullet(doc, k)

    heading(doc, "4.2. Сегодня — доминусовать новые дыры (столбик)", 2)
    add_p(
        doc,
        "Всплыло после пятницы: отель Азимут на Виноградной, батуми/грузия, чужие яхты, «круиз» "
        "(минусовали «круизы», без единственного числа).",
        space_after=4,
    )
    for phrase in MINUS_NEW:
        bullet(doc, phrase)
    add_p(
        doc,
        "Не минусовать: тайгер, tigger, азимут, гидроцикл, туры, vip, премиум, морпорт, двухпалубная.",
        space_before=4,
    )

    heading(doc, "4.3. CPA — не сейчас", 2)
    add_p(
        doc,
        "План «если ядро чистое и VIP-показов мало → телефон/MAX 350–400» формально уже про объём. "
        "Но пт 21.08 вы выключили широкие и внесли минусы. Сб–вс — 10 показов, мало для оценки. "
        "Поднять CPA в вс 23.08 = снова сдвинуть обучение. Держать 300 / 250 / 300 / 210 "
        "календарную неделю 24–30.08. Если на вс 30.08 ядро «тайгер / vip / azimut / порт» "
        "и позиция хуже 4 при живых объявлениях — тогда одно изменение: телефон и MAX 350, Telegram 300, профиль 250.",
    )

    heading(doc, "4.4. Не делать", 2)
    bullet(doc, "Не включать широкие / альтернативные / сопутствующие обратно ради 10 конв./нед.")
    bullet(doc, "Не включать РСЯ и галерею услуг.")
    bullet(doc, "Не править каюты. Не менять тип стратегии.")
    bullet(doc, "Не минусовать «азимут» целиком — режем отель/бассейн/виноградную.")

    heading(doc, "5. Клики 20–23.08 (этот файл)", 1)
    add_table(
        doc,
        ["День", "Запрос", "Кто", "Вердикт"],
        [
            ["21.08", "аренда яхты в сочи с капитаном (2 кл.)", "М 45–54, Остальные, отказ 0%", "Ядро, оставить"],
            ["21.08", "аренда яхты морской порт сочи", "М 25–34, 2–5%, глуб. 2", "Ядро, оставить"],
            ["22.08", "двухпалубная яхта в аренду в адлере", "Ж 25–34, 2–5%, отказ 0%", "Близко к Tigger"],
            ["23.08", "аренда яхты морпорт сояи", "Ж 25–34, 6–10%, отказ 100%", "Опечатка порта, интент ок"],
        ],
        col_widths=[2.2, 6.4, 5.4, 3.6],
    )

    heading(doc, "6. План на неделю 24–30.08", 1)
    add_table(
        doc,
        ["Когда", "Действие"],
        [
            ["Вс 23.08 / пн 24.08", "Плюс-фразы §4.1 + минусы §4.2. Стратегию и CPA не трогать"],
            ["Пн–вс 24–30.08", "Ежедневно: новые чужие яхты/отель Азимут — в минус"],
            ["Вс 30.08", "Выгрузка пн–вс. Смотрим: показы ядра «тайгер», конверсии, позицию"],
            ["После 30.08", "Если ядра мало и позиция ≥4 — одно повышение CPA, не раньше"],
        ],
        col_widths=[4.4, 13.2],
    )

    heading(doc, "7. Источники", 1)
    bullet(doc, "CSV: 2026-08-23_18-09-05_vitaminki21.csv — факт 20–23.08, 33/5/0.")
    bullet(doc, "Добор 17–19.08: 2026-08-21_21-41-10_vitaminki21.csv.")
    bullet(doc, "Память: Tigger_campaign_memory.md. Каюты / CPA / РСЯ — без изменений.")

    add_p(
        doc,
        "Документ внутренний, A4. Не для клиента. В очередь печати не ставился.",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
    )
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
