#!/usr/bin/env python3
"""G11 луфарь без G12 — пеламида остаётся в G5."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "G11_Lufar_Popaj_2026-08-30.docx"
OUT_LOCAL = Path(__file__).resolve().parent.parent / "локальная" / OUT.name

LANDING = "https://boat-sochi.ru/gruppovaja_ribbalka"
PHONE = "+7 918 304-40-00"
PRICE = "2 800 ₽/чел."
CAMPAIGN_NO = "713632237"
NAVY = RGBColor(0x0B, 0x3D, 0x5C)
GRAY = RGBColor(0x33, 0x33, 0x33)
RED = RGBColor(0x8B, 0x1A, 0x1A)

MINUS_G11 = """-катран
-скат
-трофейная
-трофей
-целиком
-индивидуальная
-аренда катамарана
-20000
-20 000
-барабулька
-султанка
-ставрида
-пеламида
-пеламиду
-пеламиды
-!с берега
-!с пирса
-!с буны
-форель
-форелев
-форелевое
-форелевая
-пруд
-прудовая
-абхазия
-крым
-севастополь
-анапа
-геленджик
-дельфины
-дельфинарий
-парусная
-гидроцикл
-теплоход
-круиз
-купить яхту
-магазин
-базар
-брюки
-женская
-кафе
-кура
-курка
-мужская
-птица
-ресторан
-рынок
-столовая
-штаны
-вакансия
-бесплатно
-падел
-прокат катер
-лодку на прокат"""

MINUS_ADD_G5 = """-луфарь
-луфаря
-луфарю"""

KEYS_G11 = """рыбалка на луфаря сочи
рыбалка на луфаря сириус
луфарь сочи
луфарь сириус -рыбалка
ловля луфаря в сочи
морская рыбалка луфарь
луфарь черное море сочи
рыбалка луфарь адлер
осенняя рыбалка луфарь сочи"""


def set_run(run, *, bold=False, size=11, color=GRAY, font="Arial"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font


def add_p(doc, text, *, bold=False, size=11, color=GRAY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    p.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=True, size=16 if level == 1 else 13, color=NAVY)
    p.paragraph_format.space_before = Pt(12)


def add_bullets(doc, items, *, color=GRAY):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        set_run(run, size=11, color=color)


def add_code(doc, text):
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run(run, size=10, font="Consolas")
        p.paragraph_format.space_after = Pt(0)


def build():
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("G11 · Луфарь (без G12)")
    set_run(run, bold=True, size=18, color=NAVY)
    add_p(doc, f"РК ЕПК № {CAMPAIGN_NO} · {PRICE} · «Моряк Попай» · {LANDING}", size=10)

    add_heading(doc, "Решение 30.08.2026")
    add_bullets(
        doc,
        [
            "**G12 не заводим.** Пеламида остаётся в **G5** (лето-бархат) вместе с барабулькой и ставридой.",
            "**G11** — только **луфарь** (осень): отдельная группа, 1 объявление, свои фото.",
            "Старый Word «G11+G12 схема A» — **не актуален**.",
        ],
        color=RED,
    )

    add_heading(doc, "1. Разделение рыб по группам")
    add_bullets(
        doc,
        [
            "G5 — барабулька / ставрида / **пеламида** / султанка (не трогать ключи пеламиды).",
            "G11 — только луфарь (новые ключи, только луфарь в текстах и картинках).",
            "G10 — широкий «морская рыбалка Сочи» без вида рыбы; минусы −луфарь −пеламида.",
        ],
    )

    add_heading(doc, "2. G5 — что добавить (не убирать пеламиду)")
    add_p(doc, "В минусы G5 добавить, чтобы G11 не каннибализировала:", bold=True, size=10)
    add_code(doc, MINUS_ADD_G5)
    add_p(doc, "Пеламиду из G5 **не удалять** — ключи, заголовки и фото пеламиды остаются.", size=10)

    add_heading(doc, "3. G10 — минусы (без изменений)")
    add_code(doc, MINUS_ADD_G5 + "\n-пеламида\n-пеламиду\n-пеламиды")

    add_heading(doc, "4. Группа G11 — настройки")
    add_bullets(
        doc,
        [
            "Имя: G11 - луфарь",
            "Факт: объявление № **1919656966051937203** (если уже заведено)",
            "Профиль пользователя / интересы: **ПУСТО** (как G5)",
            "Гео: Россия · автотаргет как у G5",
            "UTM: utm_content=G11_lufar",
            "Посадка: /gruppovaja_ribbalka",
        ],
    )

    add_heading(doc, "5. Ключи G11")
    add_code(doc, KEYS_G11)

    add_heading(doc, "6. Минусы G11")
    add_code(doc, MINUS_G11)

    add_heading(doc, "7. Объявление G11 — только луфарь")
    add_p(doc, "Заголовки:", bold=True, size=10)
    add_bullets(
        doc,
        [
            "Рыбалка на луфаря в Сочи — 2 800 ₽/чел.",
            "Луфарь · катамаран «Моряк Попай»",
            "Ловля луфаря · Чёрное море · Сочи",
            "3 часа в море · снасти включены",
            "Группа до 8 · Сириус / линия 2",
            f"Бронь: {PHONE}",
        ],
    )
    add_p(doc, "Не использовать заголовки «осенний сезон / 3 часа в море» без слова луфарь.", size=10, color=RED)
    add_p(doc, "Креативы: только луфарь. Не грузить баннер G5 с барабулькой/пеламидой.", size=10)

    add_heading(doc, "8. Чеклист")
    add_bullets(
        doc,
        [
            "☐ G12 — **не создавать**",
            "☐ G5 — пеламида на месте; минусы −луфарь добавлены",
            "☐ G11 — только луфарь в ключах/текстах/фото",
            "☐ G10 — минусы −луфарь −пеламида",
            "☐ Стратегию / бюджет / CPA не трогать до решения по пакету",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    try:
        OUT_LOCAL.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUT_LOCAL)
    except OSError:
        pass
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
