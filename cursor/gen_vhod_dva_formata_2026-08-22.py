#!/usr/bin/env python3
"""Word: вход «морские прогулки» — два формата в одном объявлении."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "D9E2F3")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table


def main() -> None:
    doc = Document()
    title = doc.add_heading("Два формата на один запрос", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Морские прогулки в Сириусе · группа от 1 800 ₽ и катер целиком от 15 000 ₽"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub2 = doc.add_paragraph("22 августа 2026 · аккаунт vitaminki21 · не менять CPA и бюджет")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].italic = True

    add_heading(doc, "1. Зачем так", 1)
    doc.add_paragraph(
        "Клиент пишет «морские прогулки сириус» и не знает, что есть место в группе "
        "и аренда катера целиком. Две кампании на эту фразу не покажут два объявления: "
        "Директ почти всегда выберет одно, а вы торгуетесь сами с собой."
    )
    doc.add_paragraph(
        "Нужна одна дверь: одно объявление, обе цены, главная сайта (там уже оба блока), "
        "быстрые ссылки на группу и на катер."
    )

    add_heading(doc, "2. Куда кликать в Директе", 1)
    doc.add_paragraph(
        "Кампания «Сириус — прокат катеров» № 704503370. Новая группа "
        "«Морские прогулки — два формата». Не создавать заново «Морские прогулки» 712465896. "
        "Стратегию, 12 000 ₽ и цены целей не трогать."
    )
    bullets = [
        "Группа объявлений → создать. Гео как у «Аренда в Сириусе».",
        "Автотаргет: широкие / сопутствующие / альтернативные ВЫКЛ.",
        "Посадка объявлений: https://boat-sochi.ru/?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=vybor",
        "Телефон в шапке объявления не вешать один — иначе все звонки уйдут Олегу или Наталье.",
        "В РК «Групповые прогулки» 713449981: голые фразы из раздела 3 — пауза или минус.",
    ]
    for text in bullets:
        doc.add_paragraph(text, style="List Number")

    add_heading(doc, "3. Ключи только этой группы", 1)
    doc.add_paragraph(
        '"морские прогулки сириус"\n'
        '"морская прогулка сириус"\n'
        '"прогулки на яхте сириус"\n'
        '"прогулка на яхте сириус"\n'
        '"морские прогулки имеретинский порт"\n'
        '"морская прогулка имеретинский порт"'
    )

    add_heading(doc, "4. Заголовки и тексты", 1)
    add_table(
        doc,
        ["Поле", "Текст"],
        [
            ["Заголовок 1", "Морские прогулки в Сириусе"],
            ["Заголовок 2", "Группа от 1 800 ₽"],
            ["Заголовок 3", "Катер целиком от 15 000 ₽"],
            ["Заголовок 4", "Два формата — выберите"],
            ["Заголовок 5", "Имеретинский порт"],
            ["Заголовок 6", "Парусная или катер ваш"],
            [
                "Текст 1",
                "Групповая 1,5 ч от 1 800 ₽/чел. Или катер целиком с капитаном — от 15 000 ₽ за 2 часа.",
            ],
            [
                "Текст 2",
                "Не знаете формат? На сайте два варианта: место в группе и судно только для вас.",
            ],
        ],
    )

    add_heading(doc, "5. Восемь быстрых ссылок", 1)
    doc.add_paragraph(
        "Описание — две строки через перевод строки, не вкладывать ||. "
        "В этой группе лендинг групповых можно: запрос общий."
    )
    add_table(
        doc,
        ["№", "Текст", "Строка 1", "Строка 2", "Куда"],
        [
            [
                "1",
                "Группа от 1 800 ₽",
                "Место в группе 1,5 ч",
                "До 11 человек",
                "https://boat-sochi.ru/progulki_na_yacht?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_group",
            ],
            [
                "2",
                "Катер целиком",
                "Ваш катер с капитаном",
                "От 15 000 ₽ за 2 ч",
                "https://boat-sochi.ru/?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_private",
            ],
            [
                "3",
                "Утро к дельфинам",
                "Скидка до 12:00",
                "Катер Сириус",
                "https://boat-sochi.ru/delfin?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_delfin",
            ],
            [
                "4",
                "Закат в группе",
                "Слот 18:00",
                "Билет 2 500 ₽",
                "https://boat-sochi.ru/progulki_na_yacht?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_group_sunset",
            ],
            [
                "5",
                "Наш флот",
                "Катера и яхты целиком",
                "В порту Сириус",
                "https://boat-sochi.ru/?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_fleet",
            ],
            [
                "6",
                "Причал Сириус",
                "Парусная 1",
                "Имеретинский порт",
                "https://boat-sochi.ru/?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_pier",
            ],
            [
                "7",
                "Забронировать",
                "Группа или катер",
                "Выберите на сайте",
                "https://boat-sochi.ru/?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_book",
            ],
            [
                "8",
                "Контакты",
                "Группа +7 918 304-40-00",
                "Катер +7 917 675-05-55",
                "https://boat-sochi.ru/?utm_source=yandex&utm_medium=cpc&utm_campaign=sirius_704503370&utm_content=sl_contacts",
            ],
        ],
    )

    add_heading(doc, "6. Не делать", 1)
    for text in (
        "Не включать обе РК на одну голую фразу.",
        "Не тащить сюда Tigger, Алексум и рыбалку.",
        "Не менять бюджет и CPA Сириуса и групповых в этот же день.",
        "Не обещать дельфинов. Не писать 7 000 ₽.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    out = "/workspace/cursor/Vhod_morskie_progulki_dva_formata_2026-08-22.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
