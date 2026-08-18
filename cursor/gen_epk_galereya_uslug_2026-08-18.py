#!/usr/bin/env python3
"""ЕПК Галерея услуг boat-sochi.ru — Word: как создать РК в Директе."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "EPK_Galereya_Uslug_BoatSochi_2026-08-18.docx"

PHONE = "+7 917 675-05-55"
SITE = "https://boat-sochi.ru/"
COUNTER_SITE = "94713538"
COUNTER_ORG = "103116887"
UTM = "utm_source=yandex&utm_medium=cpc&utm_campaign=epk_galereya_uslug_boatsochi"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13,
            bold=True,
            color=RGBColor(0x0B, 0x3D, 0x5C),
        )
    return p


def add_p(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if r_i % 2 == 1:
                shade_cell(cell, "F2F7FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


MINUS = [
    "вакансия",
    "работа",
    "резюме",
    "купить яхту",
    "купить катер",
    "продать",
    "кредит",
    "бесплатно",
    "дельфинарий",
    "плавать с дельфинами",
    "океанариум",
    "теплоход",
    "теплохода",
    "круизный лайнер",
    "крым",
    "абхаз",
    "абхазия",
    "турци",
    "геленджик",
    "анапа",
    "туапсе",
    "новороссийск",
    "без капитана",
    "самостоятельн",
    "гидроцикл купить",
    "форель",
    "платник",
    "магазин снастей",
    "с пирса",
    "скайпарк",
    "подкова",
    "дверь в море",
    "сириус рент",
    "наталья",
]

KEYS = [
    "аренда катера сириус",
    "аренда катера сочи",
    "аренда яхты сириус",
    "морские прогулки сириус",
    "морская прогулка сочи",
    "прогулка на катере сириус",
    "прогулка на яхте сочи",
    "аренда катера адлер",
    "прокат катера имеретинский порт",
    "яхта сириус прогулка",
]


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    add_heading(doc, "ЕПК «Галерея услуг» — boat-sochi.ru", 0)
    add_p(
        doc,
        "Аккаунт vitaminki21 · код GU · 18.08.2026. РК ещё нет в Директе. "
        "Этот документ — как завести кампанию. Агент в кабинет не входит.",
        bold=True,
    )

    add_heading(doc, "1. Зачем отдельная РК", 1)
    add_p(
        doc,
        "Галерея услуг во всех текущих кампаниях выключена: иначе билет 1 800, "
        "катер 7 500 и VIP 50 000 смешиваются в одном блоке. Эта ЕПК — единственный "
        "владелец места «Список организаций и галерея услуг». На SR, MP, ZK, GP, TG, RB "
        "галерею не включать.",
    )
    add_p(
        doc,
        "Клик открывает карточку организации Яндекс Бизнеса (звонок, мессенджер, сайт), "
        "не чужой оффер в поиске. Поисковую выдачу, РСЯ и Карты в этой РК на старте выключить.",
    )

    add_heading(doc, "2. Организация и счётчики", 1)
    add_table(
        doc,
        ["Поле", "Значение"],
        [
            ["Название РК", "EPK_Galereya_Uslug_BoatSochi"],
            ["Сайт", SITE],
            ["Телефон карточки", PHONE + " (Олег)"],
            ["Причал", "Сириус, Парусная 1, Имеретинский порт"],
            ["Адрес в профиле (сверить)", "Морской бульвар 1В/3, Сириус"],
            ["Счётчик профиля (главный)", COUNTER_ORG],
            ["Счётчик сайта", COUNTER_SITE],
            ["Бюджет старт", "8 000–10 000 ₽ / неделя"],
            ["Стратегия", "Максимум конверсий, оплата за конверсии"],
        ],
        col_widths=[6, 11],
    )

    add_heading(doc, "3. Шаги в Директе", 1)
    add_bullets(
        doc,
        [
            "Добавить → Кампанию → Режим эксперта → Единая перфоманс-кампания.",
            "Имя: EPK_Galereya_Uslug_BoatSochi. Ссылка: " + SITE,
            "Организация из Яндекс Бизнеса: карточка boat-sochi.ru. Если не нашлась — добавить вручную. Одну, без дубля.",
            "Места показа: Список организаций и галерея услуг — ВКЛ. Поиск, товарная галерея, динамические места, РСЯ, Карты — ВЫКЛ.",
            "Максимум конверсий · за конверсии · цена конверсии. Директ помогает — ВЫКЛ.",
            "Счётчики 103116887 и 94713538.",
            "Цели профиля: Позвонить 100–120 ₽, мессенджер 90–110 ₽. Сайт: телефон 140, TG/MAX 120. Профиль не дешевле сайта вдвое.",
            "Автотаргет: альтернативные и сопутствующие ВЫКЛ. Возраст младше 18: −100%. Гео: Россия.",
            "Восемь быстрых ссылок с описанием из двух строк. UTM: " + UTM,
            "Пока нет 10 конверсий за неделю Пн–Вс — не менять тип, оплату, бюджет, цены целей.",
        ],
    )

    add_heading(doc, "4. Карточка Яндекс Бизнеса до запуска", 1)
    add_bullets(
        doc,
        [
            "Сайт https://boat-sochi.ru/ совпадает с кампанией.",
            "Телефон +7 917 675-05-55, мессенджеры те же.",
            "Рубрика услуговая: прокат катеров / яхт / морские прогулки — иначе галерея не покажет.",
            "Фото реальных судов, без водяных знаков Avito.",
            "Цены как на сайте. Дельфинов не гарантировать.",
            "Одна карточка. Счётчик профиля 103116887 привязан к ней.",
        ],
    )

    add_heading(doc, "5. Фразы (ядро)", 1)
    add_p(doc, "Фразовое соответствие. Одна строка — одна фраза.")
    for k in KEYS:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(k)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(1)

    add_heading(doc, "6. Минус-слова кампании", 1)
    add_p(doc, "Копировать столбиком в поле минус-слов. Не в одну строку.")
    for m in MINUS:
        p = doc.add_paragraph()
        run = p.add_run(m)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    add_heading(doc, "7. Восемь быстрых ссылок", 1)
    add_table(
        doc,
        ["№", "Текст", "Описание", "Куда"],
        [
            ["1", "Флот и цены", "Катера и яхты\nИмеретинский порт", f"{SITE}?{UTM}&utm_content=sl_fleet"],
            ["2", "Утро к дельфинам", "Скидка до 12:00\nКатер Сириус", f"https://boat-sochi.ru/delfin?{UTM}&utm_content=sl_delfin"],
            ["3", "Групповые прогулки", "Билет от 1 800 ₽\n1,5 часа, до 11", f"https://boat-sochi.ru/progulki_na_yacht?{UTM}&utm_content=sl_group"],
            ["4", "Причал Сириус", "Парусная 1\nИмеретинский порт", f"{SITE}?{UTM}&utm_content=sl_pier"],
            ["5", "Купание в море", "Открытое море\nВдали от берега", f"{SITE}?{UTM}&utm_content=sl_swim"],
            ["6", "Закат на море", "Выход с 18:00\nБронь заранее", f"{SITE}?{UTM}&utm_content=sl_sunset#services"],
            ["7", "Забронировать", "Звонок или MAX\n+7 917 675-05-55", f"{SITE}?{UTM}&utm_content=sl_book"],
            ["8", "Контакты", "Олег, капитан\nTelegram · MAX", f"{SITE}?{UTM}&utm_content=sl_contacts"],
        ],
        col_widths=[1.2, 3.4, 4.2, 8],
    )
    add_p(doc, "В Директе описание — две строки через перевод строки, не через ||.")

    add_heading(doc, "8. Антиканнибализация", 1)
    add_table(
        doc,
        ["РК", "Что сделать"],
        [
            ["SR 704503370, MP 712465896, ZK, GP, TG, RB", "Галерея услуг / список организаций — ВЫКЛ"],
            ["MP", "Поиск «аренда / морские прогулки» остаётся у MP. GU поиск не включает"],
            ["SR", "Посадка /delfin остаётся у SR"],
            ["GP", "Не делать «1 800 ₽» единственным заголовком GU"],
            ["TG / RB", "Tigger и рыбалку в sitelinks старта не ставить"],
        ],
        col_widths=[6, 11],
    )

    add_heading(doc, "9. После создания", 1)
    add_bullets(
        doc,
        [
            "Прислать агенту GU номер кампании и скрин: стратегия, цели, места показа.",
            "Проверить, что в соседних РК галерея выключена.",
            "Смотреть звонки карточки (103116887), не CTR поиска.",
            "Не включать Поиск и РСЯ «чтобы набрать 10 конверсий».",
        ],
    )

    add_p(doc, "Handoff агента: cursor/Galereya_uslug_agent_handoff.md", size=10)
    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build()
