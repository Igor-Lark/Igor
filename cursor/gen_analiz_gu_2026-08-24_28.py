#!/usr/bin/env python3
"""Анализ GU 24–28.08.2026 (Пн–Пт, без сб–вс) → Word."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Analiz_RK_Galereya_Uslug_2026-08-24_28.docx"


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size=16 if level == 1 else 13,
            bold=True,
            color=RGBColor(0x0B, 0x3D, 0x5C),
        )
    return p


def add_p(doc, text, *, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_column(doc, items):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. ")
        set_run_font(r, size=11, bold=True)
        r2 = p.add_run(item)
        set_run_font(r2, size=11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if r_i % 2 == 1:
                shade_cell(cell, "F2F7FA")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


MINUS_NOW = [
    "пицунда",
    "яхт клуб",
    "яхт-клуб",
    "детский клуб",
    "туры из сириуса",
]

KEYS = [
    "аренда катера сириус",
    "аренда катера сочи",
    "аренда яхты сириус",
    "морские прогулки сириус",
    "морская прогулка сочи",
    "прогулка на катере сириус",
    "прогулка на яхте сочи",
    "аренда катера адлер",
    "прокат катера имеретинский порт",
    "яхта сириус прогулка",
]


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    add_heading(doc, "Анализ GU 24–28.08 — мало показов", 0)
    add_p(
        doc,
        "vitaminki21 · GU · группа Яхты Сочи 5789218554 · "
        "Пн–Пт 24–28.08.2026. Сб и вс пользователь добавит в воскресенье. "
        "Файл: 2026-08-29_07-19-29_vitaminki21.csv. Печать не нужна.",
        bold=True,
    )

    add_heading(doc, "1. Вердикт: мало — это не поломка", 1)
    add_bullets(
        doc,
        [
            "5 будних дней: 9 показов · 0 кликов · 0 конв. · 0 ₽ · CTR 0% · позиция 1,22. Все строки — список организаций, 100% автотаргет, ключевые фразы пустые.",
            "Сравни с 18–23.08 (6 дней, до минусов): 22 показа · 2 клика. После минусов 23.08 темп упал с ~3,7 до ~1,8 показа/день. Срезали мусор (торт, сапы, клубы, Нугуш) — так и задумано.",
            "Место показа узкое само по себе: только список организаций / галерея, без поисковых объявлений. 9 показов за Пн–Пт для такого блока — мало, но ожидаемо. Не открывать Поиск/РСЯ «чтобы набрать показы».",
            "Оплата за конверсии + 0 конв. = стратегия не покупает клики. Бюджет 5 000 не упирается: расход 0.",
            "Неделя неполная: нет сб–вс. Догрузка в воскресенье.",
            "Пользователь 29.08: обучение можно сбросить (конверсий нет). Делаем пакет: профиль главный + WA/TG карточки + минусы + 10 фраз.",
        ],
    )

    add_heading(doc, "2. Почему мало показов (разложить)", 1)
    add_table(
        doc,
        ["Причина", "Вес", "Что делать"],
        [
            ["Только список организаций, поиск выкл", "главный", "Так задумано для GU. Не включать выдачу"],
            ["Минусы 23.08 убрали ~половину старых показов", "сильный", "Ок. Это были торт/сапы/Нугуш"],
            ["0 конв. · оплата за конверсии · учится на сайте", "сильный", "Сброс ок. Главный счётчик → профиль 103116887"],
            ["Нет ключевых фраз, только автотаргет", "средний", "Добавить 10 фраз ядра фразовым"],
            ["Нет сб–вс в этой выгрузке", "средний", "Догрузить в воскресенье"],
            ["Карусель «галерея услуг» в отчёте снова 0", "средний", "Рубрика карточки / аукцион. Не лечить поиском"],
            ["Бюджет 5 000 мало", "нет", "Расход 0 ₽. Бюджет не причина"],
        ],
        col_widths=[7, 2.2, 7.8],
    )

    add_heading(doc, "3. Сравнение двух срезов", 1)
    add_table(
        doc,
        ["Срез", "Дни", "Показы", "Кл.", "Конв.", "₽", "Пок./день"],
        [
            ["18–23.08 (до/в день минусов)", "6 (вт–вс)", "22", "2", "0", "0", "~3,7"],
            ["24–28.08 (после минусов, без сб–вс)", "5 (пн–пт)", "9", "0", "0", "0", "~1,8"],
        ],
        col_widths=[6, 3, 2, 1.5, 1.8, 1.5, 2.2],
    )
    add_p(
        doc,
        "Чистые запросы (катер/яхта + Сириус/Адлер) в новом срезе есть: 5 из 9 показов. "
        "Мусор остался: Пицунда 2, детский яхт-клуб 1, «туры из сириуса» 1.",
    )

    add_heading(doc, "4. По дням и запросам", 1)
    add_table(
        doc,
        ["День", "Пок.", "Запрос", "Кат.", "Вердикт"],
        [
            ["24.08", "1", "аренда яхты цена адлер", "цел", "своё"],
            ["25.08", "1", "детский яхт клуб в сириусе", "цел", "минус — кружок, не аренда"],
            ["25.08", "2", "пицунда аренда яхты", "цел", "минус — Абхазия (абхаз не поймал)"],
            ["26.08", "1", "аренда катера адлер", "цел", "своё"],
            ["27.08", "1", "аренда катера", "шир", "своё, широко"],
            ["27.08", "1", "аренда катеров сириус", "цел · бренд конк.", "своё (катер/курорт)"],
            ["27.08", "1", "аренда яхт адлер", "цел", "своё"],
            ["28.08", "1", "туры из сириуса", "цел", "минус — экскурсии, не флот"],
        ],
        col_widths=[2.2, 1.4, 5.4, 2.8, 5.2],
    )

    add_heading(doc, "5. Минусы — добавить столбиком", 1)
    add_p(doc, "К тем, что уже стоят с 23.08. Не минусовать адлер / аренда / сириус / катер.")
    add_column(doc, MINUS_NOW)

    add_heading(doc, "6. Сброс обучения — что кликнуть", 1)
    add_p(
        doc,
        "Конверсий 0 за 18–28.08, обучение пустое. Сброс не теряет прогресс. "
        "Тип стратегии и оплату не менять: максимум конверсий, за конверсии, цена конверсии. "
        "Директ предупредит «обучение начнётся заново» — это и есть сброс. Подтвердить.",
        bold=True,
    )
    add_numbered(
        doc,
        [
            "Кампания GU → минус-фразы. Добавить столбик из раздела 5 к тем, что уже стоят с 23.08.",
            "Стратегия → счётчики: главный / активный → профиль 103116887. Сайт 94713538 оставить вторым, не удалять.",
            "Цели профиля 103116887: включить WhatsApp 441960276 — 100 ₽ и Telegram 441960277 — 100 ₽. Звонок 110 и кнопка действия 110 не трогать.",
            "Цели сайта (тел 140 · MAX 120 · TG 120 · автоцель 130) можно оставить включёнными, но не делать их главнее профиля.",
            "Если Директ спросит про перезапуск обучения — да. Тип стратегии, оплату и «цена конверсии» не менять.",
            "Группа «Яхты Сочи» 5789218554 → ключевые фразы. Если поле пустое — вставить 10 фраз из раздела 7, соответствие фразовое.",
            "Автотаргет: «Целевые» оставить. «Альтернативные» и «Сопутствующие» — выкл.",
            "Места показа не трогать: только «Список организаций, отели и галерея услуг». Поиск / РСЯ / Карты / товарная / динамические — выкл.",
        ],
    )
    add_p(
        doc,
        "После этих кликов часы 0/10 стартуют заново. Смотреть звонок / WA / TG карточки, не CTR поиска. "
        "Бюджет 5 000 можно оставить: расход 0, поднимать его «для показов» бессмысленно. "
        "8–9 тыс. — только если после сброса стратегия начнёт тратить и упрётся в потолок.",
    )

    add_heading(doc, "7. Десять фраз ядра — фразовое", 1)
    add_p(doc, "В группу, если ключевых ещё нет. По одной строке. Соответствие — фразовое.")
    add_column(doc, KEYS)
    add_p(doc, "Не минусовать: адлер / аренда / сириус / катер / яхта / прогулка.")

    add_heading(doc, "8. Не делать", 1)
    add_bullets(
        doc,
        [
            "Не включать Поиск / РСЯ / Карты «чтобы набрать показы». Это чужое место, не GU.",
            "Не включать галерею на SR / MP / ZK / GP / TG / RB.",
            "Не менять тип стратегии и оплату (макс. конв. · за конверсии · цена конверсии).",
            "Не менять телефон в объявлении, пока пользователь сам не попросит (остаётся +7 918).",
            "Не минусовать вип / 1800 / адлер / сириус / аренда.",
            "Не судить GU по CTR поиска — конверсий карточки нет, выборка 9 показов.",
        ],
    )

    add_heading(doc, "9. Следующая выгрузка", 1)
    add_p(
        doc,
        "Сб–вс 29–30.08 — как обещали, в воскресенье. Потом полная неделя после сброса "
        "(запросы · дни · скрин стратегии: какой счётчик главный, какие цели профиля вкл). "
        "Номер кампании в шапке, если появится.",
    )
    add_p(
        doc,
        "Handoff: cursor/Galereya_uslug_agent_handoff.md · "
        "агент https://cursor.com/agents/bc-38eae2c0-267b-4d34-ba6c-f51bef4259f5",
        size=10,
    )

    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build()
