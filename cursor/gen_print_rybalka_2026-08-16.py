#!/usr/bin/env python3
"""Инструкция печати для локального агента — срез рыбалки 16.08."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTS = [
    Path(__file__).resolve().parent / "00_PRINT_Rybalka_konkurenty_2026-08-16.docx",
    Path(__file__).resolve().parent.parent / "локальная" / "00_PRINT_Rybalka_konkurenty_2026-08-16.docx",
]
NAVY = RGBColor(0x0B, 0x3D, 0x5C)


def font(run, size=12, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add(doc, text, *, size=12, bold=False, center=False, after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=NAVY if bold else None)
    p.paragraph_format.space_after = Pt(after)
    return p


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    add(doc, "ЛОКАЛЬНОМУ АГЕНТУ — НА ПЕЧАТЬ", size=16, bold=True, center=True)
    add(doc, "Групповая рыбалка · срез конкурентов · 16.08.2026", size=13, bold=True, center=True)
    add(doc, "Односторонняя печать. Не двусторонняя.", size=12, bold=True)

    add(doc, "Что печатать", size=14, bold=True)
    add(doc, "Файл: Konkurenty_gruppovaya_rybalka_Sirius_2026-08-16.docx")
    add(doc, "Положить в D:\\CURSOR\\  (или локальная\\ рядом с этим файлом).")

    add(doc, "PowerShell", size=14, bold=True)
    add(
        doc,
        "New-Item -ItemType Directory -Force -Path \"D:\\CURSOR\" | Out-Null\n"
        "$f = \"D:\\CURSOR\\Konkurenty_gruppovaya_rybalka_Sirius_2026-08-16.docx\"\n"
        "Start-Process $f -Verb Print",
        size=11,
    )

    add(doc, "В диалоге принтера", size=14, bold=True)
    add(doc, "• Двусторонняя печать — ВЫКЛ / Print on both sides — None")
    add(doc, "• Формат A4, поля как в файле")
    add(doc, "• Цвет — можно ч/б, таблицы читаются")

    add(doc, "Не печатать файл 15.08 — устарел. Печатать только редакцию 16.08.", bold=True)

    for out in OUTS:
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)
        print(f"Wrote {out}")


if __name__ == "__main__":
    build()
