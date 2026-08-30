#!/usr/bin/env python3
"""Анализ выгрузки RB 24–30.08.2026 → Word."""

import csv
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CSV_PATH = Path(__file__).resolve().parent / "2026-08-30_vitaminki21_zaprosy.csv"
OUT = Path(__file__).resolve().parent / "Analiz_trafika_RB_Popaj_2026-08-30.docx"
OUT_LOCAL = Path(__file__).resolve().parent.parent / "локальная" / OUT.name

NAVY = RGBColor(0x0B, 0x3D, 0x5C)
GRAY = RGBColor(0x33, 0x33, 0x33)
RED = RGBColor(0x8B, 0x1A, 0x1A)

MINUS_ADD = """-лазаревск
-лазаревское
-дивноморск
-дивноморское
-кудепста
-горная рыбалка
-горной рыбалки
-форелев
-форелевое
-форелевая
-падел
-прокат катер
-прокат удоч
-лодку на прокат
-на реке
-мзымта
-щук
-ибиц
-удочки на прокат
-спускать лодку
-кораблик
-деловой рыбалки
-инструктор по рыбной
-золотая рыбка
-копченая
-как подписать
-фото как готовят"""


def num(s):
    if not s or s == "-":
        return 0.0
    return float(str(s).replace(" ", "").replace(",", "."))


def load_rows():
    rows = []
    with open(CSV_PATH, encoding="utf-8-sig") as f:
        for r in csv.DictReader(f):
            if r.get("День") == "Итого":
                continue
            rows.append(r)
    return rows


def agg(rows, key_fn):
    d = defaultdict(lambda: {"cost": 0, "imp": 0, "clk": 0, "conv": 0})
    for r in rows:
        k = key_fn(r)
        d[k]["cost"] += num(r["Расход, ₽"])
        d[k]["imp"] += int(num(r["Показы"]))
        d[k]["clk"] += int(num(r["Клики"]))
        d[k]["conv"] += int(num(r["Конверсии"]))
    return d


def set_run(run, *, bold=False, size=11, color=GRAY, font="Arial"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font


def add_p(doc, text, *, bold=False, size=11, color=GRAY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    p.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run(run, bold=True, size=16 if level == 1 else 13, color=NAVY)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc, items, *, color=GRAY):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        set_run(run, size=11, color=color)


def add_code(doc, text):
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run(run, size=10, font="Consolas")
        p.paragraph_format.space_after = Pt(0)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run(run, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def build():
    rows = load_rows()
    days = sorted(set(r["День"] for r in rows))
    total = {"imp": 0, "clk": 0, "conv": 0}
    for r in rows:
        total["imp"] += int(num(r["Показы"]))
        total["clk"] += int(num(r["Клики"]))
        total["conv"] += int(num(r["Конверсии"]))
    ctr = total["clk"] / total["imp"] * 100 if total["imp"] else 0

    by_group = agg(rows, lambda r: r["Название группы"].strip())
    by_age = agg(rows, lambda r: r.get("Возраст", "?"))
    by_gender = agg(rows, lambda r: r.get("Пол", "?"))
    by_status = agg(rows, lambda r: r.get("Статус объявления", "?"))
    by_place = agg(rows, lambda r: r.get("Тип площадки", "?"))

    g10_fish = sum(
        int(num(r["Показы"]))
        for r in rows
        if "G10" in r["Название группы"]
        and any(x in (r.get("Текст", "") + r.get("Заголовок", "")).lower()
                for x in ["барабул", "ставрид", "пеламид", "сезонн"])
    )
    g10_imp = sum(v["imp"] for k, v in by_group.items() if "G10" in k)

    g1_arch = sum(
        int(num(r["Показы"]))
        for r in rows
        if "G1-G3" in r["Название группы"] and r.get("Статус объявления") == "Архивные"
    )

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Анализ трафика RB · 24–30.08.2026")
    set_run(run, bold=True, size=18, color=NAVY)
    add_p(doc, "ЕПК «Групповая рыбалка» · vitaminki21 · № 713632237 · посадка /gruppovaja_ribbalka", size=10)
    add_p(
        doc,
        f"Период: {days[0]} — {days[-1]} ({len(days)} дн.) · "
        f"Показы {total['imp']:,} · Клики {total['clk']} · Конверсии {total['conv']} · CTR {ctr:.1f}% · "
        "Расход 0 ₽ (оплата за конверсии, конверсий нет).",
        bold=True,
    )

    add_heading(doc, "1. Главный вывод")
    add_bullets(
        doc,
        [
            "Объём вырос (858 показов / 7 дн.), но **0 конверсий** после перезапуска обучения 25.08 — "
            "стратегию / бюджет / CPA **не трогать**.",
            "G10 забрал 44% показов — основной драйвер; G11 запущен (104 п), **G12 в выгрузке нет**.",
            "**Критично:** G1–G3 крутит **архивное** объявление № 1918830392081206022 (225 показов из 231). "
            "Нужны новые активные объявления.",
            "G10: ~78% показов всё ещё с текстами «барабулька/ставрида/пеламида» — дочистить.",
            "G5: схема A **не завершена** — пеламида в заголовках обоих объявлений; G12 не заведена.",
        ],
        color=RED,
    )

    add_heading(doc, "2. Сводка по группам")
    g_rows = []
    for name, v in sorted(by_group.items(), key=lambda x: -x[1]["imp"]):
        c = v["clk"] / v["imp"] * 100 if v["imp"] else 0
        g_rows.append([name[:35], v["imp"], v["clk"], v["conv"], f"{c:.1f}%"])
    add_table(doc, ["Группа", "Показы", "Клики", "Конв.", "CTR"], g_rows, [5.5, 2, 2, 1.5, 1.5])

    add_heading(doc, "3. Динамика по дням")
    by_day = agg(rows, lambda r: r["День"])
    d_rows = []
    for d in sorted(by_day):
        v = by_day[d]
        c = v["clk"] / v["imp"] * 100 if v["imp"] else 0
        d_rows.append([d, v["imp"], v["clk"], f"{c:.1f}%"])
    add_table(doc, ["День", "Показы", "Клики", "CTR"], d_rows, [3, 2.5, 2.5, 2])
    add_p(
        doc,
        "24–27.08: 340 п / 26 к (CTR 7,6%). 28–30.08: 518 п / 16 к (CTR 3,1%) — "
        "рост объёма за счёт G10/G11 и РСЯ, качество кликов просело.",
        size=10,
    )

    add_heading(doc, "4. Площадки и тип условия")
    add_table(
        doc,
        ["Площадка", "Показы", "Клики", "CTR"],
        [
            [k, v["imp"], v["clk"], f"{v['clk']/v['imp']*100:.1f}%"]
            for k, v in sorted(by_place.items(), key=lambda x: -x[1]["imp"])
        ],
        [4, 2.5, 2.5, 2],
    )
    add_p(doc, "G10: Поиск 202 п / 21 к · РСЯ 177 п / 4 к — Поиск рабочий, РСЯ слабый.", size=10)
    add_p(
        doc,
        "G10: «Фраза» 45 п / 7 к (CTR 15,6%) · Автотаргет 212 п / 14 к · "
        "Интересы 122 п / 4 к (CTR 3,3%) — профиль даёт объём, но хуже кликает.",
        size=10,
    )

    add_heading(doc, "5. G10 — что ещё поправить")
    add_bullets(
        doc,
        [
            f"Показы с сезонными текстами (барабулька/ставрида/пеламида): **{g10_fish} из {g10_imp}**.",
            "Лучший заголовок по кликам: «Рыбалка в Чёрном море — Сочи» (103 п / 15 к).",
            "«Морская рыбалка — 3 часа…» — 145 п / только 4 к: слабый, можно убрать или переписать.",
            "Клики с мусором: лазаревское, дивноморское, аренда катеров, экскурсия с рыбалкой.",
            "UTM в объявлении: epk_gruppovaya_rybalka_G10 — ок.",
        ],
    )

    add_heading(doc, "6. G11 луфарь (новая)")
    add_bullets(
        doc,
        [
            "104 показа / 3 клика / 0 конв. · объявление № **1919656966051937203**.",
            "Мусор автотаргета: падел, форелевое хозяйство, light force, аренда катеров.",
            "Часть показов с заголовками не про луфарь («3 часа в море», «осенний сезон») — оставить только луфарь-тексты.",
            "Добавить минусы (см. §8); интересы G10 **не** копировать.",
        ],
    )

    add_heading(doc, "7. G1–G3 · G5 · G9")
    add_bullets(
        doc,
        [
            f"G1–G3: **{g1_arch}** показов на **архивном** объявлении — срочно выпустить новое активное "
            "(без ската, актуальные заголовки).",
            "G5: пеламида всё ещё в заголовках (…262936 и новое …027646) — по схеме A убрать, завести G12.",
            "G5: мусор «копченая барабулька», «как подписать», «щука» — минусовать.",
            "G9: «катамаран», «аренда катеров», «катер прогулочный» — точечные минусы, без −катамаран на кампанию.",
        ],
    )

    add_heading(doc, "8. Минус-фразы (добавить на кампанию или G10+G11+G1)", 1)
    add_code(doc, MINUS_ADD)

    add_heading(doc, "9. Корректировки по возрасту — нужны ли?")
    add_table(
        doc,
        ["Возраст", "Показы", "Клики", "CTR", "Доля показов"],
        [
            [
                k,
                v["imp"],
                v["clk"],
                f"{v['clk']/v['imp']*100:.1f}%" if v["imp"] else "—",
                f"{v['imp']/total['imp']*100:.0f}%",
            ]
            for k, v in sorted(by_age.items(), key=lambda x: -x[1]["imp"])
            if v["imp"] >= 10
        ],
        [4, 2, 2, 2, 2.5],
    )
    add_table(
        doc,
        ["Пол", "Показы", "Клики", "CTR"],
        [
            [k, v["imp"], v["clk"], f"{v['clk']/v['imp']*100:.1f}%"]
            for k, v in sorted(by_gender.items(), key=lambda x: -x[1]["imp"])
        ],
        [4, 2.5, 2.5, 2.5],
    )
    add_p(doc, "Сейчас в кампании по возрасту: только **−100% младше 18**.", bold=True)
    add_bullets(
        doc,
        [
            "**Сейчас (обучение 0/10): корректировки по возрасту НЕ добавлять.** "
            "Нет конверсий — оптимизировать CTR бессмысленно, сломаете обучение.",
            "По CTR: лучше всего **45–54** (8,8%) и **женщины 25–34 / 45–54** (9–11%). "
            "**55+** — 34% показов, CTR 3,1% (много РСЯ).",
            "**Не минусовать женщин** — у них CTR выше мужчин (7,1% vs 3,7%), типично бронируют туры.",
            "**После ≥10 конв./нед:** если 55+ без конверсий — рассмотреть **−10…−15%** только на «Старше 55»; "
            "если 45–54 конвертит — **+5…+10%** на 45–54. Пока данных нет.",
            "Платёжеспособность (+20% топ 1%, +10% 2–5%) и смартфоны (+20%) — **не трогать**.",
        ],
    )

    add_heading(doc, "10. Чеклист действий")
    add_bullets(
        doc,
        [
            "☐ G1–G3: новое **активное** объявление (архивное остановить/не показывать)",
            "☐ G10: убрать тексты барабулька/ставрида/пеламида из всех вариантов",
            "☐ G5: убрать пеламиду · завести **G12** по Word схемы A",
            "☐ G11: только луфарь в заголовках; минусы §8",
            "☐ Минусы §8 на кампанию или G10+G11+G1",
            "☐ Возраст: **ничего не менять** до выхода из обучения",
            "☐ Стратегия / 8000 ₽ / CPA — не трогать",
            "☐ Через 7 дней — новая выгрузка, смотреть конверсии по возрасту",
        ],
    )

    add_heading(doc, "11. Объявления в отчёте")
    add_table(
        doc,
        ["№ объявления", "Группа", "Статус", "Показы"],
        [
            ["1918830392081206022", "G1–G3", "Архивные ⚠", str(g1_arch)],
            ["1919233043117772392", "G10", "Активные", str(by_group.get("G10 - рыбаки", {}).get("imp", 0))],
            ["1918916858360262936", "G5", "Активные", "основное"],
            ["1919869437011027646", "G5", "Активные", "новое, 14 п"],
            ["1919656966051937203", "G11", "Активные", "104"],
            ["1918997239747074414", "G9", "Активные", "22"],
        ],
        [4.5, 3.5, 2.5, 3],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    try:
        OUT_LOCAL.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUT_LOCAL)
    except OSError:
        pass
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
