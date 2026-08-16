#!/usr/bin/env python3
"""Анализ РК «Яхта Tigger» за календарную неделю 10–16.08.2026 → Word на печать."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Analiz_RK_Tigger_nedelya_2026-08-10_16_na_pechat.docx"
NAVY = (0x1F, 0x4E, 0x79)
GRAY = (0x59, 0x59, 0x59)


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_p(doc, text, size=11, bold=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.12
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "666666")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def fill_cell(cell, text, bold=False, size=10, header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold or header, color=(255, 255, 255) if header else None)
    set_cell_border(cell)
    if header:
        shade_cell(cell, "1F4E79")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, header=True, size=10)
    for r_i, row in enumerate(rows):
        prevent_row_split(table.rows[r_i + 1])
        for c_i, val in enumerate(row):
            fill_cell(table.rows[r_i + 1].cells[c_i], val, size=10, bold=c_i == 0)
            if r_i % 2 == 1:
                shade_cell(table.rows[r_i + 1].cells[c_i], "F2F2F2")
    if col_widths:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblW = OxmlElement("w:tblW")
        total = int(sum(col_widths) * 567)  # cm → twips approx
        tblW.set(qn("w:w"), str(total))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*NAVY)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run(run, size=size)


def add_page_field(paragraph):
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    set_run(run, size=9, color=GRAY)


def setup_print(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Яхта Tigger  ·  Яндекс Директ  ·  календарная неделя 10–16.08.2026  ·  на печать")
    set_run(run, size=9, color=GRAY)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Внутренний анализ РК  ·  страница ")
    set_run(run, size=9, color=GRAY)
    add_page_field(fp)
    run = fp.add_run("  ·  файл Analiz_RK_Tigger_nedelya_2026-08-10_16_na_pechat.docx")
    set_run(run, size=9, color=GRAY)


def build():
    doc = Document()
    setup_print(doc)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_p(
        doc,
        "Анализ рекламной кампании «Яхта Tigger»",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_p(
        doc,
        "Календарная неделя Директа: понедельник 10.08 — воскресенье 16.08.2026",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_p(
        doc,
        "Кампания: «Яхта Tigger - Поиск+Карты   05-08-2026»  ·  посадка /yacht_tigger",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    heading(doc, "1. Скрины прочитаны — снимок настроек на 16.08", 1)
    add_p(
        doc,
        "Три экрана Директа сверены с памятью кампании (06.08 и 14.08). Места показа "
        "и тип стратегии не менялись. Изменилась цена цели в профиле и счётчик дней обучения.",
    )
    add_table(
        doc,
        ["Параметр", "14.08 (память)", "Скрин 16.08", "Оценка"],
        [
            ["Стратегия", "Макс. конверсий, за конверсии", "То же, обычная (не пакет)", "Не трогать"],
            ["Бюджет", "10 000 ₽ / неделя", "10 000 ₽ / неделя", "Ок"],
            ["Ограничение", "Цена конверсии", "Цена конверсии", "Ок"],
            ["Обучение", "Идёт, 0/10", "Идёт, 0/10, 3 дня стратегии", "Сброс обучения?"],
            ["Товарная галерея / выдача / динамика", "ВКЛ", "ВКЛ", "Ок"],
            ["Организации / галерея услуг", "ВЫКЛ", "ВЫКЛ", "Не включать"],
            ["РСЯ", "ВЫКЛ", "ВЫКЛ", "Не включать"],
            ["Яндекс Карты", "ВКЛ", "ВКЛ", "Ок, кликов нет"],
            ["Телефон сайт 94713538", "300 ₽", "300 ₽", "Держать"],
            ["Telegram сайт", "250 ₽", "250 ₽", "Держать"],
            ["MAX сайт", "300 ₽", "300 ₽", "Держать"],
            ["Позвонить профиль 103116887", "130 ₽", "210 ₽", "Уже подняли — стоп"],
        ],
        col_widths=[5.2, 4.0, 4.6, 3.8],
    )
    add_p(
        doc,
        "Важно. Кампания запущена 06.08, а на скрине стратегии «идёт 3 дня». "
        "Скорее всего обучение перезапустилось после правки цели профиля 130 → 210 ₽ "
        "(это как раз рекомендованный коридор 200–250). Цену больше не поднимать. "
        "Директ пишет «Не меняйте настройки» — бюджет, места, тип оплаты и CPA на этой "
        "неделе не трогаем. Неделя обучения в Директе календарная: пн–вс.",
        bold=True,
        size=11,
    )

    heading(doc, "2. Короткий вердикт", 1)
    bullet(
        doc,
        "Неделя закрыта нулём: 232 показа / 23 клика / CTR 9,91% / 0 конверсий / расход 0 ₽. "
        "Оплата за конверсии — клики бесплатные, но учат модель.",
    )
    bullet(
        doc,
        "Позиция ~3,1–3,3 и объём трафика ~80% — не «над всеми», но и не провал. "
        "Спецразмещение даёт 16 из 23 кликов. Проблема не в блоке, а в запросах.",
    )
    bullet(
        doc,
        "Главный тормоз: автотаргетинг 104 показа / 5 кликов — чужие яхты, места Сириуса, "
        "«сириус рент». Категория «целевые» врёт (яхта сириус, victor, орион).",
    )
    bullet(
        doc,
        "Что улучшить сейчас: минус-фразы + выключить альтернативные/сопутствующие + "
        "починить заголовок «3 каюты». Не РСЯ, не CPA, не смена стратегии.",
    )
    bullet(
        doc,
        "К концу недели объём падает (пн 28 → вс 19 показов). VIP-ядро есть "
        "(vip, на ночь, на неделю, закат, Azimut), но его мало относительно мусора.",
    )

    heading(doc, "3. Факты недели 10–16.08 vs срез 10–14.08", 1)
    add_table(
        doc,
        ["Метрика", "10–14.08", "Пн–вс 10–16.08", "Комментарий"],
        [
            ["Показы", "175", "232", "+57, из них сб–вс: 41"],
            ["Клики", "16", "23", "+7"],
            ["CTR", "9,14%", "9,91%", "Объявления кликабельны"],
            ["Конверсии", "0", "0", "Порог обучения 10/нед не взят"],
            ["Расход", "0 ₽", "0 ₽", "Платите только за конверсию"],
            ["Ср. позиция показа", "~3,3", "~3,1–3,3", "Стабильно средняя"],
            ["Ср. объём трафика", "~79%", "80%", "Не максимальный блок"],
            ["Отказы", "~33%", "27%", "Чуть лучше, всё ещё высоко"],
            ["Глубина", "~1,1", "1,14", "Почти сразу уходят"],
        ],
        col_widths=[4.2, 3.2, 4.2, 6.0],
    )

    heading(doc, "3.1. По дням (календарная неделя)", 2)
    add_table(
        doc,
        ["День", "Показы", "Клики", "Конв.", "Заметка"],
        [
            ["Пн 10.08", "28", "5", "0", "Старт недели, 5 кликов"],
            ["Вт 11.08", "59", "3", "0", "Пик показов, слабый CTR"],
            ["Ср 12.08", "53", "4", "0", "Ещё высокий объём"],
            ["Чт 13.08", "30", "2", "0", "Спад"],
            ["Пт 14.08", "21", "3", "0", "Спад"],
            ["Сб 15.08", "22", "5", "0", "Клики: сутки, закат, туры, victor, орион"],
            ["Вс 16.08", "19", "1", "0", "1 клик: аренда на неделю Сириус"],
            ["Итого пн–вс", "232", "23", "0", "Обучение 0/10 за неделю"],
        ],
        col_widths=[3.2, 2.2, 2.0, 2.0, 8.2],
    )

    heading(doc, "3.2. Где крутятся показы", 2)
    add_table(
        doc,
        ["Срез", "Показы", "Клики", "Вывод"],
        [
            ["Поиск", "225", "23", "Весь смысл кампании здесь"],
            ["Сети (в отчёте = Карты)", "7", "0", "Не РСЯ; карты без кликов"],
            ["Спецразмещение", "106", "16", "Клики живут здесь — ок"],
            ["Прочее", "86", "4", "Низкий CTR, тянет среднюю позицию"],
            ["Эксклюзив", "22", "3", "Мало, но кликают"],
            ["Карты (вид)", "16", "0", "Оставить, не ждать лидов"],
            ["Саджест / динамика", "1 / 1", "0 / 0", "Шум"],
            ["Фраза", "128", "18", "Рабочее ядро"],
            ["Автотаргетинг", "104", "5", "45% показов — чистить"],
        ],
        col_widths=[4.4, 2.4, 2.2, 8.6],
    )

    heading(doc, "3.3. Категории запросов Директа", 2)
    add_table(
        doc,
        ["Категория", "Показы", "Клики", "Заметка"],
        [
            ["Целевые", "97", "3", "Ложные: сириус, victor, орион, рент"],
            ["Широкие", "96", "16", "Дают клики: vip / сутки / Azimut / закат"],
            ["Альтернативные", "23", "2", "Выключить. Есть «фуршет наталья»"],
            ["Узкие", "6", "2", "1 полезный: аренда на неделю Сириус"],
            ["Не определено", "7", "0", "Карты без запроса"],
            ["Сопутствующие", "3", "0", "Выключить (азимут бассейн)"],
        ],
        col_widths=[4.0, 2.4, 2.2, 9.0],
    )

    heading(doc, "3.4. Кто кликает", 2)
    add_table(
        doc,
        ["Срез", "Показы / клики", "Комментарий"],
        [
            ["Мужчины", "101 / 14", "Основной клик"],
            ["Женщины", "120 / 9", "Больше показов, слабее CTR"],
            ["25–34 + 35–44", "90+71 / 11+6", "Ядро возраста — не резать"],
            ["Топ 1%", "31 / 3", "Есть VIP: ночь, сутки, закат"],
            ["6–10%", "40 / 6", "Рабочий сегмент"],
            ["2–5%", "50 / 4", "Ок"],
            ["Остальные", "111 / 10", "Много широкого спроса"],
        ],
        col_widths=[4.0, 3.6, 10.0],
    )

    heading(doc, "4. Разбор 23 кликов — что целевое, что мусор", 1)
    add_p(doc, "Оставлять / усиливать (VIP или private charter):", bold=True, space_after=3)
    add_table(
        doc,
        ["Запрос", "Кто", "Отказ", "Вердикт"],
        [
            ["аренда vip яхт сочи", "М 45–54, 6–10%", "0%, глуб. 2", "Эталон"],
            ["снять яхту на ночь с капитаном", "М 25–34, топ 1%", "0%", "VIP ночь"],
            ["аренда … на закате", "М 35–44, топ 1%", "0%", "VIP закат"],
            ["аренда на неделю сириус", "М 35–44, 6–10%", "0%, глуб. 2", "Неделя — ок"],
            ["моторная яхта азимут", "Ж 25–34, 2–5%", "0%", "Бренд верфи"],
            ["снять на неделю", "М 25–34, 2–5%", "50%", "Интент ок"],
            ["аренда сутки (15.08)", "М 25–34, топ 1%", "0%, глуб. 2", "Смотрел страницу"],
            ["аренда сочи морпорт", "М 25–34, 6–10%", "0%", "Гео причала"],
        ],
        col_widths=[6.2, 4.4, 3.4, 3.6],
    )
    add_p(doc, "Минусовать / не кормить автотаргет:", bold=True, space_after=3)
    add_table(
        doc,
        ["Запрос", "Показы нед.", "Почему мусор"],
        [
            ["яхта сириус / сириус л6 / iceni", "11+3+2", "Путаница с катером и чужими судами"],
            ["сириус рент", "7", "Прокат авто, не яхта"],
            ["яхта victor / орион / багира / harmony", "6+1+2+2", "Чужие названия, автотаргет"],
            ["яхта на закате … наталья", "5", "Каннибал с РК «Закат» / контакт Натальи"],
            ["морские прогулки / туры / круиз адлер", "клики 1+1+1", "Групповой дешёвый спрос"],
            ["яхта азимут interphase / 42 / бассейн", "3+3+3", "Другая модель / не чартер"],
            ["для двоих / прогулка индивидуальная", "показы есть", "Тянет 1 800 ₽, не 50 000 ₽/ч"],
            ["скайпарк, подкова, дубай, шоколад, фаворит", "пачками по 1–3", "Места Сириуса, не аренда"],
        ],
        col_widths=[6.6, 2.8, 8.2],
    )
    add_p(
        doc,
        "Не минусовать: tigger / тайгер, azimut (без 42/бассейн/леопард), vip, премиум, "
        "гидроцикл, флайбридж, на неделю, на ночь, закат без «наталья».",
        size=11,
    )

    heading(doc, "5. Что улучшить — приоритет", 1)
    heading(doc, "5.1. Сегодня (не ломает стратегию)", 2)
    add_p(doc, "A. Минус-фразы кампании — вставить пакетом:", bold=True, space_after=3)
    add_p(
        doc,
        "скайпарк  подкова  дубай  восточный квартал  дверь в море  шоколад сириус  "
        "олимпийский парк  каяк  сириус рент  seawolf  cruize  krusie  ritz  маджеста  "
        "екатерина  леопард  iceni  фаворит  мегеллан  victor  орион  багира  harmony  "
        "гармония  багратион  адель  феона  снафу  джиаванитта  внжела  алексум  "
        "interphase  азимут 42  азимут бассейн  для двоих  туры  круиз парк  "
        "самостоятельн  теплоход  морская прогулка  морские прогулки  яхт клуб  наталья",
        size=10,
        space_after=8,
    )
    add_p(
        doc,
        "B. В группе автотаргетинга выключить «Альтернативные» и «Сопутствующие», "
        "если ещё включены. В статистике они есть (23+3 показа). «Целевые» пока оставить, "
        "но минусами закрыть чужие яхты — иначе Директ продолжит считать их целевыми.",
        space_after=6,
    )
    add_p(
        doc,
        "C. Объявления: заголовок «2 палубы, 3 каюты» (33 показа / 5 кликов) — ошибка "
        "против вашей правки «4 каюты». Заменить на «4 каюты» или «3 гостевые + каюта экипажа». "
        "В сильных заголовках держать: Azimut 68 · private charter · от 50 000 ₽/час · от 2 часов · Сириус.",
        space_after=6,
    )
    add_p(
        doc,
        "D. Посадка /yacht_tigger: глубина 1,14. В первом экране цена 50 000 ₽/час и "
        "«яхта целиком, не групповой билет» — чтобы отсечь прогулки до клика по телефону. "
        "Макет премиум-лендинга уже есть в репо (cursor/landings/tigger).",
    )

    heading(doc, "5.2. Не делать на неделе 17–23.08", 2)
    bullet(doc, "Не поднимать CPA (телефон 300 / TG 250 / MAX 300 / профиль 210). Профиль уже поднят.")
    bullet(doc, "Не включать РСЯ и «список организаций / галерею услуг» ради 10 конверсий в обучении.")
    bullet(doc, "Не менять тип стратегии и оплату. Успех = VIP-лид, не зелёный статус 10/нед.")
    bullet(doc, "Не сужать гео только на Сочи — VIP часто ищет из Москвы и Петербурга.")
    bullet(doc, "Не минусовать «гидроцикл» и «азимут» целиком.")

    heading(doc, "5.3. Позиции показа — реалистичный рычаг", 2)
    add_p(
        doc,
        "В ЕПК «за конверсии» позиция не ставится ставкой. Сейчас 16/23 клика уже из спецразмещения. "
        "Средняя 3,1 — из-за «прочего» и мусорных запросов, где аукцион другой. Чистка запросов "
        "сдвинет показы в ядро (vip / azimut / сутки с капитаном / неделя), где блок выше. "
        "Следующий рычаг цены конверсии — только если после минусов на неделе 17–23.08 "
        "VIP-показов мало и позиция по ядру хуже 3–4. Тогда телефон/MAX 350–400, Telegram 300 — "
        "одним изменением, не раньше.",
    )

    heading(doc, "6. План на календарную неделю 17–23.08", 1)
    add_table(
        doc,
        ["Когда", "Действие"],
        [
            ["Пн 17.08, утро", "Минусы из §5.1 + выключить альтернативные/сопутствующие"],
            ["Пн 17.08", "Поправить «3 каюты» → 4 каюты; не трогать CPA/бюджет/места"],
            ["Пн–вс", "Каждый день: поиск по запросам → нерелевант в минус кампании"],
            ["Посадка", "Цена и «private charter» в первом экране /yacht_tigger"],
            ["Вс 23.08", "Новая выгрузка пн–вс. Смотрим: 0 конв. или нет, доля автотаргета, ядро VIP"],
            ["После 23.08", "Если ядро чистое и объёма нет — тогда +15–20% к ценам сайта"],
        ],
        col_widths=[4.0, 13.6],
    )

    heading(doc, "7. Итог на вопросы", 1)
    add_table(
        doc,
        ["Вопрос", "Ответ"],
        [
            [
                "Скрины прочитаны?",
                "Да. Места ок. Стратегия в обучении 0/10, 3 дня. Профиль «Позвонить» 210 ₽ (было 130).",
            ],
            [
                "Что улучшить?",
                "Минусы + автотаргет + заголовок 4 каюты + первый экран посадки. Не CPA и не РСЯ.",
            ],
            [
                "Почему 0 конверсий при CTR 10%?",
                "Клики с «прогулок», чужих яхт и суток без понимания 50 000 ₽/ч. Модель учится на этом.",
            ],
            [
                "Главный риск",
                "Ещё одна правка стратегии/цен на старте новой недели снова обнулит 3 дня обучения.",
            ],
        ],
        col_widths=[4.2, 13.4],
    )

    heading(doc, "8. Источники", 1)
    bullet(doc, "Скрины 16.08: места показа (ручная настройка), стратегия/обучение, цены целей.")
    bullet(doc, "CSV: 2026-08-16_19-00-11_vitaminki21.csv — запросы 10–16.08.2026, 217 строк + итого.")
    bullet(doc, "Сравнение: срез 10–14.08 в Analiz_RK_Tigger_pozicii_CPA_2026-08-14.docx.")
    bullet(doc, "Память: cursor/Tigger_campaign_memory.md.")

    add_p(
        doc,
        "Документ для печати, A4, внутренний. Не для клиента.",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
        space_after=0,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
