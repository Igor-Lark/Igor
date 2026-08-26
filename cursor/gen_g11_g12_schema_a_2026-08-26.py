#!/usr/bin/env python3
"""G11+G12 схема A: луфарь и пеламида — отдельные группы + чистка G5/G10."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "G11_G12_Lufar_Pelamida_schema_A_Popaj_2026-08-26.docx"
OUT_LOCAL = Path(__file__).resolve().parent.parent / "локальная" / OUT.name

LANDING = "https://boat-sochi.ru/gruppovaja_ribalka"
PHONE = "+7 918 304-40-00"
PRICE = "2 800 ₽/чел."
DURATION = "3 часа"
BOAT = "Моряк Попай"
CAMPAIGN_NO = "713632237"
NAVY = RGBColor(0x0B, 0x3D, 0x5C)
GRAY = RGBColor(0x33, 0x33, 0x33)
RED = RGBColor(0x8B, 0x1A, 0x1A)

MINUS_FISH_GROUPS = """-катран
-скат
-трофейная
-трофей
-целиком
-индивидуальная
-аренда катамарана
-20000
-20 000
-барабулька
-султанка
-ставрида
-!с берега
-!с пирса
-!с буны
-форель
-пруд
-прудовая
-абхазия
-крым
-севастополь
-анапа
-геленджик
-дельфины
-дельфинарий
-парусная
-гидроцикл
-теплоход
-круиз
-купить яхту
-магазин
-базар
-брюки
-женская
-кафе
-кура
-курка
-мужская
-птица
-ресторан
-рынок
-столовая
-штаны
-вакансия
-бесплатно"""

# На G11 дополнительно минус пеламида; на G12 — минус луфарь
MINUS_G11_EXTRA = """-пеламида
-пеламиду
-пеламиды"""

MINUS_G12_EXTRA = """-луфарь
-луфаря
-луфарю"""

# Что ДОБАВИТЬ минусами на G10 (и желательно G1–G3)
MINUS_ADD_G10 = """-луфарь
-луфаря
-пеламида
-пеламиду
-пеламиды"""

# Что УДАЛИТЬ из ключей G5
KEYS_REMOVE_G5 = """морская рыбалка пеламида
рыбалка на пеламиду сириус
рыбалка на пеламиду сочи
пеламида сириус -рыбалка
пеламида сочи -ловля -рыбалка
ловля пеламиды в сочи
вставки на морскую рыбалку пеламида"""

KEYS_G11 = """рыбалка на луфаря сочи
рыбалка на луфаря сириус
луфарь сочи
луфарь сириус -рыбалка
ловля луфаря в сочи
морская рыбалка луфарь
луфарь черное море сочи
рыбалка луфарь адлер
осенняя рыбалка луфарь сочи"""

KEYS_G12 = """морская рыбалка пеламида
рыбалка на пеламиду сириус
рыбалка на пеламиду сочи
пеламида сириус -рыбалка
пеламида сочи -ловля -рыбалка
ловля пеламиды в сочи
рыбалка пеламида адлер
вставки на морскую рыбалку пеламида"""


def set_run(run, *, bold=False, size=11, color=GRAY, font="Arial"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


def add_p(doc, text, *, bold=False, size=11, color=GRAY, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, bold=bold, size=size, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, bold=True, size=16 if level == 1 else 13, color=NAVY)
    return p


def add_bullets(doc, items, *, color=GRAY):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        set_run(run, size=11, color=color)
        p.paragraph_format.space_after = Pt(2)


def add_code_block(doc, text):
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        set_run(run, size=10, font="Consolas")


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run(run, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, "0B3D5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run(run, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def ad_block(doc, title, headlines, texts, creatives):
    add_heading(doc, title, 2)
    add_p(doc, "Заголовки (7):", bold=True, size=10)
    add_bullets(doc, headlines)
    add_p(doc, "Тексты (3):", bold=True, size=10)
    add_bullets(doc, texts)
    add_p(doc, "Креативы:", bold=True, size=10)
    add_bullets(doc, creatives)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    run = title.add_run("G11 + G12 · Схема A — луфарь и пеламида отдельно")
    set_run(run, bold=True, size=18, color=NAVY)

    add_p(
        doc,
        f"РК ЕПК № {CAMPAIGN_NO} · {PRICE} · {DURATION} · «{BOAT}» · {LANDING}",
        size=10,
    )
    add_p(doc, f"Бронь: {PHONE} · линия 2 · Сириус / Имеретинка", size=10)
    add_p(
        doc,
        "Схема A: две группы — жёсткая связка «рыба → текст + креатив». "
        "Стратегию / бюджет 8 000 / цены целей НЕ трогать (обучение).",
        bold=True,
    )

    add_heading(doc, "0. Порядок работ (сделать в этом порядке)", 1)
    add_bullets(
        doc,
        [
            "1) Почистить G5 — убрать пеламиду (ключи + заголовки + баннер).",
            "2) На G10 добавить минусы луфарь/пеламида (и при желании на G1–G3).",
            "3) Создать G11 — только луфарь (ключи + 1 объявление + фото луфаря).",
            "4) Создать G12 — только пеламида (ключи + 1 объявление + фото пеламиды).",
            "5) Записать номера групп/объявлений → прислать агенту.",
        ],
        color=RED,
    )

    # —— G5 ——
    add_heading(doc, "1. Что убрать из G5 (обязательно)", 1)
    add_p(
        doc,
        "G5 остаётся «лето-бархат»: барабулька / ставрида / султанка. "
        "Пеламида переезжает в G12 — иначе каннибализация.",
        bold=True,
    )
    add_p(doc, "Удалить из ключей G5 эти фразы:", bold=True, size=10)
    add_code_block(doc, KEYS_REMOVE_G5)
    add_p(doc, "В объявлении G5 убрать / заменить:", bold=True, size=10)
    add_bullets(
        doc,
        [
            "Заголовки со словом «пеламида» / «сезонный клёв на пеламиду».",
            "Тексты, где пеламида в одном ряду с барабулькой.",
            "Баннер «ставрида / барабулька / пеламида» → сделать баннер только "
            "барабулька + ставрида (или нейтральный улов).",
            "Картинку «пеламида в руке» — перенести в G12, из G5 убрать.",
        ],
        color=RED,
    )
    add_p(doc, "В минусы G5 добавить (чтобы не перехватывала G12):", bold=True, size=10)
    add_code_block(doc, MINUS_ADD_G10)  # same fish minuses
    add_p(
        doc,
        "Оставить в G5: ключи барабулька/ставрида/султанка, оффер 2 800, посадка та же.",
        size=10,
    )

    # —— G10 ——
    add_heading(doc, "2. Что сделать с G10 (обязательно)", 1)
    add_p(
        doc,
        "G10 = широкий «морская рыбалка Сочи» без вида рыбы. "
        "Ключи вида рыбы в G10 не добавлять. Минусами отсечь луфаря и пеламиду.",
        bold=True,
    )
    add_p(doc, "Добавить в минус-фразы группы G10 (№ 5791489951):", bold=True, size=10)
    add_code_block(doc, MINUS_ADD_G10)
    add_p(doc, "Проверить и убрать, если вдруг появилось:", bold=True, size=10)
    add_bullets(
        doc,
        [
            "Любые ключи с «луфарь», «пеламида», «барабулька», «ставрида» — это не G10.",
            "Заголовки/тексты с названиями рыб — вернуть к «морская рыбалка Сочи / море / 2800» "
            "(как в Word G10 от 20.08).",
            "Интересы/профиль на G10 не трогать (уже настроено).",
        ],
        color=RED,
    )
    add_p(
        doc,
        "Желательно те же минусы −луфарь −пеламида добавить на G1–G3 "
        "(группа 5788984305), чтобы ядро не перехватывало сезонные запросы.",
        size=10,
    )

    # —— Structure ——
    add_heading(doc, "3. Итоговая структура рыбных групп", 1)
    add_table(
        doc,
        ["Группа", "Рыба", "Ключи", "Объявление / креатив"],
        [
            ["G5", "барабулька, ставрида", "без пеламиды", "летний баннер без пеламиды"],
            ["G11 · NEW", "только луфарь", "блок §4.1", "1 Ad · фото только луфарь"],
            ["G12 · NEW", "только пеламида", "блок §5.1", "1 Ad · фото только пеламида"],
            ["G10", "без вида рыбы", "широкий Сочи", "заголовки Сочи/море; минусы рыб"],
            ["G9", "катран/скат", "трофей", "не трогать"],
        ],
        col_widths=[3.5, 4, 4.5, 5.5],
    )

    # —— G11 ——
    add_heading(doc, "4. Группа G11 — луфарь", 1)
    add_bullets(
        doc,
        [
            "Имя: G11 - луфарь",
            "Гео: Россия · сценарий: вся заинтересованная аудитория",
            "Профиль пользователя: ПУСТО",
            "Автотаргет: целевые/узкие/широкие ВКЛ; сопутствующие/альтернативные ВЫКЛ",
            "Бренды: свой ВКЛ, конкуренты ВЫКЛ, без бренда ВКЛ",
            "РСЯ: офферный ретаргетинг ВКЛ",
            "UTM: utm_source=yandex&utm_medium=cpc&utm_campaign=epk_gruppovaya_rybalka_popaj"
            "&utm_content=G11_lufar&utm_term={keyword}",
            "Корректировки: «Просмотр сайта» +10%; ливень/снегопад −50%",
            "Посадка: " + LANDING,
        ],
    )
    add_heading(doc, "4.1. Ключи G11", 2)
    add_code_block(doc, KEYS_G11)
    add_heading(doc, "4.2. Минусы G11", 2)
    add_code_block(doc, MINUS_FISH_GROUPS + "\n" + MINUS_G11_EXTRA)
    ad_block(
        doc,
        "4.3. Объявление G11 (одно, только луфарь)",
        [
            "Рыбалка на луфаря в Сочи — 2 800 ₽/чел.",
            f"Луфарь · катамаран «{BOAT}»",
            "Ловля луфаря · Чёрное море · Сочи",
            "3 часа в море · снасти включены",
            "Группа до 8 · Сириус / линия 2",
            "Слоты каждый день · бронь",
            f"Бронь: {PHONE}",
        ],
        [
            f"Рыбалка на луфаря в Сочи: «{BOAT}», билет {PRICE}, {DURATION}. "
            f"Снасти на борту. Сириус, линия 2. Улов не гарантируем.",
            f"Ищете луфаря в Сочи? Место в группе до 8, не весь катер. {PRICE}. Бронь {PHONE}.",
            f"Выход на луфаря из Имеретинского порта. Групповая рыбалка {PRICE}. {LANDING}",
        ],
        [
            "Фото: луфарь в руках / на палубе (крупный план).",
            "Фото: луфарь + катамаран / спиннинг с кормы.",
            "Опц. баннер: «ЛУФАРЬ · Сочи · 2800» — без других рыб.",
            "Форматы 1:1, 4:3, 16:9. В это объявление НЕ грузить пеламиду.",
        ],
    )

    # —— G12 ——
    add_heading(doc, "5. Группа G12 — пеламида", 1)
    add_bullets(
        doc,
        [
            "Имя: G12 - пеламида",
            "Настройки группы — как у G11 (гео, автотаргет, профиль пустой, РСЯ).",
            "UTM: …&utm_content=G12_pelamida&utm_term={keyword}",
            "Корректировки: те же +10% просмотр сайта; −50% ливень/снегопад",
            "Посадка: " + LANDING,
        ],
    )
    add_heading(doc, "5.1. Ключи G12", 2)
    add_code_block(doc, KEYS_G12)
    add_heading(doc, "5.2. Минусы G12", 2)
    add_code_block(doc, MINUS_FISH_GROUPS + "\n" + MINUS_G12_EXTRA)
    ad_block(
        doc,
        "5.3. Объявление G12 (одно, только пеламида)",
        [
            "Рыбалка на пеламиду в Сочи — 2 800 ₽/чел.",
            f"Пеламида · катамаран «{BOAT}»",
            "Ловля пеламиды · Чёрное море · Сочи",
            "3 часа в море · снасти включены",
            "Группа до 8 · Сириус / линия 2",
            "Слоты каждый день · бронь",
            f"Бронь: {PHONE}",
        ],
        [
            f"Рыбалка на пеламиду в Сочи: «{BOAT}», билет {PRICE}, {DURATION}. "
            f"Снасти на борту. Сириус, линия 2. Улов не гарантируем.",
            f"Ищете пеламиду в Сочи? Место в группе до 8. {PRICE}. Бронь {PHONE}.",
            f"Выход на пеламиду из Имеретинского порта. Групповая рыбалка {PRICE}. {LANDING}",
        ],
        [
            "Фото: пеламида в руках (кадр из бывшего G5 — перенести сюда).",
            "Фото: пеламида / улов на палубе или в тарелке.",
            "Опц. баннер: «ПЕЛАМИДА · Сочи · 2800» — без луфаря/барабульки.",
            "В это объявление НЕ грузить луфаря и летний микс-баннер.",
        ],
    )

    add_heading(doc, "6. Быстрые ссылки и уточнения", 1)
    add_p(
        doc,
        "Брать общие кампании (2800 / слоты / Попай / линия 2 / бронь). "
        "На группу отдельно не дублировать.",
        size=10,
    )

    add_heading(doc, "7. Чеклист перед сохранением", 1)
    add_bullets(
        doc,
        [
            "☐ G5: ключи пеламиды удалены; заголовки/баннер без пеламиды; минусы −пеламида −луфарь",
            "☐ G10: минусы −луфарь −пеламида добавлены; нет ключей/заголовков с рыбами",
            "☐ G1–G3 (желательно): те же минусы −луфарь −пеламида",
            "☐ G11: только ключи луфаря + 1 Ad + фото луфаря + минус −пеламида",
            "☐ G12: только ключи пеламиды + 1 Ad + фото пеламиды + минус −луфарь",
            "☐ Посадка везде /gruppovaja_ribalka · цена только 2 800",
            "☐ Стратегия / 8000 ₽ / CPA целей не менялись",
            "☐ Номера G11/G12 и объявлений записаны",
        ],
    )

    add_heading(doc, "8. Копипаст: минусы на G10 одной вставкой", 1)
    add_code_block(doc, MINUS_ADD_G10)

    add_heading(doc, "9. Зачем так", 1)
    add_p(
        doc,
        "Ключи в Директе — на группе. Отдельная группа = запрос «луфарь» "
        "не может вызвать объявление с пеламидой. Это и есть схема A.",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    try:
        OUT_LOCAL.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUT_LOCAL)
    except OSError:
        pass
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
