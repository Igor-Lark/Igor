#!/usr/bin/env python3
"""Sirius 704503370 slice 24–26.08.2026 morning -> print Word."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path("/workspace/cursor/Analiz_RK_Sirius_dva_dnya_2026-08-24_26_na_pechat.docx")
NAVY = RGBColor(0x1A, 0x47, 0x7A)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x7A, 0x46)
ORANGE = RGBColor(0xC0, 0x6A, 0x00)


def shade(cell, color="D9E2F3"):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(el)


def set_run(run, *, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
        r.font.name = "Calibri"
    return h


def para(doc, text, *, bold=False, size=11, space=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run(r, bold=bold, size=size, color=color)
    return p


def mixed(doc, parts, *, size=11, space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1.08
    for text, kwargs in parts:
        r = p.add_run(text)
        set_run(r, size=size, **kwargs)
    return p


def bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, size=11)
        r = p.add_run(text)
        set_run(r, size=11)
    else:
        r = p.add_run(text)
        set_run(r, size=11)
    return p


def add_table(doc, headers, rows, col_w=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(cell, "1A477A")
    for ri, row in enumerate(rows):
        bg = "F2F6FB" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cell = t.rows[ri + 1].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run(r, size=8)
            shade(cell, bg)
    if col_w:
        for row in t.rows:
            for i, w in enumerate(col_w):
                row.cells[i].width = Cm(w)
    return t


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.3)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("СИРИУС — ПРОКАТ КАТЕРОВ")
    set_run(r, bold=True, size=16, color=NAVY)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Срез 24–25.08 + утро 26.08.2026")
    set_run(r, bold=True, size=13, color=NAVY)

    para(
        doc,
        "РК 704503370 · CSV 2026-08-26_08-51-17_vitaminki21.csv. "
        "26.08 — до 08:51. Автотаргет в группах: узкие + целевые + без бренда; в одной группе ещё с брендом.",
        size=10,
    )
    mixed(
        doc,
        [
            ("Вердикт. ", {"bold": True, "color": NAVY}),
            (
                "Автотаргет в таком узком виде оставлять: он дал 7 из 11 конверсий, запросы живые. "
                "Сноуборд и «куда сходить» — не автотаргет, а фраза «отдых в сочи», которую ставили на паузу 21.08. "
                "Она жива. Удалить сегодня, не пауза.",
                {},
            ),
        ],
        space=10,
    )

    heading(doc, "1. Цифры", 1)
    add_table(
        doc,
        ["Срез", "Расход", "Пок.", "Кл.", "Конв.", "CTR %", "CPA", "Поз.", "Объём %"],
        [
            ["24.08", "698", "197", "17", "5", "8,6", "140", "3,00", "45,5"],
            ["25.08", "821", "195", "25", "5", "12,8", "164", "3,16", "54,3"],
            ["26.08 до 08:51", "193", "43", "5", "1", "11,6", "193", "2,58", "37,7"],
            ["Итого в файле", "1 712", "435", "47", "11", "10,8", "156", "3,03", "48,7"],
            ["Только 24–25", "1 519", "392", "42", "10", "10,7", "152", "—", "—"],
        ],
        col_w=[3.2, 1.8, 1.5, 1.3, 1.5, 1.5, 1.5, 1.5, 1.6],
    )
    para(
        doc,
        "Спец: позиция 2,66 · 297 показов · 41 клик · 9 конв. · CTR 13,8 % · CPA 142. "
        "Прочие места: 77 показов, позиция 5,18, 0 конв. Объём 49 % — как на прошлой неделе, не из‑за автотаргета.",
        size=10,
        space=8,
    )

    heading(doc, "2. Автотаргет — не выключать", 1)
    para(
        doc,
        "В отчёте Директа «широкие / сопутствующие / альтернативные» — это класс поискового запроса. "
        "Галки автотаргета (узкие, целевые, без бренда) — другое. Фраза «отдых в сочи» матчит аквапарк и сноуборд, "
        "даже если альтернативные в автотаргете выкл.",
        size=10,
    )
    add_table(
        doc,
        ["Источник", "Расход", "Пок.", "Кл.", "Конв.", "CPA", "Качество"],
        [
            ["Автотаргет", "927", "238", "17", "7", "132", "живые катер/яхта/Сириус"],
            ["Фразы, кроме «отдых в сочи»", "202", "90", "16", "1", "202", "«аренда яхты сириус»"],
            ["Фраза «отдых в сочи»", "583", "107", "14", "3", "194", "сноуборд, куда сходить, мусор"],
        ],
        col_w=[5.0, 1.7, 1.5, 1.3, 1.5, 1.5, 4.3],
    )
    para(doc, "", space=4)
    bullet(doc, "узкие + целевые + без бренда во всех группах — держать.")
    bullet(doc, "бренд в одной группе — ок. В срезе «бренд конкурентов» = 2 показа, 0 кликов. Не включать во все группы.")
    bullet(doc, "широкие / сопутствующие / альтернативные автотаргета — не включать.")
    bullet(doc, "Автотаргет кормит: «морская прогулка сириус», «прогулка на катере в сириусе», «индивидуальные прогулки адлер».")

    heading(doc, "3. Сегодняшний мусор — фраза, не автотаргет", 1)
    add_table(
        doc,
        ["Запрос 26.08", "₽", "Кл.", "Конв.", "Кат.", "Фраза"],
        [
            ["куда сходить в сириусе сочи", "193", "1", "1", "сопутств.", "отдых в сочи"],
            ["прокат сноубордов в сочи", "0", "1", "0*", "альтернат.", "отдых в сочи"],
        ],
        col_w=[5.8, 1.5, 1.3, 1.5, 2.2, 3.5],
    )
    para(
        doc,
        "*В выгрузке до 08:51 конверсия по сноуборду не списалась (оплата за конверсии). Клик уже есть. "
        "«Отдых в сочи» с 24.08: 107 показов — боулинг, аквапарк, квадроциклы, тарзанка, каток, пикник, банан, парашют. "
        "Пауза 21.08 не сработала или фразу вернули. Удалить.",
        size=10,
        space=8,
    )

    heading(doc, "4. Конверсии 24–26", 1)
    add_table(
        doc,
        ["Дата", "Запрос", "₽", "К", "Источник", "Качество"],
        [
            ["24", "морская прогулка сириус", "330", "2", "авто", "живая"],
            ["24", "прогулка на яхте адлер", "170", "1", "авто", "живая"],
            ["24", "прогулки в море сириус", "198", "1", "отдых в сочи", "живая, случайно"],
            ["24", "прогулка на катере сириус", "0", "1", "авто", "живая"],
            ["25", "прогулка на катере в сириусе", "198", "2", "авто", "живая"],
            ["25", "сириус аренда яхты", "202", "1", "фраза", "живая"],
            ["25", "индивидуальные прогулки адлер", "229", "1", "авто", "живая"],
            ["25", "аренда яхты в сочи", "193", "1", "отдых в сочи", "широко"],
            ["26", "куда сходить в сириусе сочи", "193", "1", "отдых в сочи", "МУСОР"],
        ],
        col_w=[1.4, 5.6, 1.4, 1.0, 2.8, 2.6],
    )
    para(doc, "Живых 8–9 из 11. Автотаргет чище фразы «отдых в сочи».", size=10, space=8)

    heading(doc, "5. Сделать сегодня", 1)
    mixed(
        doc,
        [
            ("Удалить фразу: ", {"bold": True, "color": RED}),
            ("отдых в сочи — в группе Россия. Не пауза.", {}),
        ],
    )
    para(
        doc,
        "-сноуборд -сноубор -горнолыж -боулинг -аквапарк -квадроцикл -тарзанк "
        "-канатн -каток -пикник -самокат -банан -парашют -вертолёт -куда сходить",
        bold=True,
        size=12,
        color=RED,
    )
    bullet(doc, "Автотаргет, цели 25.08, 12 000 ₽, РСЯ — не трогать.")
    para(
        doc,
        "Файл: cursor/Analiz_RK_Sirius_dva_dnya_2026-08-24_26_na_pechat.docx. Очередь печати.",
        size=9,
        color=NAVY,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
